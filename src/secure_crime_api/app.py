from __future__ import annotations

import csv
import base64
import hashlib
import html
import io
import json
import re
import secrets
import shutil
import sqlite3
import textwrap
import uuid
import zipfile
from datetime import datetime, timedelta, timezone
from pathlib import Path
from time import perf_counter
from typing import Annotated, Any
from xml.etree import ElementTree

import jwt
from cryptography.exceptions import InvalidSignature
from cryptography.hazmat.primitives.asymmetric import ec, padding, rsa
from cryptography.hazmat.primitives.hashes import SHA256
from fastapi import Depends, FastAPI, File, Form, HTTPException, Request, Response, UploadFile, status
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, RedirectResponse
from fastapi.security import HTTPAuthorizationCredentials, HTTPBearer
from fastapi.staticfiles import StaticFiles
from starlette.middleware.trustedhost import TrustedHostMiddleware

from secure_crime_api.analytics import (
    SAFEGUARDS,
    answer_query,
    build_decision_support,
    discover_crime_patterns,
    network_graph,
    trend_summary,
    visible_cases,
)
from secure_crime_api.audit import verify_audit_chain, write_audit_log
from secure_crime_api.config import Settings, get_settings
from secure_crime_api.crime_data_import import import_crime_csv
from secure_crime_api.crime_ml import advanced_crime_analytics
from secure_crime_api.demo_data import seed_demo_data
from secure_crime_api.kannada_lexicon import dictionary_status, normalize_kannada_query
from secure_crime_api.intent_taxonomy import classify_master_intent, taxonomy_status
from secure_crime_api.nlp_intelligence import answer_incident_intelligence_query
from secure_crime_api.models import (
    AdvancedCrimeAnalyticsResponse,
    AuditLog,
    AuthenticatedUser,
    CaseNote,
    CaseNoteCreate,
    CaseImportRequest,
    CaseCreate,
    CaseRecord,
    CaseSearchResponse,
    CaseStatusUpdate,
    ConversationCreate,
    ConversationExchange,
    ConversationMessage,
    ConversationMessageRequest,
    ConversationRecord,
    CrimeDataImportRequest,
    CrimeDataImportResponse,
    CrimeDataStatusResponse,
    DecisionSupportResponse,
    ExplanationResponse,
    FileUploadRecord,
    FinancialAnalysisResponse,
    FinancialTransactionImportRequest,
    ForecastResponse,
    HealthResponse,
    IntelligenceQuery,
    IntelligenceResponse,
    LoginRequest,
    ModuleStatus,
    NetworkGraph,
    RateLimitAlert,
    SociologicalInsights,
    SuspectProfileResponse,
    TokenResponse,
    TrendSummary,
    TranslationRequest,
    TranslationResponse,
    UserCreate,
    UserUpdate,
    PasswordResetRequest,
    PatternAnalyticsResponse,
    UserPublic,
    PenalCodeRecord,
    CrimeLogCreate,
    CrimeLogRecord,
)
from secure_crime_api.pdf_export import render_conversation_pdf
from secure_crime_api.platform_modules import (
    CORE_INTELLIGENCE_TASKS,
    MODULES,
    SOLUTION_FRAMEWORK,
    analyze_financial_transactions,
    explain_case,
    forecast_hotspots,
    sociological_insights,
    suspect_profile,
    translate_text,
)
from secure_crime_api.rate_limit import InMemoryRateLimitMiddleware
from secure_crime_api.rbac import can_access_case, role_has_permission
from secure_crime_api.redaction import mask_case
from secure_crime_api.security import (
    create_access_token,
    decode_access_token,
    hash_password,
    password_needs_rehash,
    verify_password,
)
from secure_crime_api.storage import Database
from secure_crime_api.storage import normalized_search_terms
from secure_crime_api.strict_search import answer_strict_case_search_query


bearer_scheme = HTTPBearer(auto_error=False)


CASE_LOOKUP_HINTS = {
    "accused",
    "case",
    "cases",
    "complainant",
    "fir",
    "history",
    "name",
    "named",
    "names",
    "person",
    "profile",
    "suspect",
    "victim",
}

CASE_LOOKUP_FILLER_TERMS = {
    "about",
    "accused",
    "accuse",
    "acuse",
    "all",
    "any",
    "are",
    "against",
    "called",
    "case",
    "cases",
    "complainant",
    "complaint",
    "copy",
    "crime",
    "crimes",
    "criminal",
    "details",
    "did",
    "does",
    "find",
    "file",
    "filed",
    "fir",
    "for",
    "give",
    "have",
    "has",
    "high",
    "history",
    "in",
    "involve",
    "involved",
    "involves",
    "is",
    "karnataka",
    "list",
    "name",
    "named",
    "names",
    "number",
    "only",
    "person",
    "profile",
    "risk",
    "risky",
    "score",
    "record",
    "records",
    "registered",
    "search",
    "see",
    "show",
    "suspect",
    "there",
    "the",
    "under",
    "victim",
    "was",
    "were",
    "what",
    "where",
    "who",
    "with",
}

PROFILE_INTENT_TERMS = {
    "assessment",
    "assess",
    "behavior",
    "behavioral",
    "criminological",
    "dangerous",
    "habitual",
    "offender",
    "profile",
    "profiling",
    "recidivism",
    "repeat",
    "risk",
    "risky",
    "score",
}

CASE_LOOKUP_BLOCKER_TERMS = {
    "decision",
    "lead",
    "leads",
    "next",
    "similar",
    "step",
    "steps",
    "summarize",
    "summary",
    "timeline",
}

SYSTEM_INTENT_TERMS = {
    "capabilities",
    "capability",
    "created",
    "creator",
    "help",
    "ksp ai",
    "what can you do",
    "who are you",
    "who created",
}

CONVERSATION_INTENT_TERMS = {
    "good afternoon",
    "good morning",
    "good night",
    "hello",
    "hey",
    "hi",
    "hlo",
    "namaskara",
    "thanks",
    "thank you",
    "\u0ca8\u0cae\u0cb8\u0ccd\u0c95\u0cbe\u0cb0",
}

ADMIN_INTENT_TERMS = {
    "access control",
    "admin",
    "audit",
    "block user",
    "delete user",
    "permission",
    "permissions",
    "rate limit",
    "roles",
    "show users",
    "user activity",
    "users",
}

NETWORK_INTENT_TERMS = {
    "associate",
    "associates",
    "connected",
    "connection",
    "connections",
    "graph",
    "link",
    "links",
    "network",
    "relationship",
    "relationships",
}

FINANCIAL_INTENT_TERMS = {
    "account",
    "accounts",
    "bank",
    "financial",
    "ifsc",
    "laundering",
    "money",
    "mule",
    "structuring",
    "transaction",
    "transactions",
    "transfer",
    "transfers",
    "trail",
}

DISTRICT_NAMES = {
    "bagalkote",
    "ballari",
    "belagavi",
    "bengaluru",
    "bengaluru rural",
    "bengaluru south",
    "bengaluru urban",
    "bidar",
    "chamarajanagar",
    "chikkaballapur",
    "chikkamagaluru",
    "chitradurga",
    "dakshina kannada",
    "davanagere",
    "dharwad",
    "gadag",
    "hassan",
    "haveri",
    "kalaburagi",
    "kodagu",
    "kolar",
    "koppal",
    "mandya",
    "mangaluru",
    "mysuru",
    "raichur",
    "shivamogga",
    "tumakuru",
    "udupi",
    "uttara kannada",
    "vijayanagara",
    "vijayapura",
    "yadgiri",
}

PHONE_RE = re.compile(r"(?<!\d)(?:\+?91[-\s]?)?[6-9]\d{9}(?!\d)")
VEHICLE_RE = re.compile(r"\bKA[-\s]?\d{2}[-\s]?[A-Z]{1,2}[-\s]?\d{4}\b", re.IGNORECASE)
ACCOUNT_RE = re.compile(r"\b[A-Z]{2,12}[-_\s]?(?:AC[-_\s]?)?\d{4,}\b", re.IGNORECASE)
DATE_RE = re.compile(
    r"\b(?:\d{1,2}[-/\s](?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*[-/\s]\d{2,4}|\d{1,2}[-/]\d{1,2}[-/]\d{2,4}|\d{4}-\d{2}-\d{2})\b",
    re.IGNORECASE,
)


def get_db(request: Request) -> Database:
    return request.app.state.db


def get_app_settings(request: Request) -> Settings:
    return request.app.state.settings


def public_user_from_row(row: dict) -> UserPublic:
    return UserPublic(
        id=row["id"],
        username=row["username"],
        full_name=row["full_name"],
        role=row["role"],
        district=row["district"],
        is_active=bool(row.get("is_active", True)),
    )


def visible_transactions(
    user: AuthenticatedUser,
    transactions: list[dict],
    cases: list[dict],
) -> list[dict]:
    if user.role in {"super_admin", "supervisor"}:
        return transactions

    visible_case_ids = {case["id"] for case in cases}
    return [
        tx
        for tx in transactions
        if tx["district"] == user.district and (tx["case_id"] is None or tx["case_id"] in visible_case_ids)
    ]


def focused_network_scope(
    *,
    cases: list[dict[str, Any]],
    transactions: list[dict[str, Any]],
    focus: str | None,
    focus_type: str = "auto",
    hops: int = 4,
) -> tuple[list[dict[str, Any]], list[dict[str, Any]], dict[str, Any]]:
    focus_text = str(focus or "").strip()
    if not focus_text:
        return cases, transactions, {
            "active": False,
            "query": "",
            "focus_type": "all",
            "alerts": [],
            "case_count": len(cases),
            "transaction_count": len(transactions),
        }

    inferred_type = infer_network_focus_type(focus_text, focus_type, cases, transactions)
    max_hops = max(1, min(6, int(hops or 4)))
    seed_cases = [case for case in cases if case_matches_network_focus(case, focus_text, inferred_type)]
    seed_case_ids = {case["id"] for case in seed_cases}
    seed_accounts: set[str] = set()
    seed_people: set[str] = set()
    if inferred_type == "account":
        seed_accounts.add(normalize_account(focus_text))
    if inferred_type == "person":
        seed_people.add(focus_text)
    for case in seed_cases:
        for field in ("suspect_name", "victim_name", "complainant_name"):
            value = str(case.get(field) or "").strip()
            if value and value.casefold() not in {"unknown", "n/a", "na"}:
                seed_people.add(value)

    selected_transactions: list[dict[str, Any]] = []
    selected_transaction_ids: set[str] = set()

    def include_transaction(tx: dict[str, Any]) -> bool:
        tx_id = str(tx.get("id") or "")
        if tx_id in selected_transaction_ids:
            return False
        selected_transaction_ids.add(tx_id)
        selected_transactions.append(tx)
        for field in ("source_account", "target_account"):
            normalized = normalize_account(str(tx.get(field) or ""))
            if normalized:
                seed_accounts.add(normalized)
        for field in ("source_account_holder", "target_account_holder"):
            person = str(tx.get(field) or "").strip()
            if person:
                seed_people.add(person)
        if tx.get("case_id"):
            seed_case_ids.add(tx["case_id"])
        return True

    for tx in transactions:
        if transaction_matches_network_focus(tx, focus_text, inferred_type, seed_case_ids, seed_accounts, seed_people):
            include_transaction(tx)

    for _hop in range(max_hops):
        grew = False
        for tx in transactions:
            if str(tx.get("id") or "") in selected_transaction_ids:
                continue
            if transaction_touches_seeds(
                tx,
                seed_case_ids,
                seed_accounts,
                seed_people,
                include_case_seed=inferred_type == "case",
            ):
                grew = include_transaction(tx) or grew
        if not grew:
            break

    scoped_case_map = {case["id"]: case for case in seed_cases}
    for case in cases:
        if case["id"] in seed_case_ids or case_touches_people(case, seed_people):
            scoped_case_map[case["id"]] = case

    scoped_cases = list(scoped_case_map.values())
    linked_firs = [case["fir_number"] for case in scoped_cases]
    alerts = build_network_focus_alerts(
        focus_text=focus_text,
        focus_type=inferred_type,
        cases=scoped_cases,
        transactions=selected_transactions,
        seed_accounts=seed_accounts,
        seed_people=seed_people,
    )
    return scoped_cases, selected_transactions, {
        "active": True,
        "query": focus_text,
        "focus_type": inferred_type,
        "hops": max_hops,
        "case_count": len(scoped_cases),
        "transaction_count": len(selected_transactions),
        "seed_accounts": sorted(account for account in seed_accounts if account)[:25],
        "seed_people": sorted(person for person in seed_people if person)[:25],
        "linked_firs": linked_firs[:25],
        "alerts": alerts,
    }


def infer_network_focus_type(
    focus: str,
    requested_type: str,
    cases: list[dict[str, Any]],
    transactions: list[dict[str, Any]],
) -> str:
    normalized_requested = str(requested_type or "auto").lower().replace("_", "-")
    if normalized_requested in {"case", "fir", "account", "person"}:
        return "case" if normalized_requested == "fir" else normalized_requested
    if any(case_matches_network_focus(case, focus, "case") for case in cases):
        return "case"
    account_norm = normalize_account(focus)
    if looks_like_account_focus(focus) or any(
        account_norm in {
            normalize_account(str(tx.get("source_account") or "")),
            normalize_account(str(tx.get("target_account") or "")),
        }
        for tx in transactions
    ):
        return "account"
    return "person"


def looks_like_account_focus(value: str) -> bool:
    normalized = normalize_account(value)
    if len(normalized) < 5:
        return False
    return any(char.isdigit() for char in normalized)


def case_matches_network_focus(case: dict[str, Any], focus: str, focus_type: str) -> bool:
    normalized_focus = focus.strip().casefold()
    compact_focus = re.sub(r"[^a-z0-9]", "", normalized_focus)
    if not normalized_focus:
        return False
    if focus_type in {"case", "auto"}:
        identifiers = [
            case.get("id"),
            case.get("fir_number"),
            case.get("source_record_id"),
            case.get("case_number"),
        ]
        for identifier in identifiers:
            text = str(identifier or "").casefold()
            compact = re.sub(r"[^a-z0-9]", "", text)
            if text == normalized_focus or compact == compact_focus:
                return True
    if focus_type in {"person", "auto"}:
        return case_touches_people(case, {focus})
    return False


def transaction_matches_network_focus(
    tx: dict[str, Any],
    focus: str,
    focus_type: str,
    seed_case_ids: set[str],
    seed_accounts: set[str],
    seed_people: set[str],
) -> bool:
    if focus_type in {"case", "person"} and tx.get("case_id") in seed_case_ids:
        return True
    if focus_type == "account":
        account_norm = normalize_account(focus)
        accounts = {
            normalize_account(str(tx.get("source_account") or "")),
            normalize_account(str(tx.get("target_account") or "")),
        }
        if account_norm and account_norm in accounts:
            return True
    if focus_type == "person":
        return transaction_touches_people(tx, {focus})
    return transaction_touches_seeds(
        tx,
        seed_case_ids,
        seed_accounts,
        seed_people,
        include_case_seed=focus_type == "case",
    )


def transaction_touches_seeds(
    tx: dict[str, Any],
    seed_case_ids: set[str],
    seed_accounts: set[str],
    seed_people: set[str],
    *,
    include_case_seed: bool = True,
) -> bool:
    if include_case_seed and tx.get("case_id") in seed_case_ids:
        return True
    accounts = {
        normalize_account(str(tx.get("source_account") or "")),
        normalize_account(str(tx.get("target_account") or "")),
    }
    if seed_accounts.intersection(accounts):
        return True
    return transaction_touches_people(tx, seed_people)


def transaction_touches_people(tx: dict[str, Any], people: set[str]) -> bool:
    if not people:
        return False
    holders = [
        str(tx.get("source_account_holder") or "").strip(),
        str(tx.get("target_account_holder") or "").strip(),
    ]
    return any(holder and same_named_person(holder, person) for holder in holders for person in people)


def case_touches_people(case: dict[str, Any], people: set[str]) -> bool:
    if not people:
        return False
    case_people = [
        str(case.get("suspect_name") or "").strip(),
        str(case.get("victim_name") or "").strip(),
        str(case.get("complainant_name") or "").strip(),
    ]
    return any(recorded and same_named_person(recorded, person) for recorded in case_people for person in people)


def build_network_focus_alerts(
    *,
    focus_text: str,
    focus_type: str,
    cases: list[dict[str, Any]],
    transactions: list[dict[str, Any]],
    seed_accounts: set[str],
    seed_people: set[str],
) -> list[dict[str, Any]]:
    alerts: list[dict[str, Any]] = []
    if transactions:
        total_amount = sum(float(tx.get("amount") or 0) for tx in transactions)
        alerts.append(
            {
                "severity": "high" if cases else "medium",
                "title": "Transaction network flagged",
                "message": (
                    f"{len(transactions)} transaction(s) linked to {focus_text}; "
                    f"total accessible amount {format_rupees(total_amount)}."
                ),
            }
        )
    if cases:
        firs = ", ".join(case["fir_number"] for case in cases[:6])
        alerts.append(
            {
                "severity": "high",
                "title": "Criminal database link found",
                "message": f"{focus_type.title()} focus touches {len(cases)} accessible FIR record(s): {firs}.",
            }
        )
    if seed_accounts and transactions:
        alerts.append(
            {
                "severity": "medium",
                "title": "Account chain extracted",
                "message": f"{len(seed_accounts)} account node(s) are in the focused transfer chain.",
            }
        )
    if seed_people:
        alerts.append(
            {
                "severity": "medium",
                "title": "Person/account holder pivots",
                "message": f"{len(seed_people)} account-holder/person pivot(s) were checked against case records.",
            }
        )
    if not alerts:
        alerts.append(
            {
                "severity": "low",
                "title": "No focused network found",
                "message": "No accessible case or transaction record matched this focus value.",
            }
        )
    return alerts


MEMORY_FIELDS = (
    "conversation_id",
    "current_case_id",
    "current_fir",
    "current_suspect",
    "current_victim",
    "current_complainant",
    "current_district",
    "current_network_focus",
    "current_financial_entity",
    "last_module",
    "last_response",
)

KANNADA_RE = re.compile(r"[\u0c80-\u0cff]")
FIR_RE = re.compile(r"\b(?=[A-Z0-9/-]*\d)[A-Z]{2,}[A-Z0-9]*[-/][A-Z0-9][A-Z0-9/-]*\b", re.IGNORECASE)
CONTEXT_FIR_RE = re.compile(
    r"\bFIR\s*(?:NO|NUMBER|#|:)?\s*((?=[A-Z0-9/-]*\d)[A-Z]{2,}[A-Z0-9]*[-/][A-Z0-9][A-Z0-9/-]*)\b",
    re.IGNORECASE,
)

KANNADA_COMPLAINANT_TERMS = (
    "ದೂರು",
    "ದೂರುದಾರ",
    "ದೂರುದಾರರು",
    "ಕಂಪ್ಲೇಂಟ್",
    "ಕಂಪ್ಲೆಂಟ್",
    "ಕಾಂಪ್ಲೇಂಟ್",
    "complaint",
)
KANNADA_WHO_TERMS = ("ಯಾರು", "ಯಾರಿಂದ", "ಯಾರ", "ಕೊಟ್ಟ", "ಕೊಟ್ಟಿದ್ದಾರೆ", "ನೀಡಿದ")
KANNADA_PERSON_ALIASES = {
    "ರವಿಕುಮಾರ್": "Ravi Kumar",
    "ರವಿ ಕುಮಾರ್": "Ravi Kumar",
    "ರವಿಕುಮಾರ": "Ravi Kumar",
    "ರವಿ ಕುಮಾರ": "Ravi Kumar",
    "ರವಿ": "Ravi",
    "ಕುಮಾರ್": "Kumar",
    "ಕುಮಾರ": "Kumar",
}
KANNADA_LOOKUP_STOP_WORDS = {
    "ಅನ್ನು",
    "ಅವರ",
    "ಆಗಿದೆ",
    "ಇದೆ",
    "ಕೊಟ್ಟ",
    "ಕೊಟ್ಟಿದ್ದಾರೆ",
    "ಕಂಪ್ಲೇಂಟ್",
    "ಕಂಪ್ಲೆಂಟ್",
    "ದಾಖಲೆ",
    "ದೂರು",
    "ದೂರುದಾರ",
    "ದೂರುದಾರರು",
    "ನೀಡಿದ",
    "ಮಾಹಿತಿ",
    "ಮೇಲೆ",
    "ಯಾರು",
    "ಯಾರಿಂದ",
    "ಯಾರ",
    "ಸಂಬಂಧಿಸಿದ",
    "ಹೆಸರಿನಲ್ಲಿ",
    "ಹೆಸರಲ್ಲಿ",
}


def detect_query_language(query: str, requested_language: str | None = None) -> str:
    if KANNADA_RE.search(query):
        return "kn"
    return "en"


def prepare_query_for_intent(query: str, language: str) -> str:
    if language != "kn":
        return query
    learned = normalize_kannada_query(query)
    normalized_query = learned["normalized_query"]
    parts = [normalized_query]
    if is_kannada_complainant_query(normalized_query):
        parts.append("complainant complaint who filed complaint against")
    person = extract_kannada_person_name(normalized_query)
    if person:
        parts.append(person)
    return " ".join(parts)


def is_kannada_complainant_query(query: str) -> bool:
    if not KANNADA_RE.search(query):
        return False
    return any(term in query for term in KANNADA_COMPLAINANT_TERMS) and any(
        term in query for term in KANNADA_WHO_TERMS
    )


def extract_kannada_person_name(query: str) -> str | None:
    compact = re.sub(r"\s+", "", query)
    for source, target in KANNADA_PERSON_ALIASES.items():
        if source.replace(" ", "") in compact:
            return target
    tokens = [token for token in re.findall(r"[\u0c80-\u0cff]+", query) if token not in KANNADA_LOOKUP_STOP_WORDS]
    candidates = []
    for token in tokens:
        if token in KANNADA_PERSON_ALIASES:
            candidates.append(KANNADA_PERSON_ALIASES[token])
    if candidates:
        return " ".join(dict.fromkeys(candidates))
    return None


def classify_meta_intent(query: str) -> str:
    text = re.sub(r"\s+", " ", query.casefold()).strip()
    terms = set(normalized_search_terms(query))
    master = classify_master_intent(query)
    if terms.intersection({"motto", "moto", "mission", "purpose"}) and terms.intersection(
        {"what", "your", "you", "assistant", "system", "chatbot"}
    ):
        return "system_identity"
    if master["intent"] == "conversation_smalltalk":
        return "conversation_smalltalk"
    if master["intent"] == "system_identity":
        return "system_identity"
    if master["intent"] in {"admin_user_management", "admin_audit_lookup"}:
        return "admin_governance"
    if text == "help" or any(
        phrase in text
        for phrase in SYSTEM_INTENT_TERMS
        if phrase != "help"
    ):
        return "system_identity"
    conversation_single_terms = {phrase for phrase in CONVERSATION_INTENT_TERMS if " " not in phrase}
    conversation_phrases = CONVERSATION_INTENT_TERMS - conversation_single_terms
    if text in conversation_single_terms or terms.intersection(conversation_single_terms) or any(
        phrase in text for phrase in conversation_phrases
    ):
        return "conversation_smalltalk"
    if terms.intersection(ADMIN_INTENT_TERMS) or any(phrase in text for phrase in ADMIN_INTENT_TERMS if " " in phrase):
        return "admin_governance"
    return "crime_data"


def answer_meta_intent(
    *,
    query: str,
    db: Database,
    user: AuthenticatedUser,
    meta_intent: str,
) -> dict[str, Any] | None:
    if meta_intent == "crime_data":
        return None
    if meta_intent == "system_identity":
        answer = "\n".join(
            [
                "I am KSP AI Crime Intelligence Assistant.",
                "My motto is: serve investigators with evidence-first intelligence, protect sensitive records, and never infer guilt.",
                "",
                "What I can do:",
                "- Search FIRs, case numbers, suspects, victims, complainants, police stations, districts, and status.",
                "- Build suspect profiles, linked FIR summaries, case timelines, and similar-case leads.",
                "- Analyze criminal networks, repeat-offender links, common associates, and relationship paths.",
                "- Trace financial fraud flows, mule accounts, circular transfers, dead ends, bank/IFSC/branch/manager links.",
                "- Discover crime patterns, hotspots, seasonal trends, MO clusters, and emerging threat signals.",
                "- Summarize uploaded FIR copies and case files after Python extraction from PDF, DOC, DOCX, CSV, and Excel files.",
                "- Support Kannada and English questions, translation, voice input, PDF export, RBAC, audit trails, and admin governance.",
                "",
                "How to ask:",
                "- \"Find Ravi Kumar\"",
                "- \"Who filed the complaint against Ravi Kumar?\"",
                "- \"Show FIR BLR-001 details\"",
                "- \"Top 10 theft cases in Bengaluru\"",
                "- \"Show money trail for A1\"",
                "- \"Open criminal network for Ravi Kumar\"",
                "- \"Forecast hotspot risk by district\"",
                "",
                "I only answer from stored records your role can access. If the data is not stored or not authorized for your role, I will say so instead of guessing.",
                "This system heading identifies the platform as Karnataka State Police Intelligent Systems By Jaswinzz Gowda.",
            ]
        )
        return meta_response(
            intent="system_identity",
            answer=answer,
            evidence_mode="system_capability_metadata",
            data_scope="system_metadata",
            filters={"routing_priority": "SYSTEM", "crime_api_called": False},
        )
    if meta_intent == "conversation_smalltalk":
        answer = "Hello. I am ready to help with authorized crime intelligence."
        return meta_response(
            intent="conversation_smalltalk",
            answer=answer,
            evidence_mode="conversation_control",
            data_scope="conversation_only",
            filters={"routing_priority": "CONVERSATION", "crime_api_called": False},
        )
    if meta_intent == "admin_governance":
        return answer_admin_intent(query=query, db=db, user=user)
    return None


def meta_response(
    *,
    intent: str,
    answer: str,
    evidence_mode: str,
    data_scope: str,
    filters: dict[str, Any],
) -> dict[str, Any]:
    return {
        "intent": intent,
        "answer": answer,
        "visible_case_count": 0,
        "sources": [],
        "query_analysis": {
            "original_query": "",
            "normalized_query": "",
            "interpreted_terms": [],
            "interpreted_filters": filters,
            "evidence_mode": evidence_mode,
            "data_scope": data_scope,
        },
        "safeguards": [
            "No crime records were queried for this non-crime control intent.",
            "No external knowledge or unstored facts are used for this answer.",
        ],
        "extracted_entities": {},
        "_non_crime_response": True,
    }


def answer_admin_intent(*, query: str, db: Database, user: AuthenticatedUser) -> dict[str, Any]:
    text = query.casefold()
    filters = {
        "routing_priority": "ADMIN",
        "rbac_role": user.role,
        "rbac_district": user.district,
    }
    if "audit" in text or "activity" in text:
        if not role_has_permission(user.role, "audit:read"):
            answer = "Your current role is not authorized to read audit or activity logs."
            filters["permission"] = "audit:read denied"
        else:
            logs = db.list_audit_logs(limit=8)
            lines = [f"Recent audit activity visible to {user.role}: {len(logs)} event(s)."]
            for item in logs[:5]:
                lines.append(
                    f"- {item['action']} / {item['status']} / actor {item.get('actor_username') or 'system'} / resource {item.get('resource_type')}"
                )
            answer = "\n".join(lines)
            filters["permission"] = "audit:read allowed"
        return meta_response(
            intent="admin_audit_lookup",
            answer=answer,
            evidence_mode="audit_log_index",
            data_scope="governance_data_only",
            filters=filters,
        )
    if "user" in text or "role" in text or "permission" in text or "access" in text or "admin" in text:
        if not role_has_permission(user.role, "user:read"):
            answer = (
                "Your current role cannot manage user accounts. "
                "Only Super Admin can create, block, reset, delete, and monitor users."
            )
            filters["permission"] = "user:read denied"
        else:
            users = db.list_users()
            role_counts: dict[str, int] = {}
            active_count = 0
            for row in users:
                role_counts[row["role"]] = role_counts.get(row["role"], 0) + 1
                active_count += 1 if row.get("is_active") else 0
            lines = [
                f"Super Admin user management has {len(users)} account(s), {active_count} active.",
                "Available controls: create user, edit role, edit district, block/unblock, reset password, view activity, delete account, and rate-limit breach monitoring.",
                "Role distribution: "
                + ", ".join(f"{role}: {count}" for role, count in sorted(role_counts.items())),
                "Passwords are never readable because they are stored as hashes; Super Admin can only reset them.",
            ]
            filters["permission"] = "user:read allowed"
            filters["user_count"] = len(users)
        return meta_response(
            intent="admin_user_management",
            answer="\n".join(lines) if "lines" in locals() else answer,
            evidence_mode="rbac_user_registry",
            data_scope="governance_data_only",
            filters=filters,
        )
    return meta_response(
        intent="admin_governance",
        answer="Admin governance can review users, permissions, audit logs, rate-limit breaches, and account activity according to RBAC.",
        evidence_mode="rbac_policy",
        data_scope="governance_data_only",
        filters=filters,
    )


def enrich_query_with_memory(query: str, memory: dict[str, Any] | None) -> str:
    if not memory:
        return query
    text = query.casefold()
    if query_has_explicit_focus(query):
        return query
    focus = memory.get("current_suspect") or memory.get("current_network_focus")
    if focus and any(term in text for term in NETWORK_INTENT_TERMS | FINANCIAL_INTENT_TERMS):
        return f"{query} for {focus}"
    if focus and any(term in text for term in {"profile", "behavior", "offender", "criminological"}):
        return f"{query} for {focus}"
    return query


def query_has_explicit_focus(query: str) -> bool:
    entities = extract_query_entities(query)
    if any(entities.get(key) for key in ["fir_number", "account", "vehicle", "phone"]):
        return True
    return bool(extract_person_name_from_query(query))


def extract_query_entities(query: str) -> dict[str, Any]:
    entities: dict[str, Any] = {}
    explicit_firs = [item.upper() for item in CONTEXT_FIR_RE.findall(query)]
    fir_candidates = explicit_firs or [item.upper() for item in FIR_RE.findall(query)]
    firs = [item for item in fir_candidates if not looks_like_account_identifier(item)]
    if firs:
        entities["fir_number"] = firs[-1]
    phones = [re.sub(r"\s+", "", item) for item in PHONE_RE.findall(query)]
    if phones:
        entities["phone"] = list(dict.fromkeys(phones))
    vehicles = [re.sub(r"[-\s]+", "", item.upper()) for item in VEHICLE_RE.findall(query)]
    if vehicles:
        entities["vehicle"] = list(dict.fromkeys(vehicles))
    accounts = []
    for item in ACCOUNT_RE.findall(query):
        cleaned = re.sub(r"\s+", "-", item.upper()).replace("_", "-")
        prefix = cleaned.split("-", 1)[0]
        if prefix in {"FIR", "PHONE", "MOBILE", "ON", "VEHICLE"}:
            continue
        if not looks_like_account_identifier(cleaned):
            continue
        if cleaned not in accounts:
            accounts.append(cleaned)
    if accounts:
        entities["account"] = accounts
    dates = DATE_RE.findall(query)
    if dates:
        entities["date"] = list(dict.fromkeys(dates))
    lowered = query.casefold()
    districts = [district.title() for district in DISTRICT_NAMES if district in lowered]
    if districts:
        entities["district"] = sorted(dict.fromkeys(districts))
    person = extract_person_name_from_query(query)
    if person and is_probable_person_focus(person):
        entities["person_name"] = person
    return entities


def looks_like_account_identifier(value: str) -> bool:
    cleaned = re.sub(r"\s+", "-", value.upper())
    return "-AC-" in cleaned or bool(re.match(r"^[A-Z]{3,12}[-_]?(?:AC[-_]?)?\d{4,}$", cleaned))


def build_conversation_memory(conversation_id: str, messages: list[dict[str, Any]]) -> dict[str, Any]:
    memory = {field: None for field in MEMORY_FIELDS}
    memory["conversation_id"] = conversation_id
    for message in messages:
        metadata = message.get("metadata") or {}
        saved = metadata.get("conversation_memory")
        if isinstance(saved, dict):
            for key in MEMORY_FIELDS:
                if key != "conversation_id" and saved.get(key):
                    memory[key] = saved[key]
        entities = metadata.get("extracted_entities")
        if isinstance(entities, dict):
            apply_entities_to_memory(memory, entities)
        content = str(message.get("content") or "")
        firs = FIR_RE.findall(content)
        if firs:
            memory["current_fir"] = firs[-1].upper()
        if message.get("role") == "assistant" and content:
            memory["last_response"] = content[:900]
            if metadata.get("selected_module"):
                memory["last_module"] = metadata["selected_module"]
    return memory


def apply_entities_to_memory(memory: dict[str, Any], entities: dict[str, Any]) -> None:
    mapping = {
        "current_case_id": ["case_id", "current_case_id"],
        "current_fir": ["fir_number", "current_fir"],
        "current_suspect": ["suspect_name", "current_suspect"],
        "current_victim": ["victim_name", "current_victim"],
        "current_complainant": ["complainant_name", "current_complainant"],
        "current_district": ["district", "current_district"],
        "current_network_focus": ["network_focus", "current_network_focus"],
        "current_financial_entity": ["financial_entity", "current_financial_entity"],
    }
    for memory_key, entity_keys in mapping.items():
        for entity_key in entity_keys:
            value = entities.get(entity_key)
            if value:
                memory[memory_key] = value[0] if isinstance(value, list) else value
                break


def response_entities(query: str, response: dict[str, Any]) -> dict[str, Any]:
    entities: dict[str, Any] = dict(extract_query_entities(query))
    for key, value in dict(response.get("extracted_entities") or {}).items():
        if value not in (None, "", []):
            if key in entities and isinstance(entities[key], list) and not isinstance(value, list):
                continue
            entities[key] = value
    sources = response.get("sources") or []
    if sources:
        first = sources[0]
        entities.setdefault("case_id", first.get("case_id"))
        entities.setdefault("fir_number", first.get("fir_number"))
        entities.setdefault("district", first.get("district"))
        excerpt = str(first.get("excerpt") or "")
        for label, key in [
            ("Suspect", "suspect_name"),
            ("Victim", "victim_name"),
            ("Complainant", "complainant_name"),
        ]:
            match = re.search(rf"{label}:\s*([^|.\n]+)", excerpt, flags=re.IGNORECASE)
            if match and match.group(1).strip():
                entities.setdefault(key, match.group(1).strip())
    firs = FIR_RE.findall(query)
    if firs:
        entities.setdefault("fir_number", firs[-1].upper())
    name = extract_person_name_from_query(query)
    if name and is_probable_person_focus(name):
        entities.setdefault("person_name", name)
    return {key: value for key, value in entities.items() if value not in (None, "", [])}


def update_conversation_memory(
    base_memory: dict[str, Any] | None,
    response: dict[str, Any],
) -> dict[str, Any]:
    memory = {field: None for field in MEMORY_FIELDS}
    if base_memory:
        for key in MEMORY_FIELDS:
            memory[key] = base_memory.get(key)
    entities = response.get("extracted_entities") or {}
    apply_entities_to_memory(memory, entities)
    if not memory.get("current_suspect") and entities.get("person_name") and response.get("intent") in {
        "accused_lookup",
        "behavioral_criminological_profiling",
        "criminal_network_analysis",
        "direct_retrieval",
        "financial_analysis",
        "money_trail",
        "network_analysis",
        "offender_profile",
        "suspect_profile",
    }:
        memory["current_suspect"] = entities["person_name"]
    memory["last_module"] = response.get("selected_module") or memory.get("last_module")
    memory["last_response"] = strip_existing_contract(str(response.get("answer") or ""))[:900]
    return memory


def extract_person_name_from_query(query: str) -> str | None:
    if KANNADA_RE.search(query):
        kannada_name = extract_kannada_person_name(query)
        if kannada_name:
            return kannada_name
    patterns = [
        r"\bis\s+(.+?)\s+(?:a\s+)?(?:high\s+)?(?:risk|risky|dangerous)\s+(?:suspect|accused|offender|person)\b",
        r"\bis\s+(.+?)\s+(?:suspect|accused|offender|person)\b",
        r"\b(?:risk|profile|profiling|offender\s+profile|suspect\s+profile)\s+(?:for|of|about)\s+(.+)",
        r"\b(?:show|find|search|expand)\s+(.+?)\s+(?:network|profile|money\s+trail|trail|associates|links)\b",
        r"\bagainst\s+(.+)",
        r"\bfor\s+(.+)",
        r"\babout\s+(.+)",
        r"\bof\s+(.+)",
        r"\b(?:suspect|accused|victim|complainant|person)\s+(?:name\s+)?(?:is|called|named)?\s*(.+)",
        r"\bname\s+(?:is|called|named)\s+(.+)",
    ]
    for pattern in patterns:
        match = re.search(pattern, query, flags=re.IGNORECASE)
        if not match:
            continue
        tail = re.split(
            r"\b(?:assessment|case|cases|dangerous|fir|file|number|status|network|money|trail|details|copy|limit|top|desc|high|risk|risky|score|profile|offender|suspect|accused)\b",
            match.group(1),
            maxsplit=1,
            flags=re.IGNORECASE,
        )[0]
        words = [
            word
            for word in re.findall(r"[^\W_]+", tail, flags=re.UNICODE)
            if word.casefold() not in CASE_LOOKUP_FILLER_TERMS and not word.isdigit()
        ]
        if words:
            return " ".join(words[:4]).strip()
    return None


def infer_targeted_case_intent(query: str, conversation_memory: dict[str, Any] | None) -> str | None:
    terms = set(normalized_search_terms(query))
    text = query.casefold()
    if looks_like_strict_sql_query(text):
        return None
    if terms.intersection(
        {
            "behavioral",
            "criminological",
            "demographic",
            "forecast",
            "hotspot",
            "hotspots",
            "intelligence",
            "network",
            "pattern",
            "prevention",
            "proactive",
            "profile",
            "profiling",
            "risk",
            "risky",
            "socio",
            "sociological",
            "trend",
        }
    ):
        return None
    has_memory_target = bool(
        conversation_memory
        and any(
            conversation_memory.get(key)
            for key in ["current_fir", "current_suspect", "current_victim", "current_complainant"]
        )
    )
    if (
        {"complainant", "complaint"}.intersection(terms)
        or "who filed" in text
        or is_kannada_complainant_query(query)
    ):
        return "complainant_lookup"
    if "victim" in terms:
        return "victim_lookup"
    if {"accused", "accuse", "acuse", "suspect", "named"}.intersection(terms):
        return "accused_lookup"
    if {"status", "open", "closed", "pending"}.intersection(terms) and (has_memory_target or explicit_case_lookup(query)):
        return "case_status"
    if {"summary", "summarize", "happened"}.intersection(terms) and (has_memory_target or explicit_case_lookup(query)):
        return "decision_support" if "summarize" in terms else "case_summary"
    if {"fir", "case", "file", "number"}.intersection(terms) and (extract_person_name_from_query(query) or has_memory_target):
        return "fir_lookup"
    return None


def looks_like_strict_sql_query(text: str) -> bool:
    return bool(
        re.search(r"\bwhere\b.*\b(?:suspect|victim|complainant|person)?\s*name\s*(?:is|=|called|named)\b", text)
        or re.search(r"\b(?:limit|top|order\s+by|desc)\s+\d{1,3}\b", text)
    )


def answer_targeted_case_question(
    *,
    query: str,
    db: Database,
    user: AuthenticatedUser,
    include_sources: bool,
    conversation_memory: dict[str, Any] | None,
    language: str = "en",
) -> dict[str, Any] | None:
    intent = infer_targeted_case_intent(query, conversation_memory)
    if not intent:
        return None

    lookup_text = targeted_lookup_text(query, conversation_memory, intent)
    if not lookup_text:
        return None
    matches = accessible_targeted_case_matches(db, user, lookup_text, intent=intent, limit=8)
    if not matches:
        return {
            "intent": intent,
            "answer": no_accessible_target_answer(intent, lookup_text, language=language),
            "visible_case_count": 0,
            "sources": [],
            "query_analysis": strict_response_query_analysis(
                original_query=query,
                working_query=lookup_text,
                evidence_mode="local_case_search_index",
                filters=targeted_query_filters(
                    intent=intent,
                    lookup_text=lookup_text,
                    memory_used=bool(conversation_memory),
                    language=language,
                    query=query,
                ),
            ),
            "safeguards": case_answer_safeguards(),
            "extracted_entities": {"person_name": extract_person_name_from_query(query) or lookup_text},
            "_answer_language": language,
        }

    masked = [mask_case(case, user.role) for case in matches]
    answer = targeted_answer_text(intent, masked, lookup_text, language=language)
    entities = entities_from_case(masked[0])
    person_name = extract_person_name_from_query(query)
    if person_name:
        entities["person_name"] = person_name
    return {
        "intent": intent,
        "answer": answer,
        "visible_case_count": len(masked),
        "sources": case_sources(matches, user.role) if include_sources else [],
        "query_analysis": strict_response_query_analysis(
            original_query=query,
            working_query=lookup_text,
            evidence_mode="local_case_search_index",
            filters=targeted_query_filters(
                intent=intent,
                lookup_text=lookup_text,
                memory_used=bool(conversation_memory),
                language=language,
                query=query,
            ),
        ),
        "safeguards": case_answer_safeguards(),
        "extracted_entities": entities,
        "_answer_language": language,
    }


def targeted_query_filters(
    *,
    intent: str,
    lookup_text: str,
    memory_used: bool,
    language: str,
    query: str,
) -> dict[str, Any]:
    filters: dict[str, Any] = {
        "intent": intent,
        "lookup": lookup_text,
        "memory_used": memory_used,
        "language": language,
        "lookup_priority": ["fir_number", "complainant_name", "police_station", "status"]
        if intent == "complainant_lookup"
        else [],
    }
    if language == "kn":
        normalized = normalize_kannada_query(query)
        filters["kannada_dictionary"] = dictionary_status()
        if normalized["corrections"]:
            filters["kannada_corrections"] = normalized["corrections"]
    return filters


def targeted_lookup_text(query: str, conversation_memory: dict[str, Any] | None, intent: str) -> str:
    identifier = FIR_RE.findall(query)
    if identifier:
        return identifier[-1]
    name = extract_person_name_from_query(query)
    if name:
        return name
    if intent in {"case_status", "case_summary", "decision_support", "fir_lookup"}:
        for key in ["current_fir", "current_suspect", "current_victim", "current_complainant", "current_case_id"]:
            if conversation_memory and conversation_memory.get(key):
                return str(conversation_memory[key])
    lookup_terms = case_lookup_terms(query)
    if lookup_terms:
        return " ".join(lookup_terms)
    for key in ["current_fir", "current_suspect", "current_victim", "current_complainant", "current_case_id"]:
        if conversation_memory and conversation_memory.get(key):
            return str(conversation_memory[key])
    return ""


def accessible_case_matches(db: Database, user: AuthenticatedUser, lookup_text: str, *, limit: int) -> list[dict[str, Any]]:
    raw_results = db.search_cases(lookup_text, limit=max(limit, 20))
    matches: list[dict[str, Any]] = []
    for item in raw_results["results"]:
        case = item["case"]
        if can_access_case(user, case):
            matches.append(case)
        if len(matches) >= limit:
            break
    return matches


def accessible_targeted_case_matches(
    db: Database,
    user: AuthenticatedUser,
    lookup_text: str,
    *,
    intent: str,
    limit: int,
) -> list[dict[str, Any]]:
    if intent in {"complainant_lookup", "victim_lookup", "accused_lookup"} and not FIR_RE.findall(lookup_text):
        matches = accessible_person_case_matches(db, user, lookup_text, intent=intent, limit=limit)
        if matches:
            return matches
    return accessible_case_matches(db, user, lookup_text, limit=limit)


def accessible_person_case_matches(
    db: Database,
    user: AuthenticatedUser,
    lookup_text: str,
    *,
    intent: str,
    limit: int,
) -> list[dict[str, Any]]:
    terms = normalized_search_terms(lookup_text)
    if not terms:
        return []
    field_groups = {
        "complainant_lookup": ("suspect_name", "victim_name", "complainant_name"),
        "victim_lookup": ("victim_name", "suspect_name", "complainant_name"),
        "accused_lookup": ("suspect_name", "victim_name", "complainant_name"),
    }
    fields = field_groups.get(intent, ("suspect_name", "victim_name", "complainant_name"))
    rows: list[dict[str, Any]] = []
    with db.connect() as conn:
        for field in fields:
            clauses = [f"LOWER({field}) LIKE ?"] * len(terms)
            params = [f"%{term.casefold()}%" for term in terms]
            found = conn.execute(
                f"""
                SELECT *
                FROM cases
                WHERE {' AND '.join(clauses)}
                ORDER BY COALESCE(incident_at, updated_at, created_at) DESC, fir_number ASC
                LIMIT ?
                """,
                (*params, max(limit * 3, 20)),
            ).fetchall()
            for row in found:
                case = dict(row)
                if can_access_case(user, case) and case["id"] not in {item["id"] for item in rows}:
                    rows.append(case)
                if len(rows) >= limit:
                    return rows
    return rows


def answer_suspect_profile_question(
    *,
    query: str,
    db: Database,
    user: AuthenticatedUser,
    include_sources: bool,
    conversation_memory: dict[str, Any] | None,
) -> dict[str, Any] | None:
    terms = set(normalized_search_terms(query))
    text = query.casefold()
    if not terms.intersection(PROFILE_INTENT_TERMS) and "high risk suspect" not in text:
        return None
    if terms.intersection({"forecast", "hotspot", "hotspots", "trend", "trends"}) and not terms.intersection(
        {"accused", "offender", "person", "profile", "suspect"}
    ):
        return None

    focus = extract_person_name_from_query(query)
    if not focus and conversation_memory:
        focus = conversation_memory.get("current_suspect") or conversation_memory.get("current_network_focus")
    if not focus or not is_probable_person_focus(focus):
        return None

    visible_case_rows = visible_cases(user, db.list_cases())
    profile = suspect_profile(focus, visible_case_rows, user.role)
    normalized_focus = focus.strip().casefold()
    matched_cases = [
        case
        for case in visible_case_rows
        if normalized_focus
        and (
            str(case.get("suspect_name") or "").casefold() == normalized_focus
            or normalized_focus in str(case.get("suspect_name") or "").casefold()
            or any(part.startswith(normalized_focus) for part in str(case.get("suspect_name") or "").casefold().split())
        )
    ]
    if not matched_cases:
        return {
            "intent": "offender_profile",
            "answer": (
                f"No accessible suspect profile was found for {focus}. "
                "The system checked authorized suspect-name fields and did not find a matching record."
            ),
            "visible_case_count": 0,
            "sources": [],
            "query_analysis": strict_response_query_analysis(
                original_query=query,
                working_query=focus,
                evidence_mode="suspect_profile_records",
                filters={"profile_focus": focus, "profile_match_count": 0},
            ),
            "safeguards": case_answer_safeguards(),
            "extracted_entities": {"person_name": focus, "suspect_name": focus},
        }

    display_name = str(matched_cases[0].get("suspect_name") or focus)
    risk_level = str(profile.get("risk_level") or "none").title()
    risk_score = int(profile.get("risk_score") or 0)
    status_summary = ", ".join(f"{item['key']} {item['count']}" for item in profile.get("statuses", [])[:4]) or "not recorded"
    district_summary = ", ".join(profile.get("districts", [])[:6]) or "not recorded"
    factors = profile.get("risk_factors") or ["No repeat-risk factor is present in the accessible records."]
    indicators = profile.get("behavioral_indicators") or []
    lines = [
        f"Risk profile for {display_name}:",
        f"Risk level: {risk_level}",
        f"Risk score: {risk_score}/100",
        f"Named in accessible FIRs: {len(matched_cases)}",
        f"Districts: {district_summary}",
        f"Case status mix: {status_summary}",
        "",
        "Risk factors from records:",
        *[f"- {factor}" for factor in factors[:5]],
    ]
    if indicators:
        lines.extend(["", "Behavioral indicators:", *[f"- {item}" for item in indicators[:4]]])
    lines.extend(
        [
            "",
            "Linked FIRs:",
            *[
                f"- {case['fir_number']} | {case['district']} | {case['status']}"
                for case in matched_cases[:8]
            ],
            "",
            "This is record-linkage risk triage from stored FIR records, not a finding of guilt or an automated prediction.",
        ]
    )
    return {
        "intent": "offender_profile",
        "answer": "\n".join(lines),
        "visible_case_count": len(matched_cases),
        "sources": case_sources(matched_cases, user.role) if include_sources else [],
        "query_analysis": strict_response_query_analysis(
            original_query=query,
            working_query=focus,
            evidence_mode="suspect_profile_records",
            filters={
                "profile_focus": display_name,
                "profile_match_count": len(matched_cases),
                "risk_level": str(profile.get("risk_level") or "none"),
                "risk_score": risk_score,
            },
        ),
        "safeguards": profile.get("safeguards") or case_answer_safeguards(),
        "extracted_entities": {
            "person_name": display_name,
            "suspect_name": display_name,
            "fir_number": matched_cases[0]["fir_number"],
            "risk_level": str(profile.get("risk_level") or "none"),
            "risk_score": risk_score,
        },
    }


def answer_person_network_question(
    *,
    query: str,
    db: Database,
    user: AuthenticatedUser,
    include_sources: bool,
    conversation_memory: dict[str, Any] | None,
) -> dict[str, Any] | None:
    terms = set(normalized_search_terms(query))
    if not terms.intersection(NETWORK_INTENT_TERMS):
        return None
    focus = extract_person_name_from_query(query)
    if not focus and conversation_memory:
        focus = conversation_memory.get("current_suspect") or conversation_memory.get("current_network_focus")
    account_focus = first_entity_value(extract_query_entities(query).get("account"))
    if not focus and not account_focus:
        return None
    if focus and not account_focus and not is_probable_person_focus(focus):
        return None

    person_cases = accessible_person_case_matches(
        db,
        user,
        focus or account_focus or "",
        intent="accused_lookup",
        limit=20,
    ) if focus else []
    linked_case_ids = {case["id"] for case in person_cases}
    visible_case_rows = visible_cases(user, db.list_cases())
    transactions = matching_transactions(
        visible_transactions(user, db.list_financial_transactions(), visible_case_rows),
        focus=focus,
        account=account_focus,
        linked_case_ids=linked_case_ids,
    )
    related_people = sorted(related_people_from_cases(person_cases, focus or ""))
    account_links = sorted(accounts_from_transactions(transactions))
    firs = [case["fir_number"] for case in person_cases]
    if not person_cases and not transactions:
        target = focus or account_focus
        return {
            "intent": "criminal_network_analysis",
            "answer": (
                f"No accessible relationship graph was found for {target}. "
                "The system checked case person fields and financial transaction links available to your role."
            ),
            "visible_case_count": 0,
            "sources": [],
            "query_analysis": strict_response_query_analysis(
                original_query=query,
                working_query=query,
                evidence_mode="case_transaction_network",
                filters={"network_focus": target, "memory_used": bool(conversation_memory)},
            ),
            "safeguards": case_answer_safeguards(),
            "extracted_entities": {"person_name": focus, "account": account_focus, "network_focus": target},
        }

    target = focus or account_focus
    fir_list = ", ".join(f"`{f}`" for f in firs[:6]) if firs else "None"
    people_list = ", ".join(f"`{p}`" for p in related_people[:8]) if related_people else "None on record"
    account_list = ", ".join(f"`{a}`" for a in account_links[:8]) if account_links else "None on record"
    lines = [
        f"### 🕸️ Criminal Relationship Network — {target}\n",
        f"**Summary of Accessible Links:**",
        f"- **FIR Links**: {len(firs)} ({fir_list})",
        f"- **Person Links**: {len(related_people)} ({people_list})",
        f"- **Account Links**: {len(account_links)} ({account_list})",
        f"- **Transaction Links**: {len(transactions)}",
    ]
    if transactions:
        lines.append("\n**Strongest Money-Flow Edges:**")
        for tx in transactions[:5]:
            lines.append(
                f"- `{tx['source_account']}` ➜ `{tx['target_account']}` | **{format_rupees(tx['amount'])}** | {tx.get('description') or 'no description'}"
            )
    lines.append("\n> **Compliance Note:** This is relationship intelligence from accessible records, not a determination of guilt.")
    return {
        "intent": "criminal_network_analysis",
        "answer": "\n".join(lines),
        "visible_case_count": len(person_cases),
        "sources": case_sources(person_cases, user.role) if include_sources else [],
        "query_analysis": strict_response_query_analysis(
            original_query=query,
            working_query=query,
            evidence_mode="case_transaction_network",
            filters={
                "network_focus": target,
                "memory_used": bool(conversation_memory),
                "fir_links": firs[:10],
                "transaction_count": len(transactions),
            },
        ),
        "safeguards": case_answer_safeguards(),
        "extracted_entities": {
            "person_name": focus,
            "account": account_focus,
            "network_focus": target,
            "fir_number": firs[0] if firs else None,
        },
    }


def answer_money_trail_question(
    *,
    query: str,
    db: Database,
    user: AuthenticatedUser,
    include_sources: bool,
    conversation_memory: dict[str, Any] | None,
) -> dict[str, Any] | None:
    terms = set(normalized_search_terms(query))
    if not terms.intersection(FINANCIAL_INTENT_TERMS):
        return None
    entities = extract_query_entities(query)
    focus = extract_person_name_from_query(query)
    account_focus = first_entity_value(entities.get("account"))
    if not focus and conversation_memory:
        focus = conversation_memory.get("current_financial_entity") or conversation_memory.get("current_suspect")
    if not focus and not account_focus:
        return None
    if focus and not account_focus and not is_probable_person_focus(focus):
        return None

    visible_case_rows = visible_cases(user, db.list_cases())
    focus_cases = accessible_person_case_matches(db, user, focus, intent="accused_lookup", limit=10) if focus else []
    transactions = matching_transactions(
        visible_transactions(user, db.list_financial_transactions(), visible_case_rows),
        focus=focus,
        account=account_focus,
        linked_case_ids={case["id"] for case in focus_cases},
    )
    linked_case_ids = {tx.get("case_id") for tx in transactions if tx.get("case_id")}
    linked_cases = [case for case in visible_case_rows if case["id"] in linked_case_ids]
    if focus and not linked_cases:
        linked_cases = focus_cases
    target = focus or account_focus
    if not transactions:
        answer = (
            f"### 💸 Money Trail — {target}\n\n"
            f"No accessible financial transaction row is linked to **{target}**.\n\n"
            "> The system checked account numbers, account-holder names, descriptions, and case-linked transactions available to your role."
        )
    else:
        lines = [
            f"### 💸 Money Trail Analysis — {target}",
            f"Money trail for {target}.",
            f"\n**{len(transactions)} accessible transaction(s) identified.**\n",
        ]
        for index, tx in enumerate(transactions[:10], 1):
            source_meta = bank_meta(tx, "source")
            target_meta = bank_meta(tx, "target")
            lines.append(
                f"**{index}.** `{tx['source_account']}` ➜ `{tx['target_account']}`  \n"
                f"   💰 **Amount**: {format_rupees(tx['amount'])} | 📅 **Date**: {tx['occurred_at']}  \n"
                f"   🏦 **Source**: {source_meta}  \n"
                f"   🏦 **Target**: {target_meta}  \n"
                f"   🔗 **Case Link**: {tx.get('case_id') or 'No linked case'}"
            )
        terminal_accounts = dead_end_accounts(transactions)
        if terminal_accounts:
            ta_list = ", ".join(f"`{a}`" for a in terminal_accounts[:8])
            lines.append(f"\n**⚠️ Dead-End Accounts (funds sink):** {ta_list}")
        lines.append("\n> **Compliance Note:** This is a money-flow lead from stored transactions, not a finding of guilt.")
        answer = "\n\n".join(lines)
    return {
        "intent": "financial_analysis",
        "answer": answer,
        "visible_case_count": len(linked_cases),
        "sources": case_sources(linked_cases, user.role) if include_sources else [],
        "query_analysis": strict_response_query_analysis(
            original_query=query,
            working_query=query,
            evidence_mode="financial_transaction_records",
            filters={
                "financial_focus": target,
                "memory_used": bool(conversation_memory),
                "transaction_count": len(transactions),
                "linked_case_count": len(linked_cases),
            },
        ),
        "safeguards": case_answer_safeguards(),
        "extracted_entities": {
            "person_name": focus,
            "account": account_focus,
            "financial_entity": target,
            "fir_number": linked_cases[0]["fir_number"] if linked_cases else None,
        },
    }


def first_entity_value(value: Any) -> str | None:
    if isinstance(value, list) and value:
        return str(value[0])
    if value:
        return str(value)
    return None


def normalized_term_set(value: str) -> set[str]:
    return set(normalized_search_terms(value))


def same_named_person(left: str, right: str) -> bool:
    left_terms = normalized_term_set(left)
    right_terms = normalized_term_set(right)
    if not left_terms or not right_terms:
        return False
    return left_terms.issubset(right_terms) or right_terms.issubset(left_terms)


def is_probable_person_focus(value: str) -> bool:
    terms = normalized_term_set(value)
    if not terms:
        return False
    non_person_terms = (
        CASE_LOOKUP_FILLER_TERMS
        | NETWORK_INTENT_TERMS
        | FINANCIAL_INTENT_TERMS
        | {
            "by",
            "crime",
            "female",
            "hotspot",
            "hotspots",
            "mo",
            "night",
            "pattern",
            "premise",
            "victims",
            "vehicle",
            "weapon",
        }
    )
    return any(term not in non_person_terms for term in terms)


def related_people_from_cases(cases: list[dict[str, Any]], focus: str) -> set[str]:
    people: set[str] = set()
    for case in cases:
        for field in ["suspect_name", "victim_name", "complainant_name"]:
            value = str(case.get(field) or "").strip()
            if not value or value.casefold() in {"unknown", "n/a", "na"}:
                continue
            if focus and same_named_person(value, focus):
                continue
            people.add(value)
    return people


def matching_transactions(
    transactions: list[dict[str, Any]],
    *,
    focus: str | None,
    account: str | None,
    linked_case_ids: set[str] | None = None,
) -> list[dict[str, Any]]:
    focus_terms = normalized_term_set(focus or "")
    account_norm = normalize_account(account or "")
    linked_case_ids = linked_case_ids or set()
    matches = []
    for tx in transactions:
        if tx.get("case_id") in linked_case_ids:
            matches.append(tx)
            continue
        accounts = {normalize_account(str(tx.get("source_account") or "")), normalize_account(str(tx.get("target_account") or ""))}
        if account_norm and account_norm in accounts:
            matches.append(tx)
            continue
        haystack = " ".join(
            str(tx.get(field) or "")
            for field in [
                "source_account_holder",
                "target_account_holder",
                "source_account",
                "target_account",
                "description",
            ]
        )
        hay_terms = normalized_term_set(haystack)
        if focus_terms and focus_terms.issubset(hay_terms):
            matches.append(tx)
    return matches


def normalize_account(value: str) -> str:
    return re.sub(r"[^A-Z0-9]", "", value.upper())


def accounts_from_transactions(transactions: list[dict[str, Any]]) -> set[str]:
    accounts: set[str] = set()
    for tx in transactions:
        for field in ["source_account", "target_account"]:
            value = str(tx.get(field) or "").strip()
            if value:
                accounts.add(value)
    return accounts


def dead_end_accounts(transactions: list[dict[str, Any]]) -> list[str]:
    sources = {normalize_account(str(tx.get("source_account") or "")) for tx in transactions}
    targets = {normalize_account(str(tx.get("target_account") or "")) for tx in transactions}
    labels = {normalize_account(str(tx.get("target_account") or "")): str(tx.get("target_account") or "") for tx in transactions}
    return [labels[item] for item in sorted(targets - sources) if item]


def bank_meta(tx: dict[str, Any], side: str) -> str:
    parts = [
        tx.get(f"{side}_account_holder"),
        tx.get(f"{side}_bank_name"),
        tx.get(f"{side}_ifsc_code"),
        tx.get(f"{side}_branch"),
        tx.get(f"{side}_bank_manager_phone"),
    ]
    return " / ".join(str(part) for part in parts if part) or "Not recorded"


def format_rupees(value: Any) -> str:
    try:
        amount = float(value)
    except (TypeError, ValueError):
        return str(value)
    return f"INR {amount:,.0f}"


def targeted_answer_text(
    intent: str,
    cases: list[dict[str, Any]],
    lookup_text: str,
    *,
    language: str = "en",
) -> str:
    case = cases[0]
    station = case.get("police_station") or "Not recorded in accessible case data"
    if intent == "complainant_lookup":
        if language == "kn":
            return kannada_complainant_answer(cases, lookup_text)
        if len(cases) > 1:
            lines = [f"{lookup_text} appears in {len(cases)} FIR records."]
            for index, item in enumerate(cases, 1):
                lines.append(
                    f"{index}. FIR {item['fir_number']}\n"
                    f"   Complainant: {item.get('complainant_name') or 'Not recorded in accessible case data'}\n"
                    f"   Status: {item['status']}"
                )
            lines.append("Please specify the FIR number for detailed information.")
            return "\n\n".join(lines)
        return "\n".join(
            [
                f"{lookup_text} appears in FIR {case['fir_number']}.",
                f"Complainant:\n{case.get('complainant_name') or 'Not recorded in accessible case data'}",
                f"FIR Number:\n{case['fir_number']}",
                f"Police Station:\n{station}",
                f"District:\n{case['district']}",
                f"Case Status:\n{case['status']}",
                f"Linked person on record:\nSuspect/accused field lists {case.get('suspect_name') or 'not recorded'}; victim field lists {case.get('victim_name') or 'not recorded'}.",
            ]
        )
    if intent == "victim_lookup":
        return "\n".join(
            [
                f"Victim:\n{case.get('victim_name') or 'Not recorded in accessible case data'}",
                f"FIR:\n{case['fir_number']}",
                f"District:\n{case['district']}",
                f"Status:\n{case['status']}",
                f"Complainant:\n{case.get('complainant_name') or 'Not recorded in accessible case data'}",
            ]
        )
    if intent == "accused_lookup":
        return "\n".join(
            [
                f"Suspect/accused field:\n{case.get('suspect_name') or 'Not recorded in accessible case data'}",
                f"FIR:\n{case['fir_number']}",
                f"Case number:\n{case.get('source_record_id') or case.get('id')}",
                f"District:\n{case['district']}",
                f"Status:\n{case['status']}",
                "This indicates the person is named in the case record; it is not a determination of guilt.",
            ]
        )
    if intent == "case_status":
        return "\n".join(
            [
                f"Current status:\n{case['status']}",
                f"FIR:\n{case['fir_number']}",
                f"District:\n{case['district']}",
                f"Suspect/accused field:\n{case.get('suspect_name') or 'Not recorded in accessible case data'}",
                f"Complainant:\n{case.get('complainant_name') or 'Not recorded in accessible case data'}",
            ]
        )
    if intent in {"case_summary", "decision_support"}:
        return "\n".join(
            [
                f"FIR:\n{case['fir_number']}",
                f"Status:\n{case['status']}",
                f"Summary:\n{case.get('summary') or 'No summary is recorded in accessible case data.'}",
            ]
        )

    lines = [f"Accessible case details for {lookup_text}:"]
    for index, item in enumerate(cases[:5], 1):
        lines.append(
            f"{index}. FIR {item['fir_number']} | Case no {item.get('source_record_id') or item.get('id')} | "
            f"{item['district']} | {item['status']} | suspect {item.get('suspect_name') or 'not recorded'} | "
            f"victim {item.get('victim_name') or 'not recorded'} | complainant {item.get('complainant_name') or 'not recorded'}."
        )
    return "\n".join(lines)


def kannada_complainant_answer(cases: list[dict[str, Any]], lookup_text: str) -> str:
    if len(cases) > 1:
        lines = [f"{lookup_text} ಸಂಬಂಧಿಸಿದ {len(cases)} FIR ದಾಖಲೆಗಳು ಕಂಡುಬಂದಿವೆ."]
        for index, case in enumerate(cases, 1):
            lines.append(
                f"{index}. FIR {case['fir_number']}\n"
                f"   ದೂರುದಾರರ ಹೆಸರು: {case.get('complainant_name') or 'ದಾಖಲೆಯಲ್ಲಿ ಲಭ್ಯವಿಲ್ಲ'}\n"
                f"   ಸ್ಥಿತಿ: {kannada_status(case['status'])}"
            )
        lines.append("ವಿವರವಾದ ಮಾಹಿತಿಗಾಗಿ FIR ಸಂಖ್ಯೆಯನ್ನು ಸೂಚಿಸಿ.")
        return "\n\n".join(lines)

    case = cases[0]
    station = case.get("police_station") or "ದಾಖಲೆಯಲ್ಲಿ ಲಭ್ಯವಿಲ್ಲ"
    return "\n".join(
        [
            f"{lookup_text} ಸಂಬಂಧಿಸಿದ FIR {case['fir_number']} ದಾಖಲೆ ಕಂಡುಬಂದಿದೆ.",
            "",
            f"ದೂರುದಾರರ ಹೆಸರು:\n{case.get('complainant_name') or 'ದಾಖಲೆಯಲ್ಲಿ ಲಭ್ಯವಿಲ್ಲ'}",
            f"FIR ಸಂಖ್ಯೆ:\n{case['fir_number']}",
            f"ಪೊಲೀಸ್ ಠಾಣೆ:\n{station}",
            f"ಜಿಲ್ಲೆ:\n{case['district']}",
            f"ಸ್ಥಿತಿ:\n{kannada_status(case['status'])}",
            f"ಮೂಲ ದಾಖಲೆ:\n{case['fir_number']}",
            "ಗಮನಿಸಿ:\nಇದು ದಾಖಲೆ ಆಧಾರಿತ ಮಾಹಿತಿ ಮಾತ್ರ. ತನಿಖೆಯ ಅಂತಿಮ ತೀರ್ಮಾನವಲ್ಲ.",
        ]
    )


def kannada_status(status: str) -> str:
    return {
        "open": "ತೆರೆದಿದೆ",
        "under_review": "ಪರಿಶೀಲನೆಯಲ್ಲಿದೆ",
        "closed": "ಮುಚ್ಚಲಾಗಿದೆ",
    }.get(status, status)


def no_accessible_target_answer(intent: str, lookup_text: str, *, language: str = "en") -> str:
    if language == "kn":
        return (
            f"{lookup_text} ಕುರಿತು ಈ ಪಾತ್ರಕ್ಕೆ ಲಭ್ಯವಿರುವ ದೂರುದಾರ/ಪ್ರಕರಣ ದಾಖಲೆ ಸಿಗಲಿಲ್ಲ. "
            "ಹೆಸರಿನ ಅಕ್ಷರಗಳು ಸರಿಯಿದೆಯೇ ನೋಡಿ ಅಥವಾ ಸಂಬಂಧಿತ ಜಿಲ್ಲೆ/ಸಂವೇದನಾಶೀಲ ದಾಖಲೆಗಳಿಗೆ ಪ್ರವೇಶವಿದೆಯೇ ದೃಢಪಡಿಸಿ."
        )
    subject = intent.replace("_", " ")
    return (
        f"No accessible case record contains {subject} information for {lookup_text}. "
        "Check spelling, try a shorter name, or confirm your role has access to the relevant district and sensitivity level."
    )


def entities_from_case(case: dict[str, Any]) -> dict[str, Any]:
    return {
        "case_id": case.get("id"),
        "fir_number": case.get("fir_number"),
        "district": case.get("district"),
        "suspect_name": case.get("suspect_name"),
        "victim_name": case.get("victim_name"),
        "complainant_name": case.get("complainant_name"),
    }


def case_answer_safeguards() -> list[str]:
    return [
        "Use this output as an investigative aid, not as a determination of guilt.",
        "Verify all findings against source records before operational action.",
        "Responses only use records the current role is allowed to access.",
        "No external knowledge or unstored facts are used for this answer.",
    ]


def build_indexed_case_lookup_response(
    *,
    query: str,
    db: Database,
    user: AuthenticatedUser,
    include_sources: bool,
) -> dict | None:
    lookup_terms = case_lookup_terms(query)
    if not lookup_terms:
        return None

    raw_results = db.search_cases(" ".join(lookup_terms), limit=20)
    matches = []
    for item in raw_results["results"]:
        case = item["case"]
        if can_access_case(user, case):
            matches.append(case)
        if len(matches) >= 5:
            break

    if not matches and explicit_case_lookup(query):
        searched_terms = ", ".join(lookup_terms)
        return {
            "intent": "direct_retrieval",
            "answer": (
                "No accessible case record answers that question yet. "
                f"The local case index was checked for: {searched_terms}. "
                "Try a FIR number, a partial or full suspect name, district, crime type, or case status from records you can access."
            ),
            "visible_case_count": 0,
            "sources": [],
            "query_analysis": strict_response_query_analysis(
                original_query=query,
                working_query=" ".join(lookup_terms),
                evidence_mode="local_case_search_index",
                filters={"terms": lookup_terms},
            ),
            "safeguards": [
                "Use this output as an investigative aid, not as a determination of guilt.",
                "Responses only use records the current role is allowed to access.",
                "No external knowledge or unstored facts are used for this answer.",
            ],
        }
    if not matches:
        return None

    masked = [mask_case(case, user.role) for case in matches]
    people = matched_people_summary(masked[0])
    if len(masked) == 1:
        case = masked[0]
        answer = (
            f"{' '.join(lookup_terms).title()} appears in FIR {case['fir_number']} in {case['district']}. "
            f"Current status is {case['status']}. {people}"
        )
    else:
        lines = [f"Accessible case files related to {' '.join(lookup_terms)}:"]
        for index, case in enumerate(masked[:5], 1):
            lines.append(
                f"{index}. FIR {case['fir_number']} | {case['district']} | {case['status']} | "
                f"suspect {case.get('suspect_name') or 'not recorded'} | "
                f"victim {case.get('victim_name') or 'not recorded'} | "
                f"complainant {case.get('complainant_name') or 'not recorded'}."
            )
        if people:
            lines.append(people)
        answer = "\n".join(lines)

    return {
        "intent": "direct_retrieval",
        "answer": answer,
        "visible_case_count": len(masked),
        "sources": case_sources(matches, user.role) if include_sources else [],
        "query_analysis": strict_response_query_analysis(
            original_query=query,
            working_query=" ".join(lookup_terms),
            evidence_mode="local_case_search_index",
            filters={"terms": lookup_terms},
        ),
        "safeguards": [
            "Use this output as an investigative aid, not as a determination of guilt.",
            "Verify all findings against source records before operational action.",
            "Responses only use records the current role is allowed to access.",
            "No external knowledge or unstored facts are used for this answer.",
        ],
    }


def case_lookup_terms(query: str) -> list[str]:
    terms = normalized_search_terms(query)
    if not terms:
        return []
    if set(terms).intersection(CASE_LOOKUP_BLOCKER_TERMS):
        return []
    has_hint = bool(set(terms).intersection(CASE_LOOKUP_HINTS))
    meaningful = [term for term in terms if term not in CASE_LOOKUP_FILLER_TERMS]
    if has_hint:
        return meaningful[:4]
    if 1 <= len(meaningful) <= 3 and all(len(term) >= 3 for term in meaningful):
        return meaningful
    return []


def explicit_case_lookup(query: str) -> bool:
    return bool(set(normalized_search_terms(query)).intersection(CASE_LOOKUP_HINTS))


def matched_people_summary(case: dict) -> str:
    parts = []
    if case.get("suspect_name"):
        parts.append(f"suspect {case['suspect_name']}")
    if case.get("victim_name"):
        parts.append(f"victim {case['victim_name']}")
    if case.get("complainant_name"):
        parts.append(f"complainant {case['complainant_name']}")
    return "Named fields: " + "; ".join(parts[:3]) + "." if parts else ""


def case_sources(cases: list[dict], role: str) -> list[dict]:
    return [
        {
            "case_id": case["id"],
            "fir_number": masked["fir_number"],
            "district": masked["district"],
            "status": masked["status"],
            "sensitivity": masked["sensitivity"],
            "excerpt": " | ".join(
                part
                for part in [
                    str(masked["summary"])[:180],
                    f"Suspect: {masked.get('suspect_name')}" if masked.get("suspect_name") else "",
                    f"Victim: {masked.get('victim_name')}" if masked.get("victim_name") else "",
                    f"Complainant: {masked.get('complainant_name')}" if masked.get("complainant_name") else "",
                ]
                if part
            ),
        }
        for case in cases
        for masked in [mask_case(case, role)]
    ]


def strict_response_query_analysis(
    *,
    original_query: str,
    working_query: str,
    evidence_mode: str,
    filters: dict[str, Any] | None = None,
) -> dict[str, Any]:
    terms = normalized_search_terms(working_query)
    interpreted_terms = [term for term in terms if term not in CASE_LOOKUP_FILLER_TERMS]
    return {
        "original_query": original_query,
        "normalized_query": " ".join(terms),
        "interpreted_terms": interpreted_terms[:12],
        "interpreted_filters": filters or ({"terms": interpreted_terms[:12]} if interpreted_terms else {}),
        "evidence_mode": evidence_mode,
        "data_scope": "crime_data_only",
    }


MODULE_ROUTING: dict[str, dict[str, str]] = {
    "system_identity": {
        "module": "SYSTEM: KSP AI Crime Intelligence Assistant",
        "api": "No crime API called",
    },
    "conversation_smalltalk": {
        "module": "SYSTEM: Conversation Controller",
        "api": "No crime API called",
    },
    "admin_governance": {
        "module": "MODULE 10: Security and Governance",
        "api": "RBAC/Audit APIs",
    },
    "admin_user_management": {
        "module": "MODULE 10: Security and Governance",
        "api": "GET /admin/users",
    },
    "admin_audit_lookup": {
        "module": "MODULE 9: Explainable AI",
        "api": "GET /audit/logs",
    },
    "strict_sql_case_search": {
        "module": "MODULE 1: Conversational Crime Intelligence",
        "api": "POST /intelligence/query",
    },
    "direct_retrieval": {
        "module": "MODULE 1: Conversational Crime Intelligence",
        "api": "POST /intelligence/query",
    },
    "complainant_lookup": {
        "module": "MODULE 1: Conversational Crime Intelligence",
        "api": "POST /intelligence/query",
    },
    "victim_lookup": {
        "module": "MODULE 1: Conversational Crime Intelligence",
        "api": "POST /intelligence/query",
    },
    "accused_lookup": {
        "module": "MODULE 1: Conversational Crime Intelligence",
        "api": "POST /intelligence/query",
    },
    "fir_lookup": {
        "module": "MODULE 1: Conversational Crime Intelligence",
        "api": "POST /intelligence/query",
    },
    "case_summary": {
        "module": "MODULE 6: Investigator Decision Support",
        "api": "GET /decision-support/cases/{id}",
    },
    "case_status": {
        "module": "MODULE 1: Conversational Crime Intelligence",
        "api": "POST /intelligence/query",
    },
    "criminal_network_analysis": {
        "module": "MODULE 2: Criminal Network Analysis",
        "api": "GET /analytics/network",
    },
    "network_analysis": {
        "module": "MODULE 2: Criminal Network Analysis",
        "api": "GET /analytics/network",
    },
    "crime_pattern_discovery": {
        "module": "MODULE 3: Crime Pattern Analytics",
        "api": "GET /analytics/patterns",
    },
    "crime_trend_intelligence": {
        "module": "MODULE 3: Crime Pattern Analytics",
        "api": "GET /analytics/advanced-crime",
    },
    "pattern_discovery": {
        "module": "MODULE 3: Crime Pattern Analytics",
        "api": "GET /analytics/patterns",
    },
    "trend_analysis": {
        "module": "MODULE 3: Crime Pattern Analytics",
        "api": "GET /analytics/trends",
    },
    "socio_demographic_crime_insight": {
        "module": "MODULE 4: Sociological Analytics",
        "api": "GET /analytics/sociological",
    },
    "sociological_insight": {
        "module": "MODULE 4: Sociological Analytics",
        "api": "GET /analytics/sociological",
    },
    "behavioral_criminological_profiling": {
        "module": "MODULE 5: Suspect Profiling",
        "api": "GET /profiles/suspects/{name}",
    },
    "offender_profile": {
        "module": "MODULE 5: Suspect Profiling",
        "api": "GET /profiles/suspects/{name}",
    },
    "decision_support": {
        "module": "MODULE 6: Investigator Decision Support",
        "api": "GET /decision-support/cases/{id}",
    },
    "financial_analysis": {
        "module": "MODULE 7: Financial Crime Analytics",
        "api": "GET /financial/analysis",
    },
    "money_trail": {
        "module": "MODULE 7: Financial Crime Analytics",
        "api": "GET /financial/analysis",
    },
    "proactive_crime_prevention_intelligence": {
        "module": "MODULE 8: Forecasting",
        "api": "GET /forecast/hotspots",
    },
    "forecasting": {
        "module": "MODULE 8: Forecasting",
        "api": "GET /forecast/hotspots",
    },
    "explainable_ai": {
        "module": "MODULE 9: Explainable AI",
        "api": "GET /explain/cases/{id}",
    },
    "security_governance": {
        "module": "MODULE 10: Security and Governance",
        "api": "Authentication/RBAC/Audit APIs",
    },
}


FOLLOW_UPS_BY_MODULE = {
    "MODULE 1: Conversational Crime Intelligence": [
        "Show full FIR details for a specific FIR number.",
        "Search another suspect, victim, or complainant name.",
        "Show case status and investigation history.",
    ],
    "MODULE 2: Criminal Network Analysis": [
        "Open the suspect network graph.",
        "Show money trails linked to this case.",
        "Find common associates or repeat connections.",
    ],
    "MODULE 3: Crime Pattern Analytics": [
        "Show district hotspot map for this crime type.",
        "Compare seasonal trends for the same MO.",
        "List emerging threat clusters by district.",
    ],
    "MODULE 4: Sociological Analytics": [
        "Break this down by age and gender.",
        "Compare migration, education, or urbanization indicators.",
        "Show district-wise demographic risk factors.",
    ],
    "MODULE 5: Suspect Profiling": [
        "Show the suspect profile for a named person.",
        "Compare modus operandi across linked cases.",
        "List repeat-offender indicators with FIR references.",
    ],
    "MODULE 6: Investigator Decision Support": [
        "Generate a case timeline.",
        "Find similar past cases and outcomes.",
        "Show potential investigative leads.",
    ],
    "MODULE 7: Financial Crime Analytics": [
        "Show circular transaction chains.",
        "Identify structuring or mule-account signals.",
        "Open the financial network graph.",
    ],
    "MODULE 8: Forecasting": [
        "Show early warning hotspots.",
        "Compare projected hotspots by district.",
        "Suggest prevention planning priorities.",
    ],
    "MODULE 9: Explainable AI": [
        "Explain why this conclusion was made.",
        "Show audit trail for this action.",
        "Show evidence references for this case.",
    ],
    "MODULE 10: Security and Governance": [
        "Show user access for this role.",
        "Check district restrictions.",
        "Show recent audit history.",
    ],
}


def first_scalar(value: Any) -> str | None:
    if isinstance(value, list):
        return str(value[0]) if value else None
    if value not in (None, "", []):
        return str(value)
    return None


def concrete_followups(
    *,
    selected_module: str,
    intent: str,
    response: dict[str, Any],
    entities: dict[str, Any],
) -> list[str]:
    sources = response.get("sources") or []
    first_source = sources[0] if sources else {}
    fir = first_scalar(entities.get("fir_number")) or first_source.get("fir_number")
    person = (
        first_scalar(entities.get("suspect_name"))
        or first_scalar(entities.get("person_name"))
        or first_scalar(entities.get("network_focus"))
        or first_scalar(entities.get("financial_entity"))
    )
    account = first_scalar(entities.get("account"))
    actions: list[str] = []

    def add(action: str | None) -> None:
        if not action:
            return
        text = " ".join(str(action).split())
        if not is_concrete_followup(text):
            return
        if text not in actions:
            actions.append(text)

    if selected_module == "MODULE 1: Conversational Crime Intelligence":
        add(f"Show full FIR details for {fir}" if fir else None)
        add(f"Show suspect profile for {person}" if person else None)
        add(f"Show criminal network for {person}" if person else None)
        add(f"Show case status for {fir}" if fir else None)
    elif selected_module == "MODULE 2: Criminal Network Analysis":
        add(f"Show money trail for {person or account}" if (person or account) else None)
        add(f"Show suspect profile for {person}" if person else None)
        add(f"Show full FIR details for {fir}" if fir else None)
    elif selected_module == "MODULE 5: Suspect Profiling":
        add(f"Show linked FIRs for {person}" if person else None)
        add(f"Show criminal network for {person}" if person else None)
        add(f"Show money trail for {person}" if person else None)
    elif selected_module == "MODULE 7: Financial Crime Analytics":
        add(f"Trace account network for {account}" if account else None)
        add(f"Show criminal network for {person}" if person else None)
        add(f"Show full FIR details for {fir}" if fir else None)
    elif fir:
        add(f"Show full FIR details for {fir}")
        add(f"Show evidence trail for {fir}")
    elif person:
        add(f"Show suspect profile for {person}")
        add(f"Show criminal network for {person}")

    if actions:
        return actions[:3]
    return [item for item in FOLLOW_UPS_BY_MODULE.get(selected_module, []) if is_concrete_followup(item)][:3]


def is_concrete_followup(action: str) -> bool:
    text = action.casefold()
    placeholders = {
        "another suspect",
        "a named person",
        "named person",
        "specific fir",
        "specific case",
        "this case",
        "this crime type",
        "same mo",
        "same case",
        "{subject}",
    }
    return bool(action.strip()) and not any(marker in text for marker in placeholders)


def apply_crime_assistant_contract(
    response: dict[str, Any],
    *,
    query: str,
    user: AuthenticatedUser,
    conversation_context: str | None,
    language: str = "en",
    conversation_memory: dict[str, Any] | None = None,
) -> dict[str, Any]:
    if response.get("_assistant_contract_applied"):
        return response
    intent = str(response.get("intent") or "direct_retrieval")
    route = MODULE_ROUTING.get(
        intent,
        {
            "module": "MODULE 1: Conversational Crime Intelligence",
            "api": "POST /intelligence/query",
        },
    )
    selected_module = route["module"]
    selected_api = route["api"]
    sources = response.get("sources") or []
    query_analysis = response.get("query_analysis") or {}
    confidence = intelligence_confidence(response)
    trace = orchestration_trace(
        intent=intent,
        module=selected_module,
        api=selected_api,
    )
    reasoning = reasoning_path(
        intent=intent,
        module=selected_module,
        api=selected_api,
        response=response,
        user=user,
        conversation_context=conversation_context,
    )
    base_answer = strip_existing_contract(str(response.get("answer") or ""))
    entities = response_entities(query, response)
    response["extracted_entities"] = entities
    followups = concrete_followups(
        selected_module=selected_module,
        intent=intent,
        response=response,
        entities=entities,
    )
    no_evidence = not sources and int(response.get("visible_case_count") or 0) == 0
    if intent == "conversation_smalltalk":
        followups = []
    elif no_evidence and not response.get("_non_crime_response"):
        followups = []
    elif response.get("_non_crime_response") and intent != "admin_user_management":
        followups = []
    presentation = response_presentation(
        user=user,
        intent=intent,
        followups=followups[:3],
        query=query,
    )
    if no_evidence and not response.get("_non_crime_response"):
        presentation["render_evidence"] = False
        presentation["render_actions"] = False
        presentation["suggested_actions"] = []
    response["answer"] = operational_answer_text(
        base_answer=base_answer,
        response=response,
        intent=intent,
        language=language,
        followups=followups[:3],
        presentation=presentation,
    )
    response["selected_module"] = selected_module
    response["selected_api"] = selected_api
    response["confidence"] = confidence
    response["reasoning"] = reasoning
    response["suggested_followups"] = followups[:3]
    response["orchestration_trace"] = trace
    response["presentation"] = presentation
    response["conversation_memory"] = update_conversation_memory(conversation_memory, response)
    response["_assistant_contract_applied"] = True
    master = classify_master_intent(query)
    query_analysis.setdefault("interpreted_filters", {})
    query_analysis["interpreted_filters"].setdefault("master_intent_category", master["category"])
    query_analysis["interpreted_filters"].setdefault("master_response_rule", master["response_rule"])
    query_analysis["interpreted_filters"].setdefault("generated_intent_variations", taxonomy_status()["total_variations"])
    query_analysis["interpreted_filters"].setdefault("selected_module", selected_module)
    query_analysis["interpreted_filters"].setdefault("selected_api", selected_api)
    query_analysis["interpreted_filters"].setdefault("confidence", confidence)
    query_analysis["interpreted_filters"].setdefault("orchestration_trace", trace)
    query_analysis["interpreted_filters"].setdefault("extracted_entities", entities)
    response["query_analysis"] = query_analysis
    return response


def strip_existing_contract(answer: str) -> str:
    match = re.search(
        r"(?:Direct Answer|Response):\n(?P<body>.*?)(?:\n\n(?:Reasoning|Suggested Follow-up Questions|Evidence Sources|Confidence):|\Z)",
        answer,
        flags=re.DOTALL,
    )
    if match:
        return match.group("body").strip()
    return re.sub(r"^Intent:\s.*?(?=\n\n(?:Direct Answer|Response):\n|\Z)", "", answer, flags=re.DOTALL).strip() or answer


def response_presentation(
    *,
    user: AuthenticatedUser,
    intent: str,
    followups: list[str],
    query: str,
) -> dict[str, Any]:
    audience = {
        "investigator": "investigator",
        "supervisor": "investigator",
        "super_admin": "administrator",
        "analyst": "analyst",
        "policymaker": "analyst",
        "viewer": "investigator",
    }.get(user.role, "investigator")
    non_record_intents = {"conversation_smalltalk", "system_identity"}
    return {
        "audience": audience,
        "show_debug_metadata": False,
        "show_intent": False,
        "show_module": False,
        "show_api": False,
        "show_confidence": False,
        "show_query_analysis": False,
        "render_evidence": intent not in non_record_intents,
        "render_actions": intent not in non_record_intents,
        "suggested_actions": [] if intent == "conversation_smalltalk" else followups,
        "master_training": taxonomy_status(),
        "recognized_category": classify_master_intent(query)["category"],
    }


def operational_answer_text(
    *,
    base_answer: str,
    response: dict[str, Any],
    intent: str,
    language: str,
    followups: list[str],
    presentation: dict[str, Any],
) -> str:
    if intent == "conversation_smalltalk":
        return base_answer.strip()

    sections = [base_answer.strip()]
    if presentation.get("render_evidence"):
        references = operational_evidence_references(response)
        if references:
            label = "ಮೂಲ ದಾಖಲೆಗಳು" if language == "kn" else "Evidence references"
            sections.append(f"{label}:\n" + "\n".join(f"- {item}" for item in references))
    if followups and presentation.get("render_actions"):
        label = "ಸೂಚಿತ ಕ್ರಮಗಳು" if language == "kn" else "Suggested actions"
        sections.append(f"{label}:\n" + "\n".join(f"- {item}" for item in followups[:3]))
    if presentation.get("render_evidence") and not re.search(
        r"not (?:a )?(?:determination|finding) of guilt|not.*guilt",
        base_answer,
        re.IGNORECASE,
    ):
        note = (
            "ಗಮನಿಸಿ: ಇದು ದಾಖಲೆ ಆಧಾರಿತ ತನಿಖಾ ಸಹಾಯ ಮಾತ್ರ; ದೋಷದ ತೀರ್ಮಾನವಲ್ಲ."
            if language == "kn"
            else "Note: This is an investigative aid from stored records, not a determination of guilt."
        )
        sections.append(note)
    return "\n\n".join(section for section in sections if section)


def operational_evidence_references(response: dict[str, Any]) -> list[str]:
    sources = response.get("sources") or []
    if sources:
        references = []
        for source in sources[:6]:
            label = source.get("fir_number") or source.get("case_id") or "record"
            parts = [label, source.get("district"), source.get("status")]
            references.append(" | ".join(str(part) for part in parts if part))
        return references
    count = int(response.get("visible_case_count") or 0)
    if count:
        return [f"{count:,} accessible aggregate record(s)"]
    return []


def evidence_source_text(sources: list[dict[str, Any]], response: dict[str, Any]) -> str:
    if response.get("_non_crime_response"):
        mode = str((response.get("query_analysis") or {}).get("evidence_mode") or "system metadata")
        return f"- No crime record required; answered from {mode}."
    if sources:
        rows = []
        for source in sources[:8]:
            label = source.get("fir_number") or source.get("case_id") or "record"
            rows.append(
                "- "
                + " | ".join(
                    str(part)
                    for part in [
                        label,
                        source.get("district"),
                        source.get("status"),
                        source.get("sensitivity"),
                    ]
                    if part
                )
            )
        return "\n".join(rows)
    count = int(response.get("visible_case_count") or 0)
    mode = str((response.get("query_analysis") or {}).get("evidence_mode") or "stored crime records")
    if count:
        return f"- Aggregate stored crime records: {count:,} accessible record(s), evidence mode {mode}."
    return "- No source record matched within the current role and district permissions."


def orchestration_trace(*, intent: str, module: str, api: str) -> list[str]:
    return [
        "User",
        "Chat Interface",
        f"LLM Intent Detection ({intent.replace('_', ' ')})",
        "Orchestrator",
        f"{module} ({api})",
        "Response Generator",
        "Chat UI",
    ]


def intelligence_confidence(response: dict[str, Any]) -> str:
    if response.get("_non_crime_response"):
        return "99% - High; no crime evidence lookup was required for this control intent."
    count = int(response.get("visible_case_count") or 0)
    sources = response.get("sources") or []
    filters = ((response.get("query_analysis") or {}).get("interpreted_filters") or {})
    corrected = bool(filters.get("spelling_corrections"))
    mode = str((response.get("query_analysis") or {}).get("evidence_mode") or "")
    if count <= 0:
        return "35% - Low; no accessible matching source record was found."
    if sources and not corrected and mode in {"indexed_sql_case_search", "local_case_search_index", "local_case_records"}:
        return "92% - High; indexed source record(s) are available to this role."
    if sources:
        return (
            "84% - Medium-High; source record(s) matched after spelling/query normalization."
            if corrected
            else "90% - High; source record(s) matched."
        )
    if count >= 1:
        return "72% - Medium; aggregate evidence is available, but individual source rows are not exposed for this module."
    return "35% - Low; insufficient stored evidence."


def reasoning_path(
    *,
    intent: str,
    module: str,
    api: str,
    response: dict[str, Any],
    user: AuthenticatedUser,
    conversation_context: str | None,
) -> list[str]:
    analysis = response.get("query_analysis") or {}
    filters = analysis.get("interpreted_filters") or {}
    terms = analysis.get("interpreted_terms") or []
    count = int(response.get("visible_case_count") or 0)
    steps = [
        f"Detected intent as {intent.replace('_', ' ')} and routed it to {module}.",
        f"Used {api} with evidence mode {analysis.get('evidence_mode', 'stored crime records')}.",
        f"Applied RBAC for role {user.role} and district {user.district}; returned only accessible records.",
    ]
    if terms:
        steps.append(f"Normalized query terms: {', '.join(str(term) for term in terms[:8])}.")
    if filters:
        steps.append(f"Applied filters: {', '.join(f'{key}={value}' for key, value in list(filters.items())[:8])}.")
    steps.append(f"Evidence reviewed: {count:,} accessible record(s).")
    steps.append(
        "Conversation memory was used from previous messages." if conversation_context else "No previous conversation context was required for this answer."
    )
    steps.append("This is investigative intelligence only and is not a determination of guilt.")
    return steps


def build_intelligence_response(
    *,
    query: str,
    language: str,
    cases: list[dict],
    user: AuthenticatedUser,
    include_sources: bool,
    db: Database | None = None,
    conversation_context: str | None = None,
    conversation_memory: dict[str, Any] | None = None,
) -> dict:
    resolved_language = detect_query_language(query, language)
    response = None
    if db is not None:
        meta_intent = classify_meta_intent(query)
        response = answer_meta_intent(query=query, db=db, user=user, meta_intent=meta_intent)
    working_query = prepare_query_for_intent(query, resolved_language)
    working_context = conversation_context
    if resolved_language == "kn" and working_query == query:
        working_query = translate_text(query, "kn", "en")["translated_text"]
        if conversation_context:
            working_context = translate_text(conversation_context, "kn", "en")["translated_text"]
    working_query = enrich_query_with_memory(working_query, conversation_memory)

    if db is not None:
        response = response or answer_suspect_profile_question(
            query=working_query,
            db=db,
            user=user,
            include_sources=include_sources,
            conversation_memory=conversation_memory,
        )
    if response is None and db is not None:
        response = answer_targeted_case_question(
            query=working_query,
            db=db,
            user=user,
            include_sources=include_sources,
            conversation_memory=conversation_memory,
            language=resolved_language,
        )
    if response is None and db is not None:
        response = answer_money_trail_question(
            query=working_query,
            db=db,
            user=user,
            include_sources=include_sources,
            conversation_memory=conversation_memory,
        )
    if response is None and db is not None:
        response = answer_person_network_question(
            query=working_query,
            db=db,
            user=user,
            include_sources=include_sources,
            conversation_memory=conversation_memory,
        )
    if response is None and db is not None:
        response = answer_strict_case_search_query(
            query=working_query,
            db=db,
            user=user,
            include_sources=include_sources,
        )
    if response is None and db is not None:
        response = answer_incident_intelligence_query(
            query=working_query,
            db=db,
            user=user,
            include_sources=include_sources,
        )
    if response is None and db is not None:
        response = build_indexed_case_lookup_response(
            query=working_query,
            db=db,
            user=user,
            include_sources=include_sources,
        )
    if response is None:
        response = answer_query(
            working_query,
            cases,
            user,
            include_sources=include_sources,
            conversation_context=working_context,
        )
    response.setdefault(
        "query_analysis",
        strict_response_query_analysis(
            original_query=query,
            working_query=working_query,
            evidence_mode="local_case_records" if response.get("sources") else "crime_data_only_fallback",
        ),
    )
    response["query_analysis"]["original_query"] = query
    response["query_analysis"]["normalized_query"] = " ".join(normalized_search_terms(working_query))
    response.setdefault("safeguards", [])
    if "No external knowledge or unstored facts are used for this answer." not in response["safeguards"]:
        response["safeguards"].append("No external knowledge or unstored facts are used for this answer.")
    if resolved_language == "kn" and response.get("_answer_language") != "kn":
        response["answer"] = translate_text(response["answer"], "en", "kn")["translated_text"]
    return apply_crime_assistant_contract(
        response,
        query=query,
        user=user,
        conversation_context=conversation_context,
        language=resolved_language,
        conversation_memory=conversation_memory,
    )


def get_current_user(
    request: Request,
    credentials: Annotated[HTTPAuthorizationCredentials | None, Depends(bearer_scheme)],
    settings: Annotated[Settings, Depends(get_app_settings)],
    db: Annotated[Database, Depends(get_db)],
) -> AuthenticatedUser:
    if credentials is None or credentials.scheme.lower() != "bearer":
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Missing bearer token",
        )

    try:
        user = decode_access_token(credentials.credentials, settings)
    except jwt.ExpiredSignatureError as exc:
        raise HTTPException(status_code=401, detail="Token expired") from exc
    except jwt.InvalidTokenError as exc:
        raise HTTPException(status_code=401, detail="Invalid token") from exc

    if db.is_token_revoked(user.jti):
        raise HTTPException(status_code=401, detail="Token revoked")

    stored = db.get_user_by_id(user.id)
    if not stored or not stored["is_active"]:
        raise HTTPException(status_code=401, detail="User disabled")

    return user


def require_permission(permission: str):
    def dependency(user: Annotated[AuthenticatedUser, Depends(get_current_user)]) -> AuthenticatedUser:
        if not role_has_permission(user.role, permission):
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail=f"Missing permission: {permission}",
            )
        return user

    return dependency


UPLOAD_EXTENSIONS = {".csv", ".xlsx", ".xls", ".pdf", ".doc", ".docx"}
CASE_SHEET_EXTENSIONS = {".csv", ".xlsx", ".xls"}
DOCUMENT_EXTENSIONS = {".pdf", ".doc", ".docx"}
UPLOAD_TYPES = {"case_sheet", "fir_copy"}
INLINE_PREVIEW_EXTENSIONS = {".pdf", ".png", ".jpg", ".jpeg", ".gif", ".webp", ".txt", ".csv", ".json", ".html", ".htm"}
PREVIEW_MIME_TYPES = {
    ".csv": "text/csv; charset=utf-8",
    ".gif": "image/gif",
    ".htm": "text/html; charset=utf-8",
    ".html": "text/html; charset=utf-8",
    ".jpeg": "image/jpeg",
    ".jpg": "image/jpeg",
    ".json": "application/json; charset=utf-8",
    ".pdf": "application/pdf",
    ".png": "image/png",
    ".txt": "text/plain; charset=utf-8",
    ".webp": "image/webp",
}
HEADER_ALIASES = {
    "fir_number": {"fir", "firno", "firnumber", "firnum", "firid", "fir_no", "fir_number"},
    "year": {"year", "caseyear"},
    "district": {"district", "districtname", "policedistrict"},
    "status": {"status", "casestatus", "investigationstatus"},
    "case_type": {"casetype", "crimetype", "crime", "offence", "offense"},
    "modus_operandi": {"modusoperandi", "mo", "mocodes", "method"},
    "incident_at": {"incidentat", "incidentdate", "dateocc", "occurrencedate", "date"},
    "complainant_name": {"complainant", "complainantname", "informant", "reporter"},
    "complainant_phone": {"complainantphone", "phone", "mobile", "contact"},
    "victim_name": {"victim", "victimname"},
    "victim_age": {"victimage", "ageofvictim"},
    "victim_gender": {"victimgender", "victimsex", "gender"},
    "suspect_name": {"suspect", "suspectname", "accused", "accusedname", "person"},
    "suspect_age": {"suspectage", "accusedage"},
    "suspect_gender": {"suspectgender", "accusedgender"},
    "summary": {"summary", "details", "description", "casebrief", "narrative", "facts"},
    "sensitivity": {"sensitivity", "classification"},
    "latitude": {"latitude", "lat"},
    "longitude": {"longitude", "lon", "lng", "longitude"},
    "source_record_id": {"sourcerecordid", "recordid", "caseid", "caseno", "casenumber"},
}


def safe_upload_filename(filename: str) -> str:
    name = Path(filename or "upload").name
    return re.sub(r"[^a-zA-Z0-9_. -]", "_", name).strip() or "upload"


def normalized_header(value: str) -> str:
    return re.sub(r"[^a-z0-9]", "", value.lower())


def parse_case_sheet(content: bytes, extension: str, source_system: str) -> tuple[list[CaseCreate], int, list[str]]:
    if extension == ".csv":
        rows = list(csv.DictReader(io.StringIO(content.decode("utf-8-sig"))))
    elif extension == ".xlsx":
        rows = parse_xlsx_rows(content)
    else:
        return [], 0, ["Legacy .xls files are stored as evidence but not parsed; save as .xlsx or .csv to import rows."]

    cases: list[CaseCreate] = []
    skipped = 0
    notes: list[str] = []
    for row_number, row in enumerate(rows, start=2):
        try:
            cases.append(case_from_upload_row(row, source_system))
        except ValueError as exc:
            skipped += 1
            if len(notes) < 8:
                notes.append(f"Row {row_number}: {exc}")
    return cases, skipped, notes


def parse_xlsx_rows(content: bytes) -> list[dict[str, str]]:
    with zipfile.ZipFile(io.BytesIO(content)) as archive:
        shared_strings = read_xlsx_shared_strings(archive)
        sheet_name = "xl/worksheets/sheet1.xml"
        if sheet_name not in archive.namelist():
            sheets = [name for name in archive.namelist() if name.startswith("xl/worksheets/sheet") and name.endswith(".xml")]
            if not sheets:
                return []
            sheet_name = sorted(sheets)[0]
        root = ElementTree.fromstring(archive.read(sheet_name))
    namespace = {"a": "http://schemas.openxmlformats.org/spreadsheetml/2006/main"}
    table: list[list[str]] = []
    for row in root.findall(".//a:sheetData/a:row", namespace):
        values: dict[int, str] = {}
        for cell in row.findall("a:c", namespace):
            ref = cell.attrib.get("r", "")
            index = xlsx_column_index(ref)
            cell_type = cell.attrib.get("t")
            raw = cell.findtext("a:v", default="", namespaces=namespace)
            if cell_type == "s" and raw:
                value = shared_strings[int(raw)] if int(raw) < len(shared_strings) else ""
            elif cell_type == "inlineStr":
                value = "".join(text.text or "" for text in cell.findall(".//a:t", namespace))
            else:
                value = raw
            values[index] = value.strip()
        if values:
            max_index = max(values)
            table.append([values.get(i, "") for i in range(max_index + 1)])
    if not table:
        return []
    headers = table[0]
    return [
        {headers[index]: value for index, value in enumerate(row) if index < len(headers)}
        for row in table[1:]
        if any(str(value).strip() for value in row)
    ]


def read_xlsx_shared_strings(archive: zipfile.ZipFile) -> list[str]:
    if "xl/sharedStrings.xml" not in archive.namelist():
        return []
    root = ElementTree.fromstring(archive.read("xl/sharedStrings.xml"))
    namespace = {"a": "http://schemas.openxmlformats.org/spreadsheetml/2006/main"}
    values = []
    for item in root.findall("a:si", namespace):
        values.append("".join(node.text or "" for node in item.findall(".//a:t", namespace)))
    return values


def xlsx_column_index(cell_reference: str) -> int:
    letters = "".join(ch for ch in cell_reference if ch.isalpha()).upper()
    index = 0
    for letter in letters:
        index = index * 26 + (ord(letter) - ord("A") + 1)
    return max(index - 1, 0)


def case_from_upload_row(row: dict[str, str], source_system: str) -> CaseCreate:
    normalized = {normalized_header(str(key)): str(value).strip() for key, value in row.items() if value is not None}

    def pick(field: str, default: str | None = None) -> str | None:
        keys = {normalized_header(field), *HEADER_ALIASES.get(field, set())}
        for key in keys:
            if normalized.get(key):
                return normalized[key]
        return default

    fir_number = pick("fir_number")
    district = pick("district")
    complainant_name = pick("complainant_name", "Unknown complainant")
    complainant_phone = pick("complainant_phone", "0000000000")
    victim_name = pick("victim_name", "Unknown victim")
    suspect_name = pick("suspect_name", "Unknown")
    summary = pick("summary")
    if not fir_number:
        raise ValueError("missing FIR number")
    if not district:
        raise ValueError("missing district")
    if not summary:
        raise ValueError("missing summary/details")
    return CaseCreate(
        fir_number=fir_number,
        year=parse_int(pick("year"), datetime.now(timezone.utc).year),
        district=district,
        status=parse_status(pick("status")),
        case_type=pick("case_type"),
        modus_operandi=pick("modus_operandi"),
        incident_at=parse_datetime_value(pick("incident_at")),
        complainant_name=complainant_name or "Unknown complainant",
        complainant_phone=complainant_phone or "0000000000",
        victim_name=victim_name or "Unknown victim",
        victim_age=parse_optional_int(pick("victim_age")),
        victim_gender=pick("victim_gender"),
        suspect_name=suspect_name or "Unknown",
        suspect_age=parse_optional_int(pick("suspect_age")),
        suspect_gender=pick("suspect_gender"),
        summary=summary,
        sensitivity=parse_sensitivity(pick("sensitivity")),
        latitude=parse_optional_float(pick("latitude")),
        longitude=parse_optional_float(pick("longitude")),
        source_system=source_system,
        source_record_id=pick("source_record_id"),
    )


def parse_status(value: str | None) -> str:
    text = normalized_header(value or "")
    if text in {"closed", "close"}:
        return "closed"
    if text in {"underreview", "review", "investigation", "investigating"}:
        return "under_review"
    return "open"


def parse_sensitivity(value: str | None) -> str:
    return "restricted" if normalized_header(value or "") in {"restricted", "sensitive", "confidential"} else "standard"


def parse_int(value: str | None, default: int) -> int:
    try:
        return int(float(str(value).strip()))
    except (TypeError, ValueError):
        return default


def parse_optional_int(value: str | None) -> int | None:
    if not value:
        return None
    parsed = parse_int(value, -1)
    return parsed if parsed >= 0 else None


def parse_optional_float(value: str | None) -> float | None:
    if not value:
        return None
    try:
        return float(str(value).strip())
    except ValueError:
        return None


def parse_datetime_value(value: str | None) -> datetime | None:
    if not value:
        return None
    text = str(value).strip()
    for fmt in ("%Y-%m-%d", "%d-%m-%Y", "%m/%d/%Y", "%d/%m/%Y", "%Y-%m-%d %H:%M:%S"):
        try:
            return datetime.strptime(text, fmt).replace(tzinfo=timezone.utc)
        except ValueError:
            pass
    try:
        return datetime.fromisoformat(text.replace("Z", "+00:00"))
    except ValueError:
        return None


MAX_EXTRACTED_TEXT_CHARS = 12000
MAX_EXTRACTED_SUMMARY_CHARS = 1600


def clean_extracted_text(text: str) -> str:
    lines = [re.sub(r"\s+", " ", line).strip() for line in text.splitlines()]
    return "\n".join(line for line in lines if line).strip()


def truncate_text(text: str, limit: int) -> str:
    if len(text) <= limit:
        return text
    return f"{text[: max(0, limit - 1)].rstrip()}."


def extract_docx_text(content: bytes) -> str | None:
    try:
        from docx import Document

        document = Document(io.BytesIO(content))
        parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        text = clean_extracted_text("\n".join(parts))
        if text:
            return text
    except Exception:
        pass
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as archive:
            if "word/document.xml" not in archive.namelist():
                return None
            document_xml = archive.read("word/document.xml")
        root = ElementTree.fromstring(document_xml)
        namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        text = clean_extracted_text("\n".join(node.text or "" for node in root.findall(".//w:t", namespace)))
        return text or None
    except zipfile.BadZipFile:
        return None
    except ElementTree.ParseError:
        try:
            decoded = document_xml.decode("utf-8", errors="ignore")
        except UnboundLocalError:
            return None
        text_runs = re.findall(r"<w:t\b[^>]*>(.*?)</w:t>", decoded, flags=re.IGNORECASE | re.DOTALL)
        text = clean_extracted_text("\n".join(html.unescape(re.sub(r"<[^>]+>", "", run)) for run in text_runs))
        return text or None


def extract_pdf_text(content: bytes) -> tuple[str | None, str]:
    try:
        from pypdf import PdfReader
    except Exception:
        return None, "pdf_extractor_unavailable"
    try:
        reader = PdfReader(io.BytesIO(content))
        if reader.is_encrypted:
            try:
                reader.decrypt("")
            except Exception:
                return None, "pdf_encrypted"
        pages = []
        for page in reader.pages[:80]:
            pages.append(page.extract_text() or "")
        text = clean_extracted_text("\n".join(pages))
        return (text or None), ("extracted" if text else "no_extractable_text")
    except Exception:
        return None, "pdf_parse_failed"


def extract_legacy_doc_text(content: bytes) -> str | None:
    candidates = []
    for encoding in ("utf-8", "utf-16le", "latin-1"):
        decoded = content.decode(encoding, errors="ignore")
        decoded = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]+", " ", decoded)
        runs = re.findall(r"[^\x00-\x1f]{4,}", decoded)
        text = clean_extracted_text("\n".join(runs))
        if text:
            candidates.append(text)
    if not candidates:
        return None
    return max(candidates, key=len)


DOCUMENT_FIELD_PATTERNS = [
    ("fir_number", "FIR number", r"\b(?:FIR|F\.I\.R\.?)\s*(?:number|no\.?|#)?\s*[:\-]?\s*([A-Z0-9][A-Z0-9/_.-]{2,})"),
    ("case_number", "Case number", r"\b(?:case|file)\s*(?:number|no\.?|#)?\s*[:\-]?\s*([A-Z0-9][A-Z0-9/_.-]{2,})"),
    ("complainant", "Complainant", r"\b(?:complainant|complaint\s+by|filed\s+by)\s*[:\-]\s*([A-Za-z .'\-]{3,80})"),
    ("suspect", "Accused/Suspect", r"\b(?:accused|suspect|a\d)\s*[:\-]\s*([A-Za-z .'\-]{3,80})"),
    ("victim", "Victim", r"\bvictim\s*[:\-]\s*([A-Za-z .'\-]{3,80})"),
    ("police_station", "Police station", r"\bpolice\s+station\s*[:\-]\s*([A-Za-z .'\-]{3,100})"),
    ("status", "Status", r"\bstatus\s*[:\-]\s*([A-Za-z _-]{3,40})"),
]

KANNADA_FIELD_LABELS = {
    "fir_number": "ಎಫ್‌ಐಆರ್ ಸಂಖ್ಯೆ",
    "case_number": "ಪ್ರಕರಣ ಸಂಖ್ಯೆ",
    "complainant": "ದೂರುದಾರರು",
    "suspect": "ಆರೋಪಿತ / ಶಂಕಿತ",
    "victim": "ಬಾಧಿತರು",
    "police_station": "ಪೊಲೀಸ್ ಠಾಣೆ",
    "status": "ಸ್ಥಿತಿ",
}

KANNADA_STATUS_VALUES = {
    "open": "ತೆರೆದಿದೆ",
    "closed": "ಮುಚ್ಚಲಾಗಿದೆ",
    "under review": "ಪರಿಶೀಲನೆಯಲ್ಲಿದೆ",
    "under_review": "ಪರಿಶೀಲನೆಯಲ್ಲಿದೆ",
}


FIR_TEMPLATE_VERSION = "ksp-official-fir-copy-v1"
FIR_REQUIRED_FIELDS = {
    "district",
    "crime_no",
    "fir_date",
    "police_station",
    "act_section",
    "place_of_occurrence",
    "complainant_name",
    "fir_contents",
    "action_taken",
}


def fir_layout_field(key: str, label: str, section: str, page: int, row: int, column: int, x: int, y: int, width: int, align: str = "left") -> dict[str, Any]:
    return {
        "key": key,
        "label": label,
        "section": section,
        "page": page,
        "row": row,
        "column": column,
        "x": x,
        "y": y,
        "w": width,
        "align": align,
        "required": key in FIR_REQUIRED_FIELDS,
    }


FIR_LAYOUT_FIELDS: list[dict[str, Any]] = [
    fir_layout_field("document_title", "FIRST INFORMATION REPORT", "header", 1, 1, 1, 36, 36, 520, "center"),
    fir_layout_field("agency", "KARNATAKA STATE POLICE", "header", 1, 2, 1, 36, 58, 520, "center"),
    fir_layout_field("court", "Before the Honourable Court", "header", 1, 3, 1, 36, 82, 520, "center"),
    fir_layout_field("legal_basis", "Legal Basis", "header", 1, 4, 1, 36, 104, 520, "center"),
    fir_layout_field("district", "1. District", "case_header", 1, 1, 1, 36, 132, 170),
    fir_layout_field("crime_no", "Crime No", "case_header", 1, 1, 2, 212, 132, 150),
    fir_layout_field("fir_date", "FIR Date", "case_header", 1, 1, 3, 368, 132, 160),
    fir_layout_field("circle_subdivision", "Circle/Sub Division", "case_header", 1, 2, 1, 36, 156, 230),
    fir_layout_field("police_station", "PS", "case_header", 1, 2, 2, 272, 156, 170),
    fir_layout_field("act_section", "Act & Section", "case_header", 1, 3, 1, 36, 180, 492),
    fir_layout_field("information_received_at_ps", "2.(b) Information received at the PS", "information", 1, 1, 1, 36, 220, 260),
    fir_layout_field("written_oral", "Written/Oral", "information", 1, 1, 2, 302, 220, 100),
    fir_layout_field("occurrence_from_date", "From Date", "occurrence", 1, 1, 1, 36, 254, 112),
    fir_layout_field("occurrence_to_date", "To Date", "occurrence", 1, 1, 2, 154, 254, 112),
    fir_layout_field("occurrence_from_time", "From Time", "occurrence", 1, 1, 3, 272, 254, 108),
    fir_layout_field("occurrence_to_time", "To Time", "occurrence", 1, 1, 4, 386, 254, 108),
    fir_layout_field("occurrence_day", "(a) Occurence of Offence Day", "occurrence", 1, 2, 1, 36, 278, 170),
    fir_layout_field("delay_reason", "(c) Reasons for Delay", "occurrence", 1, 2, 2, 212, 278, 200),
    fir_layout_field("gd_entry", "(d) General Diary reference Entry No. & Time", "occurrence", 1, 3, 1, 36, 302, 330),
    fir_layout_field("place_of_occurrence", "4.(a) Place of occurence with full address", "place", 1, 1, 1, 36, 342, 492),
    fir_layout_field("distance_from_ps", "(b) Distance from PS", "place", 1, 2, 1, 36, 390, 190),
    fir_layout_field("village", "(c) Village", "place", 1, 2, 2, 232, 390, 120),
    fir_layout_field("beat_name", "Beat Name", "place", 1, 2, 3, 358, 390, 150),
    fir_layout_field("complainant_name", "5. Complainant/Informant Name", "complainant", 1, 1, 1, 36, 440, 190),
    fir_layout_field("complainant_father_husband_name", "Father's/Husband's Name", "complainant", 1, 1, 2, 232, 440, 190),
    fir_layout_field("complainant_age", "(b) Age", "complainant", 2, 1, 1, 36, 70, 70),
    fir_layout_field("complainant_occupation", "(c) Occupation", "complainant", 2, 1, 2, 112, 70, 120),
    fir_layout_field("complainant_email", "(g) Email", "complainant", 2, 1, 3, 238, 70, 190),
    fir_layout_field("complainant_phone", "(h) Phone No.", "complainant", 2, 2, 1, 36, 104, 130),
    fir_layout_field("complainant_nationality", "(i) Nationality", "complainant", 2, 2, 2, 172, 104, 120),
    fir_layout_field("complainant_address", "(k) Address", "complainant", 2, 3, 1, 36, 138, 360),
    fir_layout_field("complainant_sex", "(l) Sex", "complainant", 2, 3, 2, 402, 138, 80),
    fir_layout_field("seen_or_heard", "(m) Seen occurrence or heard of it", "complainant", 2, 4, 1, 36, 172, 300),
    fir_layout_field("fir_contents", "10. F.I.R Contents", "narrative", 3, 1, 1, 36, 230, 492),
    fir_layout_field("read_over_copy_given", "11.(b) Read over and copy given", "action", 3, 1, 1, 36, 660, 230),
    fir_layout_field("action_taken", "11.(a) Action Taken", "action", 3, 1, 2, 272, 660, 160),
    fir_layout_field("spot_visit_reason", "11.(c) Reason / spot visit note", "action", 3, 2, 1, 36, 690, 492),
    fir_layout_field("dispatch_to_court_datetime", "13. Date and time of dispatch to the Court", "dispatch", 4, 1, 1, 36, 84, 220),
    fir_layout_field("sho_name", "Name", "dispatch", 4, 2, 1, 36, 118, 210),
    fir_layout_field("carrier_name", "14. Name of PC/HC who carried the FIR to the Court", "dispatch", 4, 2, 2, 252, 118, 250),
    fir_layout_field("read_over_datetime", "Read Over and Found to be correct Date and Time", "dispatch", 4, 3, 1, 36, 152, 260),
    fir_layout_field("copy_to", "Copy to", "dispatch", 4, 4, 1, 36, 186, 300),
]

KANNADA_FIELD_LABELS.update(
    {
        "crime_no": "ಎಫ್‌ಐಆರ್ ಸಂಖ್ಯೆ",
        "fir_date": "ಎಫ್‌ಐಆರ್ ದಿನಾಂಕ",
        "complainant_name": "ದೂರುದಾರರು",
        "accused_summary": "ಆರೋಪಿತ / ಶಂಕಿತ",
        "place_of_occurrence": "ಘಟನೆಯ ಸ್ಥಳ",
        "action_taken": "ಕೈಗೊಂಡ ಕ್ರಮ",
    }
)


def detected_document_field_pairs(text: str) -> list[dict[str, str]]:
    strict = build_strict_fir_metadata(text, "", "")
    fields = [
        {"key": key, "label": label, "value": value}
        for key, label in [
            ("crime_no", "FIR number"),
            ("district", "District"),
            ("fir_date", "FIR date"),
            ("police_station", "Police station"),
            ("complainant_name", "Complainant"),
            ("accused_summary", "Accused/Suspect"),
            ("place_of_occurrence", "Place of occurrence"),
            ("action_taken", "Action taken"),
        ]
        if (value := str(strict.get("values", {}).get(key) or "").strip())
    ]
    if fields:
        return fields[:8]
    for key, label, pattern in DOCUMENT_FIELD_PATTERNS:
        match = re.search(pattern, text, flags=re.IGNORECASE)
        if match:
            value = re.sub(r"\s+", " ", match.group(1)).strip(" .;-")
            if value:
                fields.append({"key": key, "label": label, "value": value})
    return fields[:8]


def detected_document_fields(text: str) -> list[str]:
    return [f"{field['label']}: {field['value']}" for field in detected_document_field_pairs(text)]


def kannada_field_value(field: dict[str, str]) -> str:
    value = field["value"]
    if field["key"] == "status":
        return KANNADA_STATUS_VALUES.get(value.casefold(), value)
    return value


def compact_for_fir(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def clean_fir_value(value: str | None, limit: int = 4000) -> str:
    if not value:
        return ""
    cleaned = re.sub(r"\s+", " ", value.replace("\x00", " ")).strip(" .;:-")
    return truncate_text(cleaned, limit)


def fir_regex_value(text: str, *patterns: str, limit: int = 4000) -> str:
    source = compact_for_fir(text)
    for pattern in patterns:
        match = re.search(pattern, source, flags=re.IGNORECASE | re.DOTALL)
        if match:
            return clean_fir_value(match.group(1), limit)
    return ""


def extract_occurrence_window(text: str) -> dict[str, str]:
    source = compact_for_fir(text)
    match = re.search(r"\b3\.\s*([0-9]{2}/[0-9]{2}/[0-9]{4})([0-9]{2}/[0-9]{2}/[0-9]{4})([0-9]{2}:[0-9]{2}:[0-9]{2})([0-9]{2}:[0-9]{2}:[0-9]{2})", source)
    if not match:
        return {}
    return {
        "occurrence_from_date": match.group(1),
        "occurrence_to_date": match.group(2),
        "occurrence_from_time": match.group(3),
        "occurrence_to_time": match.group(4),
    }


def extract_complainant_age_occupation(text: str) -> dict[str, str]:
    source = compact_for_fir(text)
    match = re.search(r"Age\s*\(c\)\s*Occupation\s*([0-9]{1,3})([A-Za-z][A-Za-z .'-]{2,60}?)(?=:{2,}|\(h\)|Phone|Nationality)", source, flags=re.IGNORECASE)
    if not match:
        return {}
    return {
        "complainant_age": clean_fir_value(match.group(1), 20),
        "complainant_occupation": clean_fir_value(match.group(2), 80),
    }


def split_compact_complainant_name(raw: str) -> tuple[str, str]:
    raw = clean_fir_value(raw, 120)
    if not raw:
        return "", ""
    if "/" in raw:
        name, father_husband = raw.split("/", 1)
        return clean_fir_value(name, 80), clean_fir_value(father_husband, 80)
    parts = raw.split()
    if 2 <= len(parts) <= 4:
        return raw, ""
    if len(parts) > 4:
        return " ".join(parts[:2]), " ".join(parts[2:])
    match = re.match(r"^([a-z]{4,}?)([a-z]{4,})$", raw, flags=re.IGNORECASE)
    if match:
        return match.group(1), match.group(2)
    return raw, ""


def extract_accused_rows(text: str) -> list[dict[str, str]]:
    source = compact_for_fir(text)
    block = fir_regex_value(
        source,
        r"AgeSexPerson TypeName / Father Name / Caste / AddressOccupationTypeSl\.No\.(.*?)(?=6\.Details of known|7\.\s+Details of Victims)",
        limit=5000,
    )
    rows: list[dict[str, str]] = []
    for match in re.finditer(r"(\d+)\s+Accused(Adult|Minor|Child)?(Male|Female|Other)?(.+?)(?=\d+\s+Accused|$)", block, flags=re.IGNORECASE | re.DOTALL):
        raw_particulars = clean_fir_value(match.group(4), 800)
        alias_match = re.search(r"\((A\d+)\)", raw_particulars, flags=re.IGNORECASE)
        name = clean_fir_value(re.split(r"\s*/|,", raw_particulars, maxsplit=1)[0], 140)
        rows.append(
            {
                "sl_no": match.group(1),
                "person_type": "Accused",
                "age_class": clean_fir_value(match.group(2) or ""),
                "sex": clean_fir_value(match.group(3) or ""),
                "name": name,
                "accused_code": alias_match.group(1).upper() if alias_match else "",
                "particulars": raw_particulars,
            }
        )
    return rows


def extract_property_rows(text: str) -> list[dict[str, str]]:
    source = compact_for_fir(text)
    block = fir_regex_value(source, r"8\.\s+Particulars of Property stolen/involved with value.*?(Sl\.No\..*?)(?=9\.)", limit=1800)
    rows: list[dict[str, str]] = []
    total = fir_regex_value(source, r"Total Value of the property Stolen / Involved\s*:\s*([^0-9A-Za-z]*[0-9.,]+|[^9]*?)(?=9\.)", limit=120)
    if block or total:
        rows.append({"sl_no": "1", "property_type": "", "estimated_value": total, "item_description": clean_fir_value(block, 500)})
    return rows


def build_strict_fir_metadata(text: str | None, filename: str, extension: str) -> dict[str, Any]:
    source = compact_for_fir(text or "")
    values: dict[str, str] = {
        "document_title": fir_regex_value(source, r"\b(FIRST INFORMATION REPORT)\b", limit=80),
        "agency": fir_regex_value(source, r"(KARNATAKA STATE POLICE)", limit=80),
        "court": fir_regex_value(source, r"KARNATAKA STATE POLICE\s*(.*?)(?=\(Under Section|1\.District)", limit=220),
        "legal_basis": fir_regex_value(source, r"\((Under Section .*?)\)\s*1\.District", limit=220),
        "district": fir_regex_value(source, r"District\s*:\s*(.*?)(?=Crime No\s*:)", limit=120),
        "crime_no": fir_regex_value(source, r"Crime No\s*:\s*(.*?)(?=FIR Date\s*:)", r"\bFIR\s*(?:Number|No\.?|#)?\s*[:\-]\s*([A-Z0-9][A-Z0-9/_.-]{2,})", limit=80),
        "fir_date": fir_regex_value(source, r"FIR Date\s*:\s*([0-9]{2}/[0-9]{2}/[0-9]{4})", r"\bFIR Date\s*[:\-]\s*([0-9A-Za-z /:-]{6,30})", limit=40),
        "circle_subdivision": fir_regex_value(source, r"Circle/Sub Division\s*:\s*(.*?)(?=PS\s*:)", limit=160),
        "police_station": fir_regex_value(
            source,
            r"\bPS\s*:\s*(.*?)(?=Act\s*&\s*Section\s*:)",
            r"\bPolice Station\s*[:\-]\s*(.*?)(?=\b(?:Status|Complainant|Accused|Victim|FIR Number|Investigation note)\b|$)",
            limit=140,
        ),
        "act_section": fir_regex_value(source, r"Act\s*&\s*Section\s*:\s*(.*?)(?=2\.\(b\)|Information received at the PS)", limit=420),
        "information_received_at_ps": fir_regex_value(source, r"Information received at the PS\s*:\s*(.*?)(?=Written/Oral\s*:)", limit=80),
        "written_oral": fir_regex_value(source, r"Written/Oral\s*:\s*(.*?)(?=From Time\s*:|From Date\s*:|3\.)", limit=40),
        "occurrence_day": fir_regex_value(source, r"Occurence of Offence Day\s*:\s*(.*?)(?=\(c\)\s*Reasons)", limit=60),
        "delay_reason": fir_regex_value(source, r"Reasons for Delay in reporting by the Complainant / Informant\s*(.*?)(?=\(d\)\s*General Diary)", limit=160),
        "gd_entry": fir_regex_value(source, r"General Diary reference Entry No\. & Time\s*:\s*(.*?)(?=4\.\(a\)|4\.)", limit=90),
        "place_of_occurrence": fir_regex_value(source, r"Place of occurence with full address\s*(.*?)(?=\(b\)\s*Distance from PS)", limit=500),
        "distance_from_ps": fir_regex_value(source, r"Distance from PS\s*:\s*(.*?)(?=\(c\)\s*Village|Village\s*:)", limit=120),
        "village": fir_regex_value(source, r"Village\s*:\s*(.*?)(?=Beat Name\s*:)", limit=120),
        "beat_name": fir_regex_value(source, r"Beat Name\s*:\s*(.*?)(?=District\s*:|\(d\))", limit=120),
        "complainant_email": fir_regex_value(source, r"Email.*?Fax\s*([A-Z0-9._%+\-]+@[A-Z0-9.\-]+\.[A-Z]{2,})", limit=140),
        "complainant_phone": fir_regex_value(source, r"Phone No\.\s*:\s*([0-9]{6,15})", r"\bPhone\s*[:\-]\s*([0-9]{6,15})", limit=30),
        "complainant_nationality": fir_regex_value(source, r"Nationality\s*:\s*([^:]+?)(?=:\s*Date of Issue|Date of Issue)", limit=80),
        "complainant_address": fir_regex_value(source, r"\(k\)\s*Address\s*:\s*(.*?)(?=\(l\)\s*Sex)", limit=420),
        "complainant_sex": fir_regex_value(source, r"\(l\)\s*Sex\s*:\s*(.*?)(?=\(m\))", limit=40),
        "seen_or_heard": fir_regex_value(source, r"Whether complainant has seen the occurence or merely heard of it\s*:\s*(.*?)(?=AgeSexPerson TypeName|6\.Details)", limit=80),
        "fir_contents": fir_regex_value(source, r"10\.\s*F\.I\.R Contents.*?\)\s*(.*?)(?=11\.\s*\(b\)|11\.)", limit=10000),
        "read_over_copy_given": fir_regex_value(source, r"copy given to the complainant free of cost\?\s*:\s*(.*?)(?=\(a\)\s*Action Taken)", limit=40),
        "action_taken": fir_regex_value(source, r"\(a\)\s*Action Taken\s*:\s*(.*?)(?=\(c\)\s*If the Police Officer)", r"\bStatus\s*[:\-]\s*(.*?)(?=\b(?:Investigation note|Police Station|Victim|Accused|Complainant)\b|$)", limit=80),
        "spot_visit_reason": fir_regex_value(source, r"the reasons there of should be mentioned\s*:\s*(.*?)(?=12\.\s*Signature|13\.)", limit=300),
        "dispatch_to_court_datetime": fir_regex_value(source, r"13\.\s*Date and time of dispatch to the Court\s*:\s*(.*?)(?=Name\s*:|14\.)", limit=80),
        "sho_name": fir_regex_value(
            source,
            r"13\.\s*Date and time of dispatch to the Court\s*:.*?\bName\s*:\s*(.*?)(?=14\.\s*Name of PC/HC|Read Over)",
            limit=140,
        ),
        "carrier_name": fir_regex_value(source, r"14\.\s*Name of PC/HC who carried the FIR to the Court\s*:\s*(.*?)(?=Read Over and Found)", limit=160),
        "read_over_datetime": fir_regex_value(source, r"Read Over and Found to be correct Date and Time\s*:\s*(.*?)(?=Signature of the SHO|Copy to)", limit=80),
        "copy_to": fir_regex_value(source, r"Copy to\s*:\s*(.*)$", limit=160),
    }
    header_match = re.search(
        r"Circle/Sub Division\s*:\s*(.*?)PS\s*:\s*(.*?)(?=Act\s*&\s*Section\s*:)",
        source,
        flags=re.IGNORECASE | re.DOTALL,
    )
    if header_match:
        values["circle_subdivision"] = clean_fir_value(header_match.group(1), 160)
        values["police_station"] = clean_fir_value(header_match.group(2), 140)
    values.update(extract_occurrence_window(source))
    values.update(extract_complainant_age_occupation(source))
    complainant_raw = fir_regex_value(
        source,
        r"5\.Complainant/Informant\s*:\s*(.*?)(?=\(g\)\s*Email|\(f\)\s*Fax|--- page 2|AgeSexPerson)",
        r"\bComplainant\s*[:\-]\s*(.*?)(?=\b(?:Accused|Suspect|Victim|Police Station|Status|Investigation note)\b|$)",
        limit=140,
    )
    complainant_name, father_husband = split_compact_complainant_name(complainant_raw)
    values.setdefault("complainant_name", complainant_name)
    if complainant_name:
        values["complainant_name"] = values.get("complainant_name") or complainant_name
    if father_husband:
        values["complainant_father_husband_name"] = father_husband
    accused = extract_accused_rows(source)
    if not accused:
        fallback_suspect = fir_regex_value(
            source,
            r"\b(?:Accused|Suspect)\s*[:\-]\s*(.*?)(?=\b(?:Victim|Police Station|Status|Investigation note|Complainant|FIR Number)\b|$)",
            limit=140,
        )
        if fallback_suspect:
            accused = [{"sl_no": "1", "person_type": "Accused", "age_class": "", "sex": "", "name": fallback_suspect, "accused_code": "", "particulars": fallback_suspect}]
    values["accused_summary"] = ", ".join(row["name"] for row in accused if row.get("name"))
    victims = []
    fallback_victim = fir_regex_value(
        source,
        r"\bVictim\s*[:\-]\s*(.*?)(?=\b(?:Police Station|Status|Investigation note|Complainant|Accused|Suspect|FIR Number)\b|$)",
        limit=140,
    )
    if fallback_victim:
        victims.append({"sl_no": "1", "name": fallback_victim, "sex": "", "age": "", "address": "", "injury_type": "", "occupation": ""})
    property_rows = extract_property_rows(source)
    missing_required = [key for key in FIR_REQUIRED_FIELDS if not values.get(key)]
    status_text = "strict_valid" if not missing_required else "strict_missing_required"
    aligned_values = {field["key"]: values.get(field["key"], "") for field in FIR_LAYOUT_FIELDS}
    aligned_values["accused_summary"] = values.get("accused_summary", "")
    return {
        "template_version": FIR_TEMPLATE_VERSION,
        "source_filename": filename,
        "source_extension": extension,
        "status": status_text,
        "missing_required": missing_required,
        "values": aligned_values,
        "accused": accused,
        "victims": victims,
        "property_items": property_rows,
    }


def render_strict_fir_xml(metadata: dict[str, Any]) -> str:
    root = ElementTree.Element(
        "KSPFirstInformationReport",
        {
            "templateVersion": str(metadata.get("template_version") or FIR_TEMPLATE_VERSION),
            "status": str(metadata.get("status") or "strict_missing_required"),
            "alignmentSource": "Karnataka State Police FIR PDF sample",
        },
    )
    template = ElementTree.SubElement(root, "TemplateMetadata")
    ElementTree.SubElement(template, "DocumentTitle").text = "FIRST INFORMATION REPORT"
    ElementTree.SubElement(template, "Agency").text = "KARNATAKA STATE POLICE"
    ElementTree.SubElement(template, "Rule").text = "Keep section labels, page order, row/column metadata, and alignment attributes fixed. Change only Value and narrative content nodes."
    missing = ElementTree.SubElement(template, "MissingRequiredFields")
    for key in metadata.get("missing_required") or []:
        ElementTree.SubElement(missing, "Field", {"key": str(key)}).text = str(key)

    values = metadata.get("values") or {}
    sections = ElementTree.SubElement(root, "AlignedSections")
    for layout in FIR_LAYOUT_FIELDS:
        section = sections.find(f"./Section[@name='{layout['section']}']")
        if section is None:
            section = ElementTree.SubElement(sections, "Section", {"name": layout["section"]})
        field = ElementTree.SubElement(
            section,
            "Field",
            {
                "key": str(layout["key"]),
                "label": str(layout["label"]),
                "page": str(layout["page"]),
                "row": str(layout["row"]),
                "column": str(layout["column"]),
                "x": str(layout["x"]),
                "y": str(layout["y"]),
                "width": str(layout["w"]),
                "align": str(layout["align"]),
                "required": "true" if layout.get("required") else "false",
            },
        )
        ElementTree.SubElement(field, "Value").text = str(values.get(layout["key"]) or "")

    accused_table = ElementTree.SubElement(root, "AccusedTable", {"section": "6", "page": "2"})
    for row in metadata.get("accused") or []:
        row_node = ElementTree.SubElement(accused_table, "Accused", {"slNo": str(row.get("sl_no") or "")})
        for key in ["person_type", "age_class", "sex", "name", "accused_code", "particulars"]:
            ElementTree.SubElement(row_node, key).text = str(row.get(key) or "")

    victim_table = ElementTree.SubElement(root, "VictimTable", {"section": "7", "page": "3"})
    for row in metadata.get("victims") or []:
        row_node = ElementTree.SubElement(victim_table, "Victim", {"slNo": str(row.get("sl_no") or "")})
        for key in ["name", "sex", "age", "address", "injury_type", "occupation"]:
            ElementTree.SubElement(row_node, key).text = str(row.get(key) or "")

    property_table = ElementTree.SubElement(root, "PropertyTable", {"section": "8", "page": "3"})
    for row in metadata.get("property_items") or []:
        row_node = ElementTree.SubElement(property_table, "Property", {"slNo": str(row.get("sl_no") or "")})
        for key in ["property_type", "estimated_value", "item_description"]:
            ElementTree.SubElement(row_node, key).text = str(row.get(key) or "")

    ElementTree.indent(root, space="  ")
    return "<?xml version=\"1.0\" encoding=\"UTF-8\"?>\n" + ElementTree.tostring(root, encoding="unicode")


def summarize_extracted_document_kn(text: str | None, filename: str, extension: str) -> str | None:
    if not text:
        return None
    fields = detected_document_field_pairs(text)
    sentences = [
        sentence.strip()
        for sentence in re.split(r"(?<=[.!?])\s+|\n+", text)
        if len(sentence.strip()) >= 24
    ]
    parts = [f"{filename} ({extension}) ಕಡತದಿಂದ ಸ್ವಯಂಚಾಲಿತವಾಗಿ ಪಠ್ಯ ಹೊರತೆಗೆದು ಸಾರಾಂಶ ಸಿದ್ಧಪಡಿಸಲಾಗಿದೆ."]
    if fields:
        field_text = "; ".join(
            f"{KANNADA_FIELD_LABELS.get(field['key'], field['label'])}: {kannada_field_value(field)}"
            for field in fields
        )
        parts.append(f"ಗುರುತಿಸಿದ ಕ್ಷೇತ್ರಗಳು: {field_text}.")
    if sentences:
        parts.append(f"ದಾಖಲೆಯ ಮುಖ್ಯ ಅಂಶಗಳು (ಮೂಲ ಪಠ್ಯ): {' '.join(sentences[:3])}")
    parts.append("ಇದು ಅಪ್‌ಲೋಡ್ ಮಾಡಿದ ಕಡತದಿಂದ ಸ್ವಯಂಚಾಲಿತವಾಗಿ ಹೊರತೆಗೆದ ಮಾಹಿತಿ; ಮೂಲ ದಾಖಲೆಯೊಂದಿಗೆ ಪರಿಶೀಲಿಸಿ.")
    return truncate_text(" ".join(parts), MAX_EXTRACTED_SUMMARY_CHARS)


def summarize_extracted_document(text: str | None, filename: str, extension: str) -> str | None:
    if not text:
        return None
    fields = detected_document_fields(text)
    sentences = [
        sentence.strip()
        for sentence in re.split(r"(?<=[.!?])\s+|\n+", text)
        if len(sentence.strip()) >= 24
    ]
    summary_parts = [f"Extracted text summary for {filename} ({extension})."]
    if fields:
        summary_parts.append("Detected fields: " + "; ".join(fields) + ".")
    if sentences:
        summary_parts.append("Document gist: " + " ".join(sentences[:4]))
    summary_parts.append("Extraction is automated from the uploaded file and must be verified against the original document.")
    return truncate_text(" ".join(summary_parts), MAX_EXTRACTED_SUMMARY_CHARS)


def extract_document_info(content: bytes, extension: str, filename: str) -> dict[str, Any]:
    text: str | None = None
    status_text = "not_applicable"
    if extension == ".docx":
        text = extract_docx_text(content)
        status_text = "extracted" if text else "docx_no_extractable_text"
    elif extension == ".pdf":
        text, status_text = extract_pdf_text(content)
    elif extension == ".doc":
        text = extract_legacy_doc_text(content)
        status_text = "legacy_doc_best_effort" if text else "legacy_doc_no_extractable_text"
    text = clean_extracted_text(text or "")
    preview = truncate_text(text, MAX_EXTRACTED_TEXT_CHARS) if text else None
    fir_metadata = build_strict_fir_metadata(preview or "", filename, extension) if preview and extension in DOCUMENT_EXTENSIONS else {}
    fir_xml = render_strict_fir_xml(fir_metadata) if fir_metadata else None
    return {
        "preview": preview,
        "summary": summarize_extracted_document(preview, filename, extension),
        "summary_kn": summarize_extracted_document_kn(preview, filename, extension),
        "status": status_text,
        "char_count": len(text),
        "metadata": fir_metadata,
        "fir_xml": fir_xml,
        "fir_xml_status": fir_metadata.get("status") if fir_metadata else None,
    }


def can_access_file_upload(user: AuthenticatedUser, upload: dict[str, Any], db: Database) -> bool:
    if user.role in {"super_admin", "supervisor"}:
        return True
    if upload.get("uploaded_by") == user.id:
        return True
    case_id = upload.get("case_id")
    if not case_id:
        return role_has_permission(user.role, "case:*")
    case = db.get_case(str(case_id))
    return bool(case and can_access_case(user, case))


def public_file_upload(upload: dict[str, Any]) -> dict[str, Any]:
    item = {
        key: upload.get(key)
        for key in [
            "id",
            "uploaded_at",
            "uploaded_by",
            "original_filename",
            "content_type",
            "extension",
            "sha256",
            "size_bytes",
            "case_id",
            "upload_type",
            "parsed_case_count",
            "skipped_case_count",
            "status",
            "notes",
            "extracted_preview",
            "extracted_summary",
            "extracted_summary_kn",
            "extraction_status",
            "extracted_char_count",
            "fir_reconstruction_xml",
            "fir_reconstruction_status",
            "linked_fir_number",
            "linked_district",
        ]
    }
    try:
        item["extracted_metadata"] = json.loads(str(upload.get("extracted_metadata_json") or "{}"))
    except json.JSONDecodeError:
        item["extracted_metadata"] = {}
    item["preview_supported"] = file_preview_supported(upload)
    return item


def file_preview_supported(upload: dict[str, Any]) -> bool:
    extension = str(upload.get("extension") or "").lower()
    return extension in INLINE_PREVIEW_EXTENSIONS or bool(upload.get("extracted_preview"))


def resolve_upload_path(upload: dict[str, Any], settings: Settings) -> Path:
    stored = Path(str(upload.get("stored_path") or ""))
    upload_root = (Path(settings.database_path).parent / "uploads").resolve()
    candidates: list[Path]
    if not stored.is_absolute():
        candidates = [
            Path.cwd() / stored,
            Path(settings.database_path).parent / stored,
            upload_root / stored,
            upload_root / stored.name,
        ]
        stored = next((candidate for candidate in candidates if candidate.exists()), candidates[0])
    resolved = stored.resolve()
    if upload_root not in [resolved, *resolved.parents]:
        raise HTTPException(status_code=403, detail="File path outside upload store")
    if not resolved.exists() or not resolved.is_file():
        raise HTTPException(status_code=404, detail="Uploaded file not found")
    return resolved


def upload_media_type(upload: dict[str, Any]) -> str:
    extension = str(upload.get("extension") or "").lower()
    return PREVIEW_MIME_TYPES.get(extension) or str(upload.get("content_type") or "application/octet-stream")


def render_pdf_page_png(content: bytes, page: int) -> bytes | None:
    try:
        import fitz

        with fitz.open(stream=content, filetype="pdf") as document:
            if document.page_count <= 0:
                return None
            page_index = max(0, min(page - 1, document.page_count - 1))
            pdf_page = document.load_page(page_index)
            pixmap = pdf_page.get_pixmap(matrix=fitz.Matrix(1.8, 1.8), alpha=False)
            return pixmap.tobytes("png")
    except Exception:
        return None


def count_pdf_pages(content: bytes) -> int | None:
    try:
        import fitz

        with fitz.open(stream=content, filetype="pdf") as document:
            return max(1, int(document.page_count))
    except Exception:
        pass
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(content))
        return max(1, len(reader.pages))
    except Exception:
        return None


def upload_preview_page_count(upload: dict[str, Any], content: bytes) -> int:
    extension = str(upload.get("extension") or "").lower()
    if extension == ".pdf":
        return count_pdf_pages(content) or 1
    if extension in {".doc", ".docx"}:
        return 1
    return 1


def load_preview_font(size: int, bold: bool = False, kannada: bool = False):
    from PIL import ImageFont

    candidates = []
    if kannada:
        candidates.extend([
            "C:/Windows/Fonts/Nirmala.ttc",
            "Nirmala.ttc",
            "Nirmala.ttf",
            "NotoSansKannada-Regular.ttf",
        ])
    candidates.extend([
        "arialbd.ttf" if bold else "arial.ttf",
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "DejaVuSans-Bold.ttf" if bold else "DejaVuSans.ttf",
    ])
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size)
        except OSError:
            continue
    return ImageFont.load_default()


def draw_wrapped_text(draw, text: str, xy: tuple[int, int], font, fill: str, max_width: int, line_gap: int) -> int:
    x, y = xy
    if not text:
        return y
    font_size = int(getattr(font, "size", 16))
    avg_char_width = max(7, int(draw.textlength("ABCDEFGHIJKLMNOPQRSTUVWXYZ", font=font) / 26))
    wrap_width = max(24, max_width // avg_char_width)
    for paragraph in text.splitlines() or [text]:
        wrapped = textwrap.wrap(paragraph, width=wrap_width) or [""]
        for line in wrapped:
            draw.text((x, y), line, font=font, fill=fill)
            y += font_size + line_gap
        y += max(4, line_gap)
    return y


def render_text_document_png(upload: dict[str, Any]) -> bytes:
    from PIL import Image, ImageDraw

    width, height = 1240, 1754
    margin = 90
    image = Image.new("RGB", (width, height), "#f4f7f6")
    draw = ImageDraw.Draw(image)
    page_box = (44, 44, width - 44, height - 44)
    draw.rounded_rectangle(page_box, radius=22, fill="#ffffff", outline="#d7e4df", width=3)
    title_font = load_preview_font(34, bold=True)
    meta_font = load_preview_font(20)
    label_font = load_preview_font(22, bold=True)
    body_font = load_preview_font(22)
    kannada_label_font = load_preview_font(24, bold=True, kannada=True)
    kannada_body_font = load_preview_font(22, kannada=True)

    y = margin
    filename = str(upload.get("original_filename") or "Uploaded document")
    draw.text((margin, y), filename, font=title_font, fill="#14231f")
    y += 48
    meta = " | ".join(
        part
        for part in [
            str(upload.get("upload_type") or "file").replace("_", " "),
            str(upload.get("extraction_status") or "rendered"),
            f"{int(upload.get('extracted_char_count') or 0):,} chars extracted",
        ]
        if part
    )
    draw.text((margin, y), meta, font=meta_font, fill="#60716d")
    y += 48
    draw.line((margin, y, width - margin, y), fill="#d7e4df", width=2)
    y += 34

    summary = str(upload.get("extracted_summary") or "").strip()
    if summary:
        draw.text((margin, y), "English extraction summary", font=label_font, fill="#087f75")
        y += 34
        y = draw_wrapped_text(draw, truncate_text(summary, 950), (margin, y), body_font, "#24332f", width - (margin * 2), 8)
        y += 18

    summary_kn = str(upload.get("extracted_summary_kn") or "").strip()
    if not summary_kn and upload.get("extracted_preview"):
        summary_kn = summarize_extracted_document_kn(
            str(upload.get("extracted_preview") or ""),
            str(upload.get("original_filename") or "Uploaded document"),
            str(upload.get("extension") or ""),
        ) or ""
    if summary_kn:
        draw.rounded_rectangle((margin - 16, y - 8, width - margin + 16, min(height - 126, y + 260)), radius=12, fill="#f6fbf9")
        draw.text((margin, y), "ಕನ್ನಡ ಸಾರಾಂಶ", font=kannada_label_font, fill="#087f75")
        y += 38
        y = draw_wrapped_text(draw, truncate_text(summary_kn, 950), (margin, y), kannada_body_font, "#24332f", width - (margin * 2), 9)
        y += 24

    preview = str(upload.get("extracted_preview") or "").strip()
    if preview:
        draw.text((margin, y), "Extracted document text / ಹೊರತೆಗೆದ ಪಠ್ಯ", font=label_font, fill="#087f75")
        y += 34
        body_font_size = int(getattr(body_font, "size", 22))
        available_lines = max(8, (height - y - 110) // (body_font_size + 8))
        body_text = "\n".join(preview.splitlines()[:available_lines])
        draw_wrapped_text(draw, truncate_text(body_text, 5200), (margin, y), body_font, "#172421", width - (margin * 2), 8)
    else:
        draw.text((margin, y), "No extractable text was found. Open the original file for visual inspection.", font=body_font, fill="#60716d")

    footer = "PNG page generated from the uploaded file for preview. Verify against the original evidence file."
    draw.text((margin, height - 78), footer, font=meta_font, fill="#60716d")
    buffer = io.BytesIO()
    image.save(buffer, format="PNG", optimize=True)
    return buffer.getvalue()


def render_upload_png(upload: dict[str, Any], content: bytes, page: int) -> bytes:
    extension = str(upload.get("extension") or "").lower()
    if extension == ".pdf":
        rendered = render_pdf_page_png(content, page)
        if rendered:
            return rendered
    return render_text_document_png(upload)


def b64url_encode(data: bytes) -> str:
    return base64.urlsafe_b64encode(data).decode("ascii").rstrip("=")


def b64url_decode(value: str) -> bytes:
    padding_needed = (-len(value)) % 4
    return base64.urlsafe_b64decode(value + ("=" * padding_needed))


def request_origin(request: Request) -> str:
    return f"{request.url.scheme}://{request.url.netloc}"


def request_rp_id(request: Request) -> str:
    return request.url.hostname or "localhost"


class CborReader:
    def __init__(self, data: bytes):
        self.data = data
        self.index = 0

    def read(self):
        if self.index >= len(self.data):
            raise ValueError("truncated CBOR")
        initial = self.data[self.index]
        self.index += 1
        major = initial >> 5
        additional = initial & 0x1F
        length = self._read_length(additional)
        if major == 0:
            return length
        if major == 1:
            return -1 - length
        if major == 2:
            return self._read_bytes(length)
        if major == 3:
            return self._read_bytes(length).decode("utf-8")
        if major == 4:
            return [self.read() for _ in range(length)]
        if major == 5:
            return {self.read(): self.read() for _ in range(length)}
        if major == 7:
            if additional == 20:
                return False
            if additional == 21:
                return True
            if additional in {22, 23}:
                return None
        raise ValueError("unsupported CBOR item")

    def _read_length(self, additional: int) -> int:
        if additional < 24:
            return additional
        if additional == 24:
            return self._read_int(1)
        if additional == 25:
            return self._read_int(2)
        if additional == 26:
            return self._read_int(4)
        if additional == 27:
            return self._read_int(8)
        raise ValueError("unsupported CBOR length")

    def _read_int(self, size: int) -> int:
        return int.from_bytes(self._read_bytes(size), "big")

    def _read_bytes(self, size: int) -> bytes:
        if self.index + size > len(self.data):
            raise ValueError("truncated CBOR bytes")
        value = self.data[self.index : self.index + size]
        self.index += size
        return value


def cbor_decode(data: bytes):
    reader = CborReader(data)
    value = reader.read()
    return value, reader.index


def parse_registration_auth_data(auth_data: bytes) -> dict[str, object]:
    if len(auth_data) < 37:
        raise ValueError("Authenticator data too short")
    flags = auth_data[32]
    sign_count = int.from_bytes(auth_data[33:37], "big")
    if not flags & 0x01:
        raise ValueError("User presence was not confirmed")
    if not flags & 0x04:
        raise ValueError("Fingerprint/user verification was not confirmed")
    if not flags & 0x40:
        raise ValueError("Attested credential data is missing")
    index = 37 + 16
    credential_length = int.from_bytes(auth_data[index : index + 2], "big")
    index += 2
    credential_id = auth_data[index : index + credential_length]
    index += credential_length
    public_key_cose = auth_data[index:]
    cbor_decode(public_key_cose)
    return {
        "flags": flags,
        "sign_count": sign_count,
        "credential_id": credential_id,
        "public_key_cose": public_key_cose,
    }


def parse_assertion_auth_data(auth_data: bytes) -> dict[str, int | bytes]:
    if len(auth_data) < 37:
        raise ValueError("Authenticator data too short")
    flags = auth_data[32]
    if not flags & 0x01:
        raise ValueError("User presence was not confirmed")
    if not flags & 0x04:
        raise ValueError("Fingerprint/user verification was not confirmed")
    return {
        "rp_id_hash": auth_data[:32],
        "flags": flags,
        "sign_count": int.from_bytes(auth_data[33:37], "big"),
    }


def cose_public_key(public_key_cose: bytes):
    cose, _ = cbor_decode(public_key_cose)
    if not isinstance(cose, dict):
        raise ValueError("Invalid COSE key")
    key_type = cose.get(1)
    alg = cose.get(3)
    if key_type == 2 and alg == -7:
        x = cose.get(-2)
        y = cose.get(-3)
        if not isinstance(x, bytes) or not isinstance(y, bytes):
            raise ValueError("Invalid EC COSE key")
        numbers = ec.EllipticCurvePublicNumbers(int.from_bytes(x, "big"), int.from_bytes(y, "big"), ec.SECP256R1())
        return numbers.public_key()
    if key_type == 3 and alg in {-257, -258, -259}:
        modulus = cose.get(-1)
        exponent = cose.get(-2)
        if not isinstance(modulus, bytes) or not isinstance(exponent, bytes):
            raise ValueError("Invalid RSA COSE key")
        numbers = rsa.RSAPublicNumbers(int.from_bytes(exponent, "big"), int.from_bytes(modulus, "big"))
        return numbers.public_key()
    raise ValueError("Unsupported biometric credential key type")


def verify_webauthn_client_data(client_data_json: bytes, *, expected_type: str, purpose: str, db: Database) -> tuple[dict, dict]:
    try:
        client_data = json.loads(client_data_json.decode("utf-8"))
    except (UnicodeDecodeError, json.JSONDecodeError) as exc:
        raise HTTPException(status_code=400, detail="Invalid WebAuthn client data") from exc
    if client_data.get("type") != expected_type:
        raise HTTPException(status_code=400, detail="Invalid WebAuthn response type")
    challenge = client_data.get("challenge")
    if not isinstance(challenge, str):
        raise HTTPException(status_code=400, detail="Missing WebAuthn challenge")
    challenge_row = db.consume_biometric_challenge(challenge, purpose)
    if not challenge_row:
        raise HTTPException(status_code=400, detail="Fingerprint challenge expired or already used")
    if client_data.get("origin") != challenge_row["origin"]:
        raise HTTPException(status_code=400, detail="Fingerprint origin mismatch")
    return client_data, challenge_row


def verify_rp_id_hash(auth_data: bytes, rp_id: str) -> None:
    expected = hashlib.sha256(rp_id.encode("utf-8")).digest()
    if auth_data[:32] != expected:
        raise HTTPException(status_code=400, detail="Fingerprint relying-party mismatch")


def add_security_middleware(app: FastAPI, settings: Settings) -> None:
    app.add_middleware(InMemoryRateLimitMiddleware, requests_per_minute=settings.rate_limit_per_minute)
    app.add_middleware(TrustedHostMiddleware, allowed_hosts=settings.allowed_host_list)
    if settings.cors_origin_list:
        app.add_middleware(
            CORSMiddleware,
            allow_origins=settings.cors_origin_list,
            allow_credentials=False,
            allow_methods=["GET", "POST", "PATCH", "DELETE"],
            allow_headers=["Authorization", "Content-Type", "X-Request-ID"],
        )


def prepare_database_path(settings: Settings) -> Path:
    database_path = Path(settings.database_path)
    bundled_path = Path(settings.bundled_database_path) if settings.bundled_database_path else None
    if not bundled_path:
        return database_path

    if not bundled_path.exists():
        raise RuntimeError(f"Bundled database not found: {bundled_path}")

    database_path.parent.mkdir(parents=True, exist_ok=True)
    should_copy = not database_path.exists()
    if not should_copy:
        source_stat = bundled_path.stat()
        target_stat = database_path.stat()
        should_copy = source_stat.st_size != target_stat.st_size
    if should_copy:
        shutil.copy2(bundled_path, database_path)
    return database_path


def create_app(settings: Settings | None = None, database: Database | None = None) -> FastAPI:
    settings = settings or get_settings()

    db = database or Database(prepare_database_path(settings))
    db.init_schema()
    db.seed_penal_codes_if_empty()
    if settings.demo_mode:
        seed_demo_data(db, demo_password=settings.demo_password)
    if settings.bootstrap_username and settings.bootstrap_password:
        db.bootstrap_user_if_empty(
            UserCreate(
                username=settings.bootstrap_username,
                password=settings.bootstrap_password,
                full_name=settings.bootstrap_full_name,
                role=settings.bootstrap_role,  # type: ignore[arg-type]
                district=settings.bootstrap_district,
            )
        )

    app = FastAPI(
        title="Secure Crime Intelligence API",
        version="0.1.0",
        docs_url="/docs" if not settings.is_production else None,
        redoc_url=None,
    )
    app.state.settings = settings
    app.state.db = db
    add_security_middleware(app, settings)
    static_dir = Path(__file__).resolve().parent / "static"
    if static_dir.exists():
        app.mount("/static", StaticFiles(directory=static_dir), name="static")
        app.mount("/assets", StaticFiles(directory=static_dir / "assets"), name="root-assets")
        app.mount("/vendor", StaticFiles(directory=static_dir / "vendor"), name="root-vendor")

    @app.middleware("http")
    async def secure_defaults(request: Request, call_next):
        request_id = request.headers.get("X-Request-ID") or str(uuid.uuid4())
        request.state.request_id = request_id

        content_length = request.headers.get("content-length")
        if content_length and int(content_length) > settings.max_request_bytes:
            return Response(
                content='{"detail":"Request body too large"}',
                status_code=413,
                media_type="application/json",
                headers={"X-Request-ID": request_id},
            )

        response = await call_next(request)
        response.headers["X-Request-ID"] = request_id
        response.headers["X-Content-Type-Options"] = "nosniff"
        response.headers["X-Frame-Options"] = "SAMEORIGIN"
        response.headers["Referrer-Policy"] = "no-referrer"
        response.headers["Permissions-Policy"] = "camera=(), microphone=(self), geolocation=()"
        response.headers["Cache-Control"] = "no-store"
        if request.url.scheme == "https":
            response.headers["Strict-Transport-Security"] = "max-age=31536000; includeSubDomains"
        return response

    @app.get("/health", response_model=HealthResponse, tags=["system"])
    def health() -> HealthResponse:
        return HealthResponse(status="ok", service="secure-crime-api")

    @app.get("/", include_in_schema=False)
    def frontend(request: Request):
        sensitive_params = {"username", "password", "token", "access_token", "jwt", "secret"}
        if any(name.lower() in sensitive_params for name in request.query_params):
            return RedirectResponse(url="/", status_code=status.HTTP_303_SEE_OTHER)
        index_path = static_dir / "index.html"
        if not index_path.exists():
            raise HTTPException(status_code=404, detail="Frontend not built")
        return FileResponse(index_path)

    def frontend_root_asset(asset_name: str):
        asset_path = static_dir / asset_name
        if not asset_path.exists():
            raise HTTPException(status_code=404, detail="Asset not found")
        return FileResponse(asset_path)

    app.add_api_route("/app.js", lambda: frontend_root_asset("app.js"), include_in_schema=False)
    app.add_api_route("/styles.css", lambda: frontend_root_asset("styles.css"), include_in_schema=False)
    app.add_api_route("/demo-data.json", lambda: frontend_root_asset("demo-data.json"), include_in_schema=False)

    dashboard_dir = Path.cwd() / "outputs" / "criminal_network"
    dashboard_files = {
        "criminal_network_dashboard.png",
        "criminal_nodes.csv",
        "criminal_edges.csv",
        "criminal_network_metrics.json",
        "link_predictions.csv",
        "shortest_paths.csv",
    }

    @app.get("/dashboards/criminal-network", include_in_schema=False)
    def criminal_network_dashboard() -> FileResponse:
        dashboard_path = dashboard_dir / "criminal_network_dashboard.html"
        if not dashboard_path.exists():
            raise HTTPException(
                status_code=404,
                detail="Criminal network dashboard has not been generated. Run scripts/criminal_network_analysis_dashboard.py.",
            )
        return FileResponse(dashboard_path, media_type="text/html")

    @app.get("/dashboards/criminal-network/files/{filename}", include_in_schema=False)
    def criminal_network_dashboard_file(filename: str) -> FileResponse:
        if filename not in dashboard_files:
            raise HTTPException(status_code=404, detail="Dashboard file not found")
        file_path = dashboard_dir / filename
        if not file_path.exists():
            raise HTTPException(status_code=404, detail="Dashboard file not generated")
        return FileResponse(file_path)

    @app.get("/modules", response_model=list[ModuleStatus], tags=["system"])
    def modules() -> list[dict]:
        return MODULES

    @app.get("/framework", tags=["system"])
    def framework() -> dict[str, Any]:
        return {
            "title": "Intelligent Conversational AI and Crime Analytics Platform",
            "mission": (
                "Enable investigators, analysts, supervisors, and policymakers to query crime records, "
                "discover relationships, and generate preventive intelligence with transparent evidence trails."
            ),
            "core_tasks": CORE_INTELLIGENCE_TASKS,
            "solution_framework": SOLUTION_FRAMEWORK,
            "governance": {
                "access": "Role-based, district-aware, sensitivity-aware access control",
                "traceability": "Audit log with hash-chain integrity checks",
                "safeguards": SAFEGUARDS,
            },
        }

    @app.post("/crime-data/import", response_model=CrimeDataImportResponse, tags=["analytics"])
    def import_crime_data(
        payload: CrimeDataImportRequest,
        request: Request,
        user: Annotated[AuthenticatedUser, Depends(require_permission("admin:*"))],
        db: Annotated[Database, Depends(get_db)],
    ) -> dict:
        path = Path(payload.path).expanduser()
        try:
            result = import_crime_csv(
                db,
                path,
                source_system=payload.source_system,
                limit=payload.limit,
                reset_source=payload.reset_source,
            )
        except FileNotFoundError as exc:
            raise HTTPException(status_code=404, detail=f"Crime data file not found: {path}") from exc
        write_audit_log(
            db,
            action="CRIME_DATA_IMPORT",
            resource_type="crime_data",
            status="SUCCESS",
            request=request,
            user=user,
            detail={
                "source_system": payload.source_system,
                "imported": result["imported"],
                "skipped": result["skipped"],
                "limit": payload.limit,
            },
        )
        return result

    @app.get("/crime-data/status", response_model=CrimeDataStatusResponse, tags=["analytics"])
    def crime_data_status(
        request: Request,
        user: Annotated[AuthenticatedUser, Depends(require_permission("analytics:read"))],
        db: Annotated[Database, Depends(get_db)],
    ) -> dict:
        response = db.crime_data_status()
        write_audit_log(
            db,
            action="CRIME_DATA_STATUS",
            resource_type="crime_data",
            status="SUCCESS",
            request=request,
            user=user,
            detail={"imported_count": response["imported_count"]},
        )
        return response

    @app.post("/auth/login", response_model=TokenResponse, tags=["auth"])
    def login(
        payload: LoginRequest,
        request: Request,
        settings: Annotated[Settings, Depends(get_app_settings)],
        db: Annotated[Database, Depends(get_db)],
    ) -> TokenResponse:
        user_row = db.get_user_by_username(payload.username)
        if not user_row or not user_row["is_active"]:
            write_audit_log(
                db,
                action="LOGIN_FAILED",
                resource_type="authentication",
                status="FAILED",
                request=request,
                detail={"username": payload.username, "reason": "invalid_credentials"},
            )
            raise HTTPException(status_code=401, detail="Invalid credentials")

        if not verify_password(payload.password, user_row["password_hash"]):
            write_audit_log(
                db,
                action="LOGIN_FAILED",
                resource_type="authentication",
                status="FAILED",
                request=request,
                detail={"username": payload.username, "reason": "invalid_credentials"},
            )
            raise HTTPException(status_code=401, detail="Invalid credentials")

        if password_needs_rehash(user_row["password_hash"]):
            db.update_password_hash(user_row["id"], hash_password(payload.password))

        public_user = public_user_from_row(user_row)
        token, expires_at, _jti = create_access_token(public_user, settings)
        write_audit_log(
            db,
            action="LOGIN_SUCCESS",
            resource_type="authentication",
            status="SUCCESS",
            request=request,
            user=AuthenticatedUser(**public_user.model_dump(), jti="issued"),
        )
        return TokenResponse(access_token=token, expires_at=expires_at, user=public_user)

    @app.post("/auth/logout", tags=["auth"])
    def logout(
        request: Request,
        credentials: Annotated[HTTPAuthorizationCredentials | None, Depends(bearer_scheme)],
        user: Annotated[AuthenticatedUser, Depends(get_current_user)],
        settings: Annotated[Settings, Depends(get_app_settings)],
        db: Annotated[Database, Depends(get_db)],
    ) -> dict[str, str]:
        if credentials is None:
            raise HTTPException(status_code=401, detail="Missing bearer token")
        decoded = jwt.decode(
            credentials.credentials,
            settings.jwt_secret,
            algorithms=["HS256"],
            issuer=settings.jwt_issuer,
            options={"verify_exp": False},
        )
        expires_at = datetime.fromtimestamp(int(decoded["exp"]), timezone.utc).isoformat()
        db.revoke_token(user.jti, expires_at)
        write_audit_log(
            db,
            action="LOGOUT",
            resource_type="authentication",
            status="SUCCESS",
            request=request,
            user=user,
        )
        return {"status": "revoked"}

    @app.post("/auth/fingerprint/register/options", tags=["auth"])
    def fingerprint_register_options(
        request: Request,
        user: Annotated[AuthenticatedUser, Depends(get_current_user)],
        db: Annotated[Database, Depends(get_db)],
    ) -> dict:
        challenge = b64url_encode(secrets.token_bytes(32))
        rp_id = request_rp_id(request)
        origin = request_origin(request)
        db.create_biometric_challenge(
            challenge=challenge,
            user_id=user.id,
            username=user.username,
            purpose="registration",
            rp_id=rp_id,
            origin=origin,
            expires_at=(datetime.now(timezone.utc) + timedelta(minutes=5)).isoformat(),
        )
        credentials = db.list_biometric_credentials_for_user(user.id)
        write_audit_log(
            db,
            action="FINGERPRINT_REGISTER_OPTIONS",
            resource_type="biometric",
            status="SUCCESS",
            request=request,
            user=user,
            detail={"credential_count": len(credentials)},
        )
        return {
            "publicKey": {
                "challenge": challenge,
                "rp": {"name": "KSP Secure Intelligence", "id": rp_id},
                "user": {
                    "id": b64url_encode(user.id.encode("utf-8")),
                    "name": user.username,
                    "displayName": user.full_name,
                },
                "pubKeyCredParams": [
                    {"type": "public-key", "alg": -7},
                    {"type": "public-key", "alg": -257},
                ],
                "timeout": 60000,
                "attestation": "none",
                "authenticatorSelection": {
                    "residentKey": "preferred",
                    "userVerification": "required",
                },
                "excludeCredentials": [
                    {
                        "type": "public-key",
                        "id": row["credential_id"],
                        "transports": json.loads(row.get("transports_json") or "[]"),
                    }
                    for row in credentials
                ],
            }
        }

    @app.post("/auth/fingerprint/register/verify", tags=["auth"])
    def fingerprint_register_verify(
        payload: dict,
        request: Request,
        user: Annotated[AuthenticatedUser, Depends(get_current_user)],
        db: Annotated[Database, Depends(get_db)],
    ) -> dict:
        response = payload.get("response") or {}
        client_data_json = b64url_decode(str(response.get("clientDataJSON") or ""))
        _, challenge_row = verify_webauthn_client_data(
            client_data_json,
            expected_type="webauthn.create",
            purpose="registration",
            db=db,
        )
        if challenge_row["user_id"] != user.id:
            raise HTTPException(status_code=403, detail="Fingerprint challenge user mismatch")
        attestation_object = b64url_decode(str(response.get("attestationObject") or ""))
        try:
            attestation, _ = cbor_decode(attestation_object)
            auth_data = attestation.get("authData") if isinstance(attestation, dict) else None
            if not isinstance(auth_data, bytes):
                raise ValueError("Missing authenticator data")
            verify_rp_id_hash(auth_data, challenge_row["rp_id"])
            parsed = parse_registration_auth_data(auth_data)
        except ValueError as exc:
            raise HTTPException(status_code=400, detail=str(exc)) from exc
        credential_id = b64url_encode(parsed["credential_id"])  # type: ignore[arg-type]
        raw_id = str(payload.get("rawId") or payload.get("id") or "")
        if raw_id and raw_id != credential_id:
            raise HTTPException(status_code=400, detail="Fingerprint credential id mismatch")
        transports = response.get("transports")
        if not isinstance(transports, list):
            transports = payload.get("transports")
        if not isinstance(transports, list):
            transports = []
        try:
            credential = db.create_biometric_credential(
                user_id=user.id,
                credential_id=credential_id,
                public_key_cose=parsed["public_key_cose"],  # type: ignore[arg-type]
                sign_count=int(parsed["sign_count"]),
                transports=transports,
                label=payload.get("label") or "Fingerprint device",
            )
        except sqlite3.IntegrityError as exc:
            raise HTTPException(status_code=409, detail="Fingerprint credential already registered") from exc
        write_audit_log(
            db,
            action="FINGERPRINT_REGISTER",
            resource_type="biometric",
            resource_id=credential["id"],
            status="SUCCESS",
            request=request,
            user=user,
            detail={"credential_id": credential_id},
        )
        return {"status": "registered", "credential_id": credential_id}

    @app.post("/auth/fingerprint/login/options", tags=["auth"])
    def fingerprint_login_options(
        payload: dict,
        request: Request,
        db: Annotated[Database, Depends(get_db)],
    ) -> dict:
        username = str(payload.get("username") or "").strip()
        if not username:
            raise HTTPException(status_code=400, detail="Username is required for fingerprint login")
        user_row = db.get_user_by_username(username)
        if not user_row or not user_row["is_active"]:
            raise HTTPException(status_code=401, detail="Fingerprint login is not enrolled")
        credentials = db.list_biometric_credentials_for_user(user_row["id"])
        if not credentials:
            raise HTTPException(status_code=401, detail="Fingerprint login is not enrolled")
        challenge = b64url_encode(secrets.token_bytes(32))
        rp_id = request_rp_id(request)
        db.create_biometric_challenge(
            challenge=challenge,
            user_id=user_row["id"],
            username=username,
            purpose="authentication",
            rp_id=rp_id,
            origin=request_origin(request),
            expires_at=(datetime.now(timezone.utc) + timedelta(minutes=5)).isoformat(),
        )
        write_audit_log(
            db,
            action="FINGERPRINT_LOGIN_OPTIONS",
            resource_type="biometric",
            status="SUCCESS",
            request=request,
            detail={"username": username, "credential_count": len(credentials)},
        )
        return {
            "publicKey": {
                "challenge": challenge,
                "rpId": rp_id,
                "timeout": 60000,
                "userVerification": "required",
                "allowCredentials": [
                    {
                        "type": "public-key",
                        "id": row["credential_id"],
                        "transports": json.loads(row.get("transports_json") or "[]"),
                    }
                    for row in credentials
                ],
            }
        }

    @app.post("/auth/fingerprint/login/verify", response_model=TokenResponse, tags=["auth"])
    def fingerprint_login_verify(
        payload: dict,
        request: Request,
        settings: Annotated[Settings, Depends(get_app_settings)],
        db: Annotated[Database, Depends(get_db)],
    ) -> TokenResponse:
        response = payload.get("response") or {}
        credential_id = str(payload.get("rawId") or payload.get("id") or "")
        credential = db.get_biometric_credential_by_credential_id(credential_id)
        if not credential:
            raise HTTPException(status_code=401, detail="Fingerprint credential not recognized")
        client_data_json = b64url_decode(str(response.get("clientDataJSON") or ""))
        _, challenge_row = verify_webauthn_client_data(
            client_data_json,
            expected_type="webauthn.get",
            purpose="authentication",
            db=db,
        )
        if challenge_row["user_id"] != credential["user_id"]:
            raise HTTPException(status_code=403, detail="Fingerprint challenge user mismatch")
        user_row = db.get_user_by_id(credential["user_id"])
        if not user_row or not user_row["is_active"]:
            raise HTTPException(status_code=401, detail="User disabled")
        auth_data = b64url_decode(str(response.get("authenticatorData") or ""))
        signature = b64url_decode(str(response.get("signature") or ""))
        verify_rp_id_hash(auth_data, challenge_row["rp_id"])
        try:
            assertion = parse_assertion_auth_data(auth_data)
            signed_bytes = auth_data + hashlib.sha256(client_data_json).digest()
            public_key = cose_public_key(bytes(credential["public_key_cose"]))
            if isinstance(public_key, ec.EllipticCurvePublicKey):
                public_key.verify(signature, signed_bytes, ec.ECDSA(SHA256()))
            elif isinstance(public_key, rsa.RSAPublicKey):
                public_key.verify(signature, signed_bytes, padding.PKCS1v15(), SHA256())
            else:
                raise ValueError("Unsupported public key")
        except (InvalidSignature, ValueError) as exc:
            write_audit_log(
                db,
                action="FINGERPRINT_LOGIN_FAILED",
                resource_type="biometric",
                status="FAILED",
                request=request,
                detail={"credential_id": credential_id},
            )
            raise HTTPException(status_code=401, detail="Fingerprint verification failed") from exc
        old_count = int(credential.get("sign_count") or 0)
        new_count = int(assertion["sign_count"])
        if old_count and new_count and new_count <= old_count:
            raise HTTPException(status_code=401, detail="Fingerprint authenticator counter replay detected")
        db.update_biometric_credential_use(credential_id, max(old_count, new_count))
        public_user = public_user_from_row(user_row)
        token, expires_at, _jti = create_access_token(public_user, settings)
        write_audit_log(
            db,
            action="FINGERPRINT_LOGIN_SUCCESS",
            resource_type="biometric",
            status="SUCCESS",
            request=request,
            user=AuthenticatedUser(**public_user.model_dump(), jti="issued"),
            detail={"credential_id": credential_id},
        )
        return TokenResponse(access_token=token, expires_at=expires_at, user=public_user)

    @app.get("/auth/me", response_model=UserPublic, tags=["auth"])
    def me(user: Annotated[AuthenticatedUser, Depends(get_current_user)]) -> UserPublic:
        return UserPublic(**user.model_dump(exclude={"jti"}))

    @app.get("/admin/users", response_model=list[UserPublic], tags=["admin"])
    def admin_list_users(
        request: Request,
        user: Annotated[AuthenticatedUser, Depends(require_permission("user:read"))],
        db: Annotated[Database, Depends(get_db)],
    ) -> list[dict]:
        users = db.list_users()
        write_audit_log(
            db,
            action="USER_LIST",
            resource_type="user",
            status="SUCCESS",
            request=request,
            user=user,
            detail={"returned": len(users)},
        )
        return [{**row, "is_active": bool(row["is_active"])} for row in users]

    @app.post("/admin/users", response_model=UserPublic, tags=["admin"])
    def admin_create_user(
        payload: UserCreate,
        request: Request,
        user: Annotated[AuthenticatedUser, Depends(require_permission("user:create"))],
        db: Annotated[Database, Depends(get_db)],
    ) -> dict:
        created = db.create_user(payload)
        write_audit_log(
            db,
            action="USER_CREATE",
            resource_type="user",
            resource_id=created["id"],
            status="SUCCESS",
            request=request,
            user=user,
            detail={"username": created["username"], "role": created["role"]},
        )
        return {**created, "is_active": bool(created["is_active"])}

    @app.patch("/admin/users/{user_id}", response_model=UserPublic, tags=["admin"])
    def admin_update_user(
        user_id: str,
        payload: UserUpdate,
        request: Request,
        user: Annotated[AuthenticatedUser, Depends(require_permission("user:update"))],
        db: Annotated[Database, Depends(get_db)],
    ) -> dict:
        updated = db.update_user(
            user_id,
            full_name=payload.full_name,
            role=payload.role,
            district=payload.district,
            is_active=payload.is_active,
        )
        if not updated:
            raise HTTPException(status_code=404, detail="User not found")
        write_audit_log(
            db,
            action="USER_UPDATE",
            resource_type="user",
            resource_id=user_id,
            status="SUCCESS",
            request=request,
            user=user,
            detail=payload.model_dump(exclude_none=True),
        )
        return {**updated, "is_active": bool(updated["is_active"])}

    @app.get("/admin/users/{user_id}/activity", tags=["admin"])
    def admin_user_activity(
        user_id: str,
        request: Request,
        user: Annotated[AuthenticatedUser, Depends(require_permission("audit:read"))],
        db: Annotated[Database, Depends(get_db)],
        limit: int = 50,
    ) -> dict:
        target = db.get_user_by_id(user_id)
        if not target:
            raise HTTPException(status_code=404, detail="User not found")
        events = db.list_user_activity(user_id, limit)
        write_audit_log(
            db,
            action="USER_ACTIVITY_READ",
            resource_type="user",
            resource_id=user_id,
            status="SUCCESS",
            request=request,
            user=user,
            detail={"returned": len(events), "target_username": target["username"]},
        )
        return {
            "user": {**target, "is_active": bool(target["is_active"])},
            "count": len(events),
            "events": events,
        }

    @app.post("/admin/users/{user_id}/reset-password", tags=["admin"])
    def admin_reset_password(
        user_id: str,
        payload: PasswordResetRequest,
        request: Request,
        user: Annotated[AuthenticatedUser, Depends(require_permission("user:reset_password"))],
        db: Annotated[Database, Depends(get_db)],
    ) -> dict:
        target = db.get_user_by_id(user_id)
        if not target:
            raise HTTPException(status_code=404, detail="User not found")
        db.update_password_hash(user_id, hash_password(payload.password))
        write_audit_log(
            db,
            action="USER_PASSWORD_RESET",
            resource_type="user",
            resource_id=user_id,
            status="SUCCESS",
            request=request,
            user=user,
        )
        return {"status": "password_reset"}

    @app.delete("/admin/users/{user_id}", tags=["admin"])
    def admin_delete_user(
        user_id: str,
        request: Request,
        user: Annotated[AuthenticatedUser, Depends(require_permission("user:delete"))],
        db: Annotated[Database, Depends(get_db)],
    ) -> dict:
        if user_id == user.id:
            raise HTTPException(status_code=400, detail="Super admin cannot delete the active session account")
        target = db.get_user_by_id(user_id)
        if not target:
            raise HTTPException(status_code=404, detail="User not found")
        deleted = db.delete_user(user_id)
        if not deleted:
            raise HTTPException(status_code=404, detail="User not found")
        write_audit_log(
            db,
            action="USER_DELETE",
            resource_type="user",
            resource_id=user_id,
            status="SUCCESS",
            request=request,
            user=user,
            detail={"username": target["username"], "mode": deleted["delete_mode"]},
        )
        return {
            "status": deleted["delete_mode"],
            "user_id": user_id,
            "username": target["username"],
        }

    @app.get("/admin/rate-limit-alerts", response_model=list[RateLimitAlert], tags=["admin"])
    def admin_rate_limit_alerts(
        request: Request,
        user: Annotated[AuthenticatedUser, Depends(require_permission("admin:*"))],
        db: Annotated[Database, Depends(get_db)],
        limit: int = 50,
        include_acknowledged: bool = False,
    ) -> list[dict]:
        alerts = db.list_rate_limit_alerts(limit, include_acknowledged=include_acknowledged)
        write_audit_log(
            db,
            action="RATE_LIMIT_ALERT_LIST",
            resource_type="rate_limit",
            status="SUCCESS",
            request=request,
            user=user,
            detail={"returned": len(alerts), "include_acknowledged": include_acknowledged},
        )
        return alerts

    @app.post("/admin/rate-limit-alerts/{alert_id}/acknowledge", response_model=RateLimitAlert, tags=["admin"])
    def admin_acknowledge_rate_limit_alert(
        alert_id: int,
        request: Request,
        user: Annotated[AuthenticatedUser, Depends(require_permission("admin:*"))],
        db: Annotated[Database, Depends(get_db)],
    ) -> dict:
        alert = db.acknowledge_rate_limit_alert(alert_id, user.username)
        if not alert:
            raise HTTPException(status_code=404, detail="Rate limit alert not found")
        write_audit_log(
            db,
            action="RATE_LIMIT_ALERT_ACKNOWLEDGE",
            resource_type="rate_limit",
            resource_id=str(alert_id),
            status="SUCCESS",
            request=request,
            user=user,
        )
        return alert

    @app.post("/translate", response_model=TranslationResponse, tags=["intelligence"])
    def translate(
        payload: TranslationRequest,
        request: Request,
        user: Annotated[AuthenticatedUser, Depends(require_permission("translation:use"))],
        db: Annotated[Database, Depends(get_db)],
    ) -> dict:
        response = translate_text(payload.text, payload.source_language, payload.target_language)
        write_audit_log(
            db,
            action="TRANSLATION",
            resource_type="translation",
            status="SUCCESS",
            request=request,
            user=user,
            detail={"source": payload.source_language, "target": payload.target_language},
        )
        return response

    @app.post("/conversations", response_model=ConversationRecord, tags=["intelligence"])
    def create_conversation(
        payload: ConversationCreate,
        request: Request,
        user: Annotated[AuthenticatedUser, Depends(require_permission("conversation:*"))],
        db: Annotated[Database, Depends(get_db)],
    ) -> dict:
        conversation = db.create_conversation(user.id, payload.title)
        write_audit_log(
            db,
            action="CONVERSATION_CREATE",
            resource_type="conversation",
            resource_id=conversation["id"],
            status="SUCCESS",
            request=request,
            user=user,
        )
        return conversation

    @app.post(
        "/conversations/{conversation_id}/messages",
        response_model=ConversationExchange,
        tags=["intelligence"],
    )
    def conversation_message(
        conversation_id: str,
        payload: ConversationMessageRequest,
        request: Request,
        user: Annotated[AuthenticatedUser, Depends(require_permission("conversation:*"))],
        db: Annotated[Database, Depends(get_db)],
    ) -> dict:
        conversation = db.get_conversation(conversation_id)
        if not conversation:
            raise HTTPException(status_code=404, detail="Conversation not found")
        if conversation["user_id"] != user.id:
            raise HTTPException(status_code=403, detail="Conversation access denied")

        previous_messages = db.list_conversation_messages(conversation_id)
        context_window = " ".join(
            message["content"]
            for message in previous_messages[-8:]
            if message["role"] == "user"
        )
        conversation_memory = build_conversation_memory(conversation_id, previous_messages)
        user_message = db.add_conversation_message(conversation_id, "user", payload.query)
        intelligence = build_intelligence_response(
            query=payload.query,
            language=payload.language,
            cases=db.list_cases(),
            user=user,
            include_sources=True,
            db=db,
            conversation_context=context_window,
            conversation_memory=conversation_memory,
        )
        assistant_message = db.add_conversation_message(
            conversation_id,
            "assistant",
            intelligence["answer"],
            metadata={
                "intent": intelligence["intent"],
                "language": detect_query_language(payload.query, payload.language),
                "selected_module": intelligence.get("selected_module"),
                "extracted_entities": intelligence.get("extracted_entities", {}),
                "conversation_memory": intelligence.get("conversation_memory", {}),
                "orchestration_trace": intelligence.get("orchestration_trace", []),
            },
        )
        conversation = db.get_conversation(conversation_id)
        write_audit_log(
            db,
            action="CONVERSATION_MESSAGE",
            resource_type="conversation",
            resource_id=conversation_id,
            status="SUCCESS",
            request=request,
            user=user,
            detail={"intent": intelligence["intent"]},
        )
        return {
            "conversation": conversation,
            "user_message": user_message,
            "assistant_message": assistant_message,
            "intelligence": intelligence,
        }

    @app.get("/conversations/{conversation_id}/export.pdf", tags=["intelligence"])
    def export_conversation_pdf(
        conversation_id: str,
        request: Request,
        user: Annotated[AuthenticatedUser, Depends(require_permission("conversation:*"))],
        db: Annotated[Database, Depends(get_db)],
    ) -> Response:
        conversation = db.get_conversation(conversation_id)
        if not conversation:
            raise HTTPException(status_code=404, detail="Conversation not found")
        if conversation["user_id"] != user.id:
            raise HTTPException(status_code=403, detail="Conversation access denied")

        messages = db.list_conversation_messages(conversation_id)
        pdf = render_conversation_pdf(conversation["title"], messages)
        write_audit_log(
            db,
            action="CONVERSATION_EXPORT_PDF",
            resource_type="conversation",
            resource_id=conversation_id,
            status="SUCCESS",
            request=request,
            user=user,
        )
        return Response(
            content=pdf,
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{conversation_id}.pdf"'},
        )

    @app.get("/cases", response_model=list[CaseRecord], tags=["cases"])
    def list_cases(
        request: Request,
        user: Annotated[AuthenticatedUser, Depends(require_permission("case:read"))],
        db: Annotated[Database, Depends(get_db)],
    ) -> list[dict]:
        visible = visible_cases(user, db.list_cases())
        write_audit_log(
            db,
            action="CASE_LIST",
            resource_type="case",
            status="SUCCESS",
            request=request,
            user=user,
            detail={"returned": len(visible)},
        )
        return [mask_case(case, user.role) for case in visible]

    @app.get("/cases/search", response_model=CaseSearchResponse, tags=["cases"])
    def search_cases(
        q: str,
        request: Request,
        user: Annotated[AuthenticatedUser, Depends(require_permission("search:case"))],
        db: Annotated[Database, Depends(get_db)],
        limit: int = 20,
    ) -> dict:
        started = perf_counter()
        raw_results = db.search_cases(q, limit)
        filtered = []
        for item in raw_results["results"]:
            case = item["case"]
            if can_access_case(user, case):
                filtered.append({"case": mask_case(case, user.role), "score": item["score"]})
            if len(filtered) >= max(1, min(limit, 50)):
                break
        elapsed_ms = (perf_counter() - started) * 1000
        write_audit_log(
            db,
            action="CASE_SEARCH",
            resource_type="case",
            status="SUCCESS",
            request=request,
            user=user,
            detail={"terms": raw_results["terms"], "returned": len(filtered), "elapsed_ms": round(elapsed_ms, 3)},
        )
        return {
            "query": q,
            "normalized_terms": raw_results["terms"],
            "result_count": len(filtered),
            "elapsed_ms": elapsed_ms,
            "results": filtered,
        }

    @app.post("/cases/import", tags=["cases"])
    def import_cases(
        payload: CaseImportRequest,
        request: Request,
        user: Annotated[AuthenticatedUser, Depends(require_permission("case:*"))],
        db: Annotated[Database, Depends(get_db)],
    ) -> dict:
        imported = 0
        for record in payload.records:
            record.source_system = record.source_system or payload.source_system
            db.create_case(record)
            imported += 1
        write_audit_log(
            db,
            action="CASE_IMPORT",
            resource_type="case",
            status="SUCCESS",
            request=request,
            user=user,
            detail={"source_system": payload.source_system, "imported": imported},
        )
        return {"imported": imported, "source_system": payload.source_system, "indexed": True}

    @app.post("/files/upload", tags=["cases"])
    async def upload_case_file(
        request: Request,
        user: Annotated[AuthenticatedUser, Depends(get_current_user)],
        db: Annotated[Database, Depends(get_db)],
        settings: Annotated[Settings, Depends(get_app_settings)],
        file: UploadFile = File(...),
        upload_type: str = Form("case_sheet"),
        case_id: str | None = Form(default=None),
        source_system: str = Form("uploaded-case-file"),
        auto_import: bool = Form(True),
    ) -> dict:
        if upload_type not in UPLOAD_TYPES:
            raise HTTPException(status_code=400, detail="upload_type must be case_sheet or fir_copy")
        original_filename = safe_upload_filename(file.filename or "upload")
        extension = Path(original_filename).suffix.lower()
        if extension not in UPLOAD_EXTENSIONS:
            raise HTTPException(status_code=400, detail="Unsupported file type. Use .csv, .xlsx, .xls, .pdf, .doc, or .docx")
        if upload_type == "case_sheet" and not role_has_permission(user.role, "case:*"):
            raise HTTPException(status_code=403, detail="Missing permission: case:*")
        if upload_type == "fir_copy" and not (
            role_has_permission(user.role, "case:*") or role_has_permission(user.role, "case:note")
        ):
            raise HTTPException(status_code=403, detail="Missing permission: case:note")
        if upload_type == "case_sheet" and extension not in CASE_SHEET_EXTENSIONS:
            raise HTTPException(status_code=400, detail="Case sheet uploads must be .csv, .xlsx, or .xls")
        if upload_type == "fir_copy" and extension not in DOCUMENT_EXTENSIONS:
            raise HTTPException(status_code=400, detail="FIR copy uploads must be .pdf, .doc, or .docx")

        linked_case = None
        if case_id:
            linked_case = db.get_case(case_id)
            if not linked_case:
                raise HTTPException(status_code=404, detail="Linked case not found")
            if not can_access_case(user, linked_case):
                raise HTTPException(status_code=403, detail="Case access denied")

        content = await file.read(settings.max_request_bytes + 1)
        await file.close()
        if len(content) > settings.max_request_bytes:
            raise HTTPException(status_code=413, detail="Uploaded file too large")

        sha256 = hashlib.sha256(content).hexdigest()
        upload_dir = Path(settings.database_path).parent / "uploads"
        upload_dir.mkdir(parents=True, exist_ok=True)
        stored_name = f"{uuid.uuid4()}{extension}"
        stored_path = upload_dir / stored_name
        stored_path.write_bytes(content)

        imported = 0
        skipped = 0
        notes: list[str] = []
        document_info = extract_document_info(content, extension, original_filename) if extension in DOCUMENT_EXTENSIONS else {
            "preview": None,
            "summary": None,
            "summary_kn": None,
            "status": "not_applicable",
            "char_count": 0,
            "metadata": {},
            "fir_xml": None,
            "fir_xml_status": None,
        }
        extracted_preview = document_info["preview"]
        extracted_summary = document_info["summary"]
        extracted_summary_kn = document_info["summary_kn"]
        extraction_status = document_info["status"]
        extracted_char_count = document_info["char_count"]
        extracted_metadata = document_info["metadata"]
        fir_reconstruction_xml = document_info["fir_xml"]
        fir_reconstruction_status = document_info["fir_xml_status"]
        if upload_type == "case_sheet" and auto_import:
            records, skipped, parse_notes = parse_case_sheet(content, extension, source_system)
            notes.extend(parse_notes)
            for record in records:
                try:
                    db.create_case(record)
                    imported += 1
                except sqlite3.IntegrityError:
                    skipped += 1
                    if len(notes) < 8:
                        notes.append(f"Duplicate FIR/year skipped: {record.fir_number}/{record.year}")
        elif upload_type == "case_sheet":
            notes.append("Case sheet was stored without importing rows because auto_import is disabled.")
        else:
            notes.append("FIR/document copy stored and hashed. Attach a case id to link it to a case record.")
            if linked_case:
                notes[-1] = f"FIR/document copy stored and linked to FIR {linked_case['fir_number']}."
            if extracted_summary:
                notes.append(f"Document text extracted and summarized ({extracted_char_count} character(s)).")
                if fir_reconstruction_status:
                    missing = extracted_metadata.get("missing_required") or []
                    suffix = f"; missing required: {', '.join(missing)}" if missing else ""
                    notes.append(f"Strict FIR XML reconstruction: {fir_reconstruction_status}{suffix}.")
            elif extraction_status != "not_applicable":
                notes.append(f"Document extraction status: {extraction_status}.")

        status_text = "stored"
        if upload_type == "case_sheet":
            status_text = "imported" if imported else "stored"
            if skipped:
                status_text = f"{status_text}_with_skips"

        upload = db.create_file_upload(
            uploaded_by=user.id,
            original_filename=original_filename,
            stored_path=str(stored_path.resolve()),
            content_type=file.content_type,
            extension=extension,
            sha256=sha256,
            size_bytes=len(content),
            case_id=case_id,
            upload_type=upload_type,
            parsed_case_count=imported,
            skipped_case_count=skipped,
            status=status_text,
            notes="; ".join(notes)[:2000],
            extracted_preview=extracted_preview,
            extracted_summary=extracted_summary,
            extracted_summary_kn=extracted_summary_kn,
            extraction_status=extraction_status,
            extracted_char_count=extracted_char_count,
            extracted_metadata=extracted_metadata,
            fir_reconstruction_xml=fir_reconstruction_xml,
            fir_reconstruction_status=fir_reconstruction_status,
        )
        write_audit_log(
            db,
            action="FILE_UPLOAD",
            resource_type="file_upload",
            resource_id=upload["id"],
            status="SUCCESS",
            request=request,
            user=user,
            detail={
                "filename": original_filename,
                "extension": extension,
                "upload_type": upload_type,
                "size_bytes": len(content),
                "sha256": sha256,
                "imported": imported,
                "skipped": skipped,
                "case_id": case_id,
                "extraction_status": extraction_status,
                "extracted_char_count": extracted_char_count,
                "fir_reconstruction_status": fir_reconstruction_status,
                "strict_missing_required": extracted_metadata.get("missing_required") if extracted_metadata else [],
            },
        )
        return {
            "upload_id": upload["id"],
            "filename": original_filename,
            "upload_type": upload_type,
            "size_bytes": len(content),
            "sha256": sha256,
            "status": status_text,
            "imported": imported,
            "skipped": skipped,
            "case_id": case_id,
            "notes": notes,
            "extracted_preview": extracted_preview,
            "extracted_summary": extracted_summary,
            "extracted_summary_kn": extracted_summary_kn,
            "extraction_status": extraction_status,
            "extracted_char_count": extracted_char_count,
            "extracted_metadata": extracted_metadata,
            "fir_reconstruction_xml": fir_reconstruction_xml,
            "fir_reconstruction_status": fir_reconstruction_status,
        }

    @app.get("/files/uploads", response_model=list[FileUploadRecord], tags=["cases"])
    def list_file_uploads(
        request: Request,
        user: Annotated[AuthenticatedUser, Depends(require_permission("case:read"))],
        db: Annotated[Database, Depends(get_db)],
        q: str = "",
        limit: int = 50,
    ) -> list[dict[str, Any]]:
        uploads = [
            public_file_upload(upload)
            for upload in db.list_file_uploads(q, limit)
            if can_access_file_upload(user, upload, db)
        ]
        write_audit_log(
            db,
            action="FILE_UPLOAD_LIST",
            resource_type="file_upload",
            status="SUCCESS",
            request=request,
            user=user,
            detail={"returned": len(uploads), "query": q[:80]},
        )
        return uploads

    @app.get("/files/uploads/{upload_id}", response_model=FileUploadRecord, tags=["cases"])
    def get_file_upload(
        upload_id: str,
        request: Request,
        user: Annotated[AuthenticatedUser, Depends(require_permission("case:read"))],
        db: Annotated[Database, Depends(get_db)],
    ) -> dict[str, Any]:
        upload = db.get_file_upload(upload_id)
        if not upload:
            raise HTTPException(status_code=404, detail="Uploaded file not found")
        if not can_access_file_upload(user, upload, db):
            raise HTTPException(status_code=403, detail="File access denied")
        write_audit_log(
            db,
            action="FILE_UPLOAD_READ",
            resource_type="file_upload",
            resource_id=upload_id,
            status="SUCCESS",
            request=request,
            user=user,
        )
        return public_file_upload(upload)

    @app.get("/files/uploads/{upload_id}/content", tags=["cases"])
    def preview_file_upload(
        upload_id: str,
        request: Request,
        user: Annotated[AuthenticatedUser, Depends(require_permission("case:read"))],
        db: Annotated[Database, Depends(get_db)],
        settings: Annotated[Settings, Depends(get_app_settings)],
    ) -> FileResponse:
        upload = db.get_file_upload(upload_id)
        if not upload:
            raise HTTPException(status_code=404, detail="Uploaded file not found")
        if not can_access_file_upload(user, upload, db):
            raise HTTPException(status_code=403, detail="File access denied")
        file_path = resolve_upload_path(upload, settings)
        write_audit_log(
            db,
            action="FILE_UPLOAD_PREVIEW",
            resource_type="file_upload",
            resource_id=upload_id,
            status="SUCCESS",
            request=request,
            user=user,
            detail={"filename": upload.get("original_filename"), "extension": upload.get("extension")},
        )
        return FileResponse(
            file_path,
            media_type=upload_media_type(upload),
            filename=upload.get("original_filename") or file_path.name,
            headers={"Content-Disposition": f'inline; filename="{safe_upload_filename(upload.get("original_filename") or file_path.name)}"'},
        )

    @app.get("/files/uploads/{upload_id}/pages", tags=["cases"])
    def file_upload_preview_pages(
        upload_id: str,
        request: Request,
        user: Annotated[AuthenticatedUser, Depends(require_permission("case:read"))],
        db: Annotated[Database, Depends(get_db)],
        settings: Annotated[Settings, Depends(get_app_settings)],
    ) -> dict[str, Any]:
        upload = db.get_file_upload(upload_id)
        if not upload:
            raise HTTPException(status_code=404, detail="Uploaded file not found")
        if not can_access_file_upload(user, upload, db):
            raise HTTPException(status_code=403, detail="File access denied")
        if str(upload.get("extension") or "").lower() not in DOCUMENT_EXTENSIONS:
            return {"upload_id": upload_id, "page_count": 1, "visible_pages": [1], "preview_supported": file_preview_supported(upload)}
        file_path = resolve_upload_path(upload, settings)
        page_count = upload_preview_page_count(upload, file_path.read_bytes())
        visible_pages = list(range(1, min(5, page_count) + 1))
        write_audit_log(
            db,
            action="FILE_UPLOAD_PAGE_METADATA",
            resource_type="file_upload",
            resource_id=upload_id,
            status="SUCCESS",
            request=request,
            user=user,
            detail={"filename": upload.get("original_filename"), "page_count": page_count},
        )
        return {
            "upload_id": upload_id,
            "page_count": page_count,
            "visible_pages": visible_pages,
            "preview_supported": file_preview_supported(upload),
        }

    @app.get("/files/uploads/{upload_id}/render.png", tags=["cases"])
    def render_file_upload_png(
        upload_id: str,
        request: Request,
        user: Annotated[AuthenticatedUser, Depends(require_permission("case:read"))],
        db: Annotated[Database, Depends(get_db)],
        settings: Annotated[Settings, Depends(get_app_settings)],
        page: int = 1,
    ) -> Response:
        upload = db.get_file_upload(upload_id)
        if not upload:
            raise HTTPException(status_code=404, detail="Uploaded file not found")
        if not can_access_file_upload(user, upload, db):
            raise HTTPException(status_code=403, detail="File access denied")
        file_path = resolve_upload_path(upload, settings)
        extension = str(upload.get("extension") or "").lower()
        if extension not in DOCUMENT_EXTENSIONS:
            raise HTTPException(status_code=400, detail="PNG rendering is available for .pdf, .doc, and .docx uploads")
        content = file_path.read_bytes()
        png = render_upload_png(upload, content, max(1, page))
        write_audit_log(
            db,
            action="FILE_UPLOAD_RENDER_PNG",
            resource_type="file_upload",
            resource_id=upload_id,
            status="SUCCESS",
            request=request,
            user=user,
            detail={"filename": upload.get("original_filename"), "extension": extension, "page": page},
        )
        return Response(
            content=png,
            media_type="image/png",
            headers={"Content-Disposition": f'inline; filename="{safe_upload_filename(upload.get("original_filename") or upload_id)}.png"'},
        )

    @app.get("/files/uploads/{upload_id}/fir.xml", tags=["cases"])
    def download_fir_reconstruction_xml(
        upload_id: str,
        request: Request,
        user: Annotated[AuthenticatedUser, Depends(require_permission("case:read"))],
        db: Annotated[Database, Depends(get_db)],
    ) -> Response:
        upload = db.get_file_upload(upload_id)
        if not upload:
            raise HTTPException(status_code=404, detail="Uploaded file not found")
        if not can_access_file_upload(user, upload, db):
            raise HTTPException(status_code=403, detail="File access denied")
        xml = str(upload.get("fir_reconstruction_xml") or "").strip()
        if not xml:
            raise HTTPException(status_code=404, detail="Strict FIR XML reconstruction is not available for this upload")
        write_audit_log(
            db,
            action="FILE_UPLOAD_FIR_XML",
            resource_type="file_upload",
            resource_id=upload_id,
            status="SUCCESS",
            request=request,
            user=user,
            detail={"filename": upload.get("original_filename"), "status": upload.get("fir_reconstruction_status")},
        )
        filename = f"{Path(str(upload.get('original_filename') or upload_id)).stem}-fir-reconstruction.xml"
        return Response(
            content=xml.encode("utf-8"),
            media_type="application/xml; charset=utf-8",
            headers={"Content-Disposition": f'attachment; filename="{safe_upload_filename(filename)}"'},
        )

    @app.post("/intelligence/query", response_model=IntelligenceResponse, tags=["intelligence"])
    def intelligence_query(
        payload: IntelligenceQuery,
        request: Request,
        user: Annotated[AuthenticatedUser, Depends(require_permission("intelligence:query"))],
        db: Annotated[Database, Depends(get_db)],
    ) -> dict:
        response = build_intelligence_response(
            query=payload.query,
            language=payload.language,
            cases=db.list_cases(),
            user=user,
            include_sources=payload.include_sources,
            db=db,
        )
        write_audit_log(
            db,
            action="INTELLIGENCE_QUERY",
            resource_type="intelligence",
            status="SUCCESS",
            request=request,
            user=user,
            detail={"intent": response["intent"], "source_count": len(response["sources"])},
        )
        return response

    @app.get("/analytics/trends", response_model=TrendSummary, tags=["analytics"])
    def analytics_trends(
        request: Request,
        user: Annotated[AuthenticatedUser, Depends(require_permission("analytics:read"))],
        db: Annotated[Database, Depends(get_db)],
    ) -> dict:
        cases = visible_cases(user, db.list_cases())
        response = trend_summary(cases)
        write_audit_log(
            db,
            action="ANALYTICS_TRENDS",
            resource_type="analytics",
            status="SUCCESS",
            request=request,
            user=user,
            detail={"case_count": len(cases)},
        )
        return response

    @app.get("/analytics/network", response_model=NetworkGraph, tags=["analytics"])
    def analytics_network(
        request: Request,
        user: Annotated[AuthenticatedUser, Depends(require_permission("analytics:read"))],
        db: Annotated[Database, Depends(get_db)],
        focus: str | None = None,
        focus_type: str = "auto",
        hops: int = 4,
    ) -> dict:
        cases = visible_cases(user, db.list_cases())
        transactions = visible_transactions(user, db.list_financial_transactions(), cases)
        scoped_cases, scoped_transactions, focus_metadata = focused_network_scope(
            cases=cases,
            transactions=transactions,
            focus=focus,
            focus_type=focus_type,
            hops=hops,
        )
        response = network_graph(scoped_cases, scoped_transactions, user.role, focus_metadata=focus_metadata)
        write_audit_log(
            db,
            action="ANALYTICS_NETWORK",
            resource_type="analytics",
            status="SUCCESS",
            request=request,
            user=user,
            detail={
                "node_count": len(response["nodes"]),
                "link_count": len(response["links"]),
                "focus": focus_metadata.get("query") if focus_metadata else None,
                "focus_type": focus_metadata.get("focus_type") if focus_metadata else "all",
            },
        )
        return response

    @app.get("/analytics/patterns", response_model=PatternAnalyticsResponse, tags=["analytics"])
    def analytics_patterns(
        request: Request,
        user: Annotated[AuthenticatedUser, Depends(require_permission("analytics:read"))],
        db: Annotated[Database, Depends(get_db)],
    ) -> dict:
        cases = visible_cases(user, db.list_cases())
        response = discover_crime_patterns(cases)
        write_audit_log(
            db,
            action="ANALYTICS_PATTERNS",
            resource_type="analytics",
            status="SUCCESS",
            request=request,
            user=user,
            detail={"case_count": len(cases), "cluster_count": len(response["clusters"])},
        )
        return response

    @app.get("/analytics/advanced-crime", response_model=AdvancedCrimeAnalyticsResponse, tags=["analytics"])
    def analytics_advanced_crime(
        request: Request,
        user: Annotated[AuthenticatedUser, Depends(require_permission("analytics:read"))],
        db: Annotated[Database, Depends(get_db)],
    ) -> dict:
        response = advanced_crime_analytics(db)
        write_audit_log(
            db,
            action="ANALYTICS_ADVANCED_CRIME",
            resource_type="analytics",
            status="SUCCESS",
            request=request,
            user=user,
            detail={
                "imported_count": response["imported_count"],
                "heatmap_points": len(response["heatmap_points"]),
                "risk_areas": len(response["risk_areas"]),
                "anomalies": len(response["anomalies"]),
            },
        )
        return response

    @app.get("/analytics/sociological", response_model=SociologicalInsights, tags=["analytics"])
    def analytics_sociological(
        request: Request,
        user: Annotated[AuthenticatedUser, Depends(require_permission("sociological:read"))],
        db: Annotated[Database, Depends(get_db)],
    ) -> dict:
        cases = visible_cases(user, db.list_cases())
        response = sociological_insights(cases)
        write_audit_log(
            db,
            action="ANALYTICS_SOCIOLOGICAL",
            resource_type="analytics",
            status="SUCCESS",
            request=request,
            user=user,
            detail={"case_count": len(cases)},
        )
        return response

    @app.get("/profiles/suspects/{suspect_name}", response_model=SuspectProfileResponse, tags=["analytics"])
    def profile_suspect(
        suspect_name: str,
        request: Request,
        user: Annotated[AuthenticatedUser, Depends(require_permission("profile:read"))],
        db: Annotated[Database, Depends(get_db)],
    ) -> dict:
        cases = visible_cases(user, db.list_cases())
        response = suspect_profile(suspect_name, cases, user.role)
        write_audit_log(
            db,
            action="SUSPECT_PROFILE",
            resource_type="profile",
            status="SUCCESS",
            request=request,
            user=user,
            detail={"suspect_name": suspect_name, "case_count": response["named_in_case_count"]},
        )
        return response

    @app.get("/financial/analysis", response_model=FinancialAnalysisResponse, tags=["analytics"])
    def financial_analysis(
        request: Request,
        user: Annotated[AuthenticatedUser, Depends(require_permission("financial:read"))],
        db: Annotated[Database, Depends(get_db)],
    ) -> dict:
        cases = visible_cases(user, db.list_cases())
        transactions = visible_transactions(user, db.list_financial_transactions(), cases)
        response = analyze_financial_transactions(transactions, user.role)
        write_audit_log(
            db,
            action="FINANCIAL_ANALYSIS",
            resource_type="financial",
            status="SUCCESS",
            request=request,
            user=user,
            detail={"transaction_count": len(transactions), "finding_count": len(response["findings"])},
        )
        return response

    @app.post("/financial/transactions/import", tags=["analytics"])
    def import_financial_transactions(
        payload: FinancialTransactionImportRequest,
        request: Request,
        user: Annotated[AuthenticatedUser, Depends(require_permission("financial:import"))],
        db: Annotated[Database, Depends(get_db)],
    ) -> dict:
        all_cases = db.list_cases()
        accessible_cases = visible_cases(user, all_cases)
        accessible_case_ids = {case["id"] for case in accessible_cases}
        imported = 0
        for record in payload.records:
            if record.case_id:
                case = db.get_case(record.case_id)
                if not case:
                    raise HTTPException(status_code=404, detail=f"Linked case not found: {record.case_id}")
                if record.case_id not in accessible_case_ids:
                    raise HTTPException(status_code=403, detail=f"Case access denied: {record.case_id}")
            if user.role not in {"super_admin", "supervisor"} and record.district != user.district:
                raise HTTPException(status_code=403, detail=f"District access denied: {record.district}")
            db.create_financial_transaction(
                occurred_at=record.occurred_at.isoformat(),
                source_account=record.source_account,
                target_account=record.target_account,
                source_account_holder=record.source_account_holder,
                target_account_holder=record.target_account_holder,
                source_bank_name=record.source_bank_name,
                source_ifsc_code=record.source_ifsc_code,
                source_branch=record.source_branch,
                source_bank_manager_phone=record.source_bank_manager_phone,
                target_bank_name=record.target_bank_name,
                target_ifsc_code=record.target_ifsc_code,
                target_branch=record.target_branch,
                target_bank_manager_phone=record.target_bank_manager_phone,
                amount=record.amount,
                currency=record.currency,
                district=record.district,
                case_id=record.case_id,
                description=record.description,
            )
            imported += 1
        write_audit_log(
            db,
            action="FINANCIAL_TRANSACTION_IMPORT",
            resource_type="financial",
            status="SUCCESS",
            request=request,
            user=user,
            detail={"source_system": payload.source_system, "imported": imported},
        )
        return {"imported": imported, "source_system": payload.source_system}

    @app.get("/forecast/hotspots", response_model=ForecastResponse, tags=["analytics"])
    def forecast(
        request: Request,
        user: Annotated[AuthenticatedUser, Depends(require_permission("forecast:read"))],
        db: Annotated[Database, Depends(get_db)],
    ) -> dict:
        cases = visible_cases(user, db.list_cases())
        response = forecast_hotspots(cases)
        write_audit_log(
            db,
            action="FORECAST_HOTSPOTS",
            resource_type="forecast",
            status="SUCCESS",
            request=request,
            user=user,
            detail={"hotspot_count": len(response["hotspots"])},
        )
        return response

    @app.get("/cases/{case_id}", response_model=CaseRecord, tags=["cases"])
    def get_case(
        case_id: str,
        request: Request,
        user: Annotated[AuthenticatedUser, Depends(require_permission("case:read"))],
        db: Annotated[Database, Depends(get_db)],
    ) -> dict:
        case = db.get_case(case_id)
        if not case:
            raise HTTPException(status_code=404, detail="Case not found")
        if not can_access_case(user, case):
            write_audit_log(
                db,
                action="CASE_ACCESS_DENIED",
                resource_type="case",
                resource_id=case_id,
                status="DENIED",
                request=request,
                user=user,
            )
            raise HTTPException(status_code=403, detail="Case access denied")

        write_audit_log(
            db,
            action="CASE_READ",
            resource_type="case",
            resource_id=case_id,
            status="SUCCESS",
            request=request,
            user=user,
        )
        return mask_case(case, user.role)

    @app.get(
        "/decision-support/cases/{case_id}",
        response_model=DecisionSupportResponse,
        tags=["intelligence"],
    )
    def decision_support(
        case_id: str,
        request: Request,
        user: Annotated[AuthenticatedUser, Depends(require_permission("intelligence:query"))],
        db: Annotated[Database, Depends(get_db)],
    ) -> dict:
        case = db.get_case(case_id)
        if not case:
            raise HTTPException(status_code=404, detail="Case not found")
        cases = visible_cases(user, db.list_cases())
        if not any(visible["id"] == case_id for visible in cases):
            write_audit_log(
                db,
                action="DECISION_SUPPORT_DENIED",
                resource_type="case",
                resource_id=case_id,
                status="DENIED",
                request=request,
                user=user,
            )
            raise HTTPException(status_code=403, detail="Case access denied")

        response = build_decision_support(case, cases, user.role, notes=db.list_case_notes(case_id))
        write_audit_log(
            db,
            action="DECISION_SUPPORT",
            resource_type="case",
            resource_id=case_id,
            status="SUCCESS",
            request=request,
            user=user,
        )
        return response

    @app.get("/explain/cases/{case_id}", response_model=ExplanationResponse, tags=["analytics"])
    def explain_case_endpoint(
        case_id: str,
        request: Request,
        user: Annotated[AuthenticatedUser, Depends(require_permission("explain:read"))],
        db: Annotated[Database, Depends(get_db)],
    ) -> dict:
        case = db.get_case(case_id)
        if not case:
            raise HTTPException(status_code=404, detail="Case not found")
        if not can_access_case(user, case):
            write_audit_log(
                db,
                action="EXPLAIN_DENIED",
                resource_type="case",
                resource_id=case_id,
                status="DENIED",
                request=request,
                user=user,
            )
            raise HTTPException(status_code=403, detail="Case access denied")
        response = explain_case(case, db.latest_audit_hash())
        write_audit_log(
            db,
            action="EXPLAIN_CASE",
            resource_type="case",
            resource_id=case_id,
            status="SUCCESS",
            request=request,
            user=user,
        )
        return response

    @app.patch("/cases/{case_id}/status", response_model=CaseRecord, tags=["cases"])
    def update_case_status(
        case_id: str,
        payload: CaseStatusUpdate,
        request: Request,
        user: Annotated[AuthenticatedUser, Depends(require_permission("case:update"))],
        db: Annotated[Database, Depends(get_db)],
    ) -> dict:
        case = db.get_case(case_id)
        if not case:
            raise HTTPException(status_code=404, detail="Case not found")
        if not can_access_case(user, case):
            write_audit_log(
                db,
                action="CASE_UPDATE_DENIED",
                resource_type="case",
                resource_id=case_id,
                status="DENIED",
                request=request,
                user=user,
            )
            raise HTTPException(status_code=403, detail="Case access denied")

        updated = db.update_case_status(case_id, payload.status)
        write_audit_log(
            db,
            action="CASE_STATUS_UPDATE",
            resource_type="case",
            resource_id=case_id,
            status="SUCCESS",
            request=request,
            user=user,
            detail={"new_status": payload.status},
        )
        return mask_case(updated, user.role)

    @app.post("/cases/{case_id}/notes", response_model=CaseNote, tags=["cases"])
    def create_case_note(
        case_id: str,
        payload: CaseNoteCreate,
        request: Request,
        user: Annotated[AuthenticatedUser, Depends(require_permission("case:note"))],
        db: Annotated[Database, Depends(get_db)],
    ) -> dict:
        case = db.get_case(case_id)
        if not case:
            raise HTTPException(status_code=404, detail="Case not found")
        if not can_access_case(user, case):
            write_audit_log(
                db,
                action="CASE_NOTE_DENIED",
                resource_type="case",
                resource_id=case_id,
                status="DENIED",
                request=request,
                user=user,
            )
            raise HTTPException(status_code=403, detail="Case access denied")

        note = db.create_case_note(case_id, user.id, payload.body)
        write_audit_log(
            db,
            action="CASE_NOTE_CREATE",
            resource_type="case",
            resource_id=case_id,
            status="SUCCESS",
            request=request,
            user=user,
        )
        return note

    @app.get("/audit/logs", response_model=list[AuditLog], tags=["audit"])
    def audit_logs(
        user: Annotated[AuthenticatedUser, Depends(require_permission("audit:read"))],
        db: Annotated[Database, Depends(get_db)],
        limit: int = 100,
    ) -> list[dict]:
        _ = user
        return db.list_audit_logs(limit)

    @app.get("/audit/integrity", tags=["audit"])
    def audit_integrity(
        user: Annotated[AuthenticatedUser, Depends(require_permission("audit:verify"))],
        db: Annotated[Database, Depends(get_db)],
    ) -> dict:
        _ = user
        return verify_audit_chain(db)

    @app.get("/penal-codes", response_model=list[PenalCodeRecord], tags=["reference"])
    def list_penal_codes(
        user: Annotated[AuthenticatedUser, Depends(require_permission("case:read"))],
        db: Annotated[Database, Depends(get_db)],
    ) -> list[dict]:
        _ = user
        return db.list_penal_codes()

    @app.post("/crime-logs", response_model=CrimeLogRecord, tags=["cases"])
    def create_crime_log(
        payload: CrimeLogCreate,
        request: Request,
        user: Annotated[AuthenticatedUser, Depends(require_permission("case:write"))],
        db: Annotated[Database, Depends(get_db)],
    ) -> dict:
        log = db.create_crime_log(payload.model_dump(), user.id)
        write_audit_log(
            db,
            action="CRIME_LOG_CREATE",
            resource_type="crime_log",
            resource_id=log["id"],
            status="SUCCESS",
            request=request,
            user=user,
        )
        return log

    @app.get("/crime-logs", response_model=list[CrimeLogRecord], tags=["cases"])
    def list_crime_logs(
        user: Annotated[AuthenticatedUser, Depends(require_permission("case:read"))],
        db: Annotated[Database, Depends(get_db)],
        limit: int = 100,
    ) -> list[dict]:
        _ = user
        return db.list_crime_logs(limit)

    return app


app = create_app()
